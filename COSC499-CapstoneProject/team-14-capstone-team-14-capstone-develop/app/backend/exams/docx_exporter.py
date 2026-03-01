"""
DOCX Export functionality for exam variants
"""

import datetime
from io import BytesIO
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

from .models import Exam


class DOCXExporter:
    """Handles DOCX export of exam variants with UBC professional formatting"""

    @staticmethod
    def export_variants_to_docx_files(
        exam: Exam, variant_ids: Optional[List[int]] = None
    ) -> List[Tuple[str, bytes]]:
        """
        Export variants to DOCX format with UBC professional formatting.
        Returns a list of tuples: [(filename, docx_bytes), ...]
        """
        result_files = []

        # If variant_ids is None, export all variants
        if variant_ids is None:
            variant_ids = list(exam.variants.values_list("id", flat=True))

        for vid in variant_ids:
            variant = exam.variants.get(id=vid)
            buffer = BytesIO()
            doc = Document()

            # Set up document styles - UBC professional formatting
            styles = doc.styles
            normal_style = styles["Normal"]
            normal_style.font.size = Pt(12)
            normal_style.font.name = "Times New Roman"

            # Set margins to 1 inch
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # 1. UBC-STYLE HEADER AREA
            # Course code and title
            if hasattr(exam, "course") and exam.course:
                course_para = doc.add_paragraph()
                course_run = course_para.add_run(f"[{exam.course.code}] {exam.title}")
                course_run.font.size = Pt(14)
                course_run.font.bold = True
                course_run.font.name = "Times New Roman"
                course_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Course title
            if hasattr(exam, "course") and exam.course:
                course_title_para = doc.add_paragraph()
                course_title_run = course_title_para.add_run(
                    f"Course Title: {exam.course.name}"
                )
                course_title_run.font.size = Pt(12)
                course_title_run.font.name = "Times New Roman"
                course_title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Date
            exam_date = getattr(exam, "available_from", None)
            if not exam_date:
                exam_date = getattr(exam, "created_at", datetime.datetime.now())
            if isinstance(exam_date, str):
                exam_date = datetime.datetime.fromisoformat(
                    exam_date.replace("Z", "+00:00")
                )
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Date: {exam_date.strftime('%B %Y')}")
            date_run.font.size = Pt(12)
            date_run.font.name = "Times New Roman"
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Variant label
            variant_para = doc.add_paragraph()
            variant_run = variant_para.add_run(f"Variant {variant.version_label}")
            variant_run.font.size = Pt(14)
            variant_run.font.bold = True
            variant_run.font.name = "Times New Roman"
            variant_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add spacing before student info
            doc.add_paragraph()
            doc.add_paragraph()

            # 2. STUDENT INFO SECTION
            # Name and Student ID table
            table = doc.add_table(rows=2, cols=2)
            table.style = "Table Grid"

            # Name row
            name_cell = table.cell(0, 0)
            name_cell.text = "Name:"
            name_cell.paragraphs[0].runs[0].font.bold = True

            name_blank_cell = table.cell(0, 1)
            name_blank_cell.text = "_____________________________"

            # Student ID row
            id_cell = table.cell(1, 0)
            id_cell.text = "Student ID:"
            id_cell.paragraphs[0].runs[0].font.bold = True

            id_blank_cell = table.cell(1, 1)
            id_blank_cell.text = "_____________________________"

            # Signature
            doc.add_paragraph()
            sig_para = doc.add_paragraph("Signature: ____________________________")
            sig_para.runs[0].font.bold = True

            doc.add_paragraph()
            doc.add_paragraph()

            # 3. INSTRUCTIONS BLOCK (only if exam_instructions is not empty)
            if hasattr(exam, "exam_instructions") and exam.exam_instructions.strip():
                instructions_para = doc.add_paragraph()
                instructions_para.style = "Normal"

                instructions_run = instructions_para.add_run(exam.exam_instructions)
                instructions_run.font.name = "Times New Roman"
                instructions_run.font.size = Pt(12)
                instructions_run.font.italic = True

                doc.add_paragraph()

            # Time and marks info (only if both are set and non-zero)
            time_limit = getattr(exam, "time_limit", 0)
            total_points = getattr(exam, "total_points", 0)

            if time_limit and total_points and time_limit > 0 and total_points > 0:
                time_marks_para = doc.add_paragraph()
                time_marks_text = f"• You have {time_limit} minutes to write this exam. A total of {total_points} marks are available."
                time_marks_run = time_marks_para.add_run(time_marks_text)
                time_marks_run.font.name = "Times New Roman"
                time_marks_run.font.size = Pt(12)

                doc.add_paragraph()

            # 4. ACADEMIC INTEGRITY STATEMENT
            if (
                hasattr(exam, "include_academic_integrity")
                and exam.include_academic_integrity
            ):
                integrity_para = doc.add_paragraph()
                integrity_run = integrity_para.add_run("Academic Integrity Statement:")
                integrity_run.font.bold = True
                integrity_run.font.size = Pt(12)
                integrity_run.font.name = "Times New Roman"

                doc.add_paragraph()

                # Use custom statement if provided, otherwise use default
                custom_statement = getattr(exam, "academic_integrity_statement", "")
                if custom_statement and custom_statement.strip():
                    statement = custom_statement
                else:
                    statement = "I confirm that I will not give or receive unauthorized aid on this examination. I understand that academic dishonesty includes, but is not limited to, cheating, plagiarism, and the unauthorized use of materials or devices."

                statement_para = doc.add_paragraph()
                statement_run = statement_para.add_run(statement)
                statement_run.font.name = "Times New Roman"
                statement_run.font.size = Pt(12)

                doc.add_paragraph()
                doc.add_paragraph()

            # 5. QUESTIONS SECTION
            variant_questions = variant.variantquestion_set.select_related(
                "question", "section"
            ).order_by("order")

            # Group questions by section if sections exist
            if hasattr(exam, "sections") and exam.sections.exists():
                # Group questions by section using the stored section information
                questions_by_section = {}
                questions_without_section = []

                for vq in variant_questions:
                    section = vq.section
                    if section:
                        if section.id not in questions_by_section:
                            questions_by_section[section.id] = {
                                "section": section,
                                "questions": [],
                            }
                        questions_by_section[section.id]["questions"].append(vq)
                    else:
                        # Questions without sections go to a separate list
                        questions_without_section.append(vq)

                # Sort sections by order and render questions
                question_number = 1

                for section_id in sorted(
                    questions_by_section.keys(),
                    key=lambda x: questions_by_section[x]["section"].order,
                ):
                    section_data = questions_by_section[section_id]
                    section = section_data["section"]

                    # Section header
                    # Check if title already starts with "Section" to avoid duplication
                    if section.title.startswith("Section "):
                        section_para = doc.add_paragraph(
                            f"{section.title} — Multiple Choice"
                        )
                    else:
                        section_para = doc.add_paragraph(
                            f"Section {section.title} — Multiple Choice"
                        )
                    section_para.runs[0].font.bold = True
                    section_para.runs[0].font.size = Pt(12)
                    section_para.runs[0].font.name = "Times New Roman"

                    # Add section point information based on marking scheme
                    if hasattr(exam, "marking_scheme") and exam.marking_scheme:
                        section_weighting = exam.marking_scheme.get(
                            "sectionWeighting", {}
                        )
                        # Try both string and number keys for section ID
                        section_id_str = str(section.id)
                        section_id_num = section.id
                        weight = section_weighting.get(
                            section_id_str, section_weighting.get(section_id_num, 1.0)
                        )

                        # Debug logging
                        print(f"🔍 DOCX Export - Section {section.id}:")
                        print(f"   Section weighting: {section_weighting}")
                        print(f"   Section ID (str): {section_id_str}")
                        print(f"   Section ID (num): {section_id_num}")
                        print(f"   Weight: {weight}")

                        points_info = doc.add_paragraph(
                            f"Each question in this section is worth {weight} point{'s' if weight != 1 else ''}."
                        )
                        points_info.runs[0].font.name = "Times New Roman"
                        points_info.runs[0].font.size = Pt(11)
                        points_info.runs[0].font.italic = True
                        doc.add_paragraph()  # Empty line

                    # Section instructions if provided
                    if section.instructions:
                        section_instructions_para = doc.add_paragraph(
                            section.instructions
                        )
                        section_instructions_para.runs[0].font.name = "Times New Roman"
                        section_instructions_para.runs[0].font.size = Pt(12)
                        section_instructions_para.runs[0].font.italic = True
                        doc.add_paragraph()  # Empty line

                    # Questions in this section
                    for vq in section_data["questions"]:
                        question = vq.question

                        # Add question (without individual marks)
                        q_para = doc.add_paragraph()
                        q_run = q_para.add_run(f"{question_number}. {question.prompt}")
                        q_run.font.bold = True
                        q_run.font.size = Pt(12)
                        q_run.font.name = "Times New Roman"

                        # Get choices - use randomized if available
                        if vq.randomized_choices:
                            choices = vq.randomized_choices
                        else:
                            choices = question.choices

                        # Add choices with proper indentation
                        if isinstance(choices, dict):
                            for label, text in choices.items():
                                c_para = doc.add_paragraph(f"{label}. {text}")
                                c_para.paragraph_format.left_indent = Inches(0.5)
                                c_para.runs[0].font.name = "Times New Roman"
                                c_para.runs[0].font.size = Pt(12)

                        # Space between questions
                        doc.add_paragraph()
                        question_number += 1

                # Handle questions without sections
                if questions_without_section:
                    # Section header for questions without sections
                    section_para = doc.add_paragraph(
                        "Additional Questions — Multiple Choice"
                    )
                    section_para.runs[0].font.bold = True
                    section_para.runs[0].font.size = Pt(12)
                    section_para.runs[0].font.name = "Times New Roman"

                    for vq in questions_without_section:
                        question = vq.question

                        # Add question (without individual marks)
                        q_para = doc.add_paragraph()
                        q_run = q_para.add_run(f"{question_number}. {question.prompt}")
                        q_run.font.bold = True
                        q_run.font.size = Pt(12)
                        q_run.font.name = "Times New Roman"

                        # Get choices - use randomized if available
                        if vq.randomized_choices:
                            choices = vq.randomized_choices
                        else:
                            choices = question.choices

                        # Add choices with proper indentation
                        if isinstance(choices, dict):
                            for label, text in choices.items():
                                c_para = doc.add_paragraph(f"{label}. {text}")
                                c_para.paragraph_format.left_indent = Inches(0.5)
                                c_para.runs[0].font.name = "Times New Roman"
                                c_para.runs[0].font.size = Pt(12)

                        # Space between questions
                        doc.add_paragraph()
                        question_number += 1
            else:
                # No sections - just add questions sequentially
                for idx, vq in enumerate(variant_questions, start=1):
                    question = vq.question

                    # Add question (without individual marks)
                    q_para = doc.add_paragraph()
                    q_run = q_para.add_run(f"{idx}. {question.prompt}")
                    q_run.font.bold = True
                    q_run.font.size = Pt(12)
                    q_run.font.name = "Times New Roman"

                    # Get choices - use randomized if available
                    if vq.randomized_choices:
                        choices = vq.randomized_choices
                    else:
                        choices = question.choices

                    # Add choices with proper indentation
                    if isinstance(choices, dict):
                        for label, text in choices.items():
                            c_para = doc.add_paragraph(f"{label}. {text}")
                            c_para.paragraph_format.left_indent = Inches(0.5)
                            c_para.runs[0].font.name = "Times New Roman"
                            c_para.runs[0].font.size = Pt(12)

                    # Space between questions
                    doc.add_paragraph()

            # 6. FOOTER (if provided)
            if hasattr(exam, "footer_text") and exam.footer_text:
                # Add a page break to ensure footer is on a new page
                doc.add_page_break()
                footer_para = doc.add_paragraph(exam.footer_text)
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                footer_para.runs[0].font.italic = True
                footer_para.runs[0].font.size = Pt(10)
                footer_para.runs[0].font.name = "Times New Roman"

            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)

            filename = f"{exam.title}_Variant_{variant.version_label}.docx"
            result_files.append((filename, buffer.getvalue()))

        return result_files
