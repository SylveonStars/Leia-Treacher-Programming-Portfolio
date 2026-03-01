"""
Services for exam operations - separates business logic from models and views
"""

import csv
import io
import logging
import random
from typing import Any, Dict, List, Optional, Tuple
import zipfile

from questions.models import Question

from .docx_exporter import DOCXExporter
from .exam_algorithms import (
    generate_section_based_variants_reuse,
    generate_section_based_variants_unique,
)
from .models import Exam, ExamExportHistory, Variant
from .pdf_exporter import PDFExporter

logger = logging.getLogger(__name__)


class ExamValidationService:
    """Service for validating exam configuration"""

    @staticmethod
    def validate_exam_config(exam: Exam) -> List[str]:
        """Validate exam configuration and return list of errors"""
        errors = []

        # Validate difficulty distribution
        total = (
            int(getattr(exam, "easy_percentage", 0))
            + int(getattr(exam, "medium_percentage", 0))
            + int(getattr(exam, "hard_percentage", 0))
        )
        # Allow sum <= 100 (remainder is Unknown)
        if total > 100 or total < 0:
            errors.append(
                (
                    "Difficulty distribution must add up to 100% or less "
                    "(remainder will be filled with questions of unknown difficulty)"
                )
            )

        # Validate question count from sections
        available_questions = []
        for section in exam.sections.all():
            for question_bank in section.question_banks.all():
                bank_questions = Question.objects.filter(bank=question_bank)
                available_questions.extend(bank_questions)

        # Remove duplicates
        available_question_ids = set(q.id for q in available_questions)
        total_available = len(available_question_ids)

        if total_available < exam.questions_per_variant:
            errors.append(
                (
                    f"Not enough questions in selected banks. Need {exam.questions_per_variant}, "
                    f"have {total_available}"
                )
            )

        # Validate mandatory questions
        mandatory_count = getattr(exam, "mandatory_questions").all().count()
        if exam.questions_per_variant < mandatory_count:
            errors.append(
                (
                    f"Number of questions per variant ({exam.questions_per_variant}) "
                    f"is less than mandatory questions ({mandatory_count})"
                )
            )

        # Validate variant count
        if exam.num_variants < 1 or exam.num_variants > 5:
            errors.append("Number of variants must be between 1 and 5")

        return errors

    @staticmethod
    def validate_variant_generation(exam: Exam) -> List[str]:
        """Validate if variants can be generated"""
        errors = ExamValidationService.validate_exam_config(exam)

        if errors:
            return errors

        # Check if we have enough questions for unique mode
        if not exam.allow_reuse:
            mandatory_count = getattr(exam, "mandatory_questions").all().count()
            non_mandatory_needed = exam.questions_per_variant - mandatory_count
            total_non_mandatory_needed = non_mandatory_needed * exam.num_variants

            # Get all available non-mandatory questions
            mandatory_question_ids = set(
                getattr(exam, "mandatory_questions").values_list("id", flat=True)
            )
            all_question_ids = set()

            for section in exam.sections.all():
                for question_bank in section.question_banks.all():
                    bank_question_ids = Question.objects.filter(
                        bank=question_bank
                    ).values_list("id", flat=True)
                    all_question_ids.update(bank_question_ids)

            available_non_mandatory = all_question_ids - mandatory_question_ids

            if len(available_non_mandatory) < total_non_mandatory_needed:
                errors.append(
                    f"Not enough unique questions for {exam.num_variants} variants. "
                    f"Need {total_non_mandatory_needed} unique non-mandatory questions, "
                    f"have {len(available_non_mandatory)}"
                )

        return errors


class VariantGenerationService:
    """Service for generating exam variants"""

    @staticmethod
    def _apply_anti_cheating_measures(variants: List[Variant]) -> None:
        """Apply anti-cheating measures to variants"""
        if len(variants) < 2:
            return

        # Get all variant questions for analysis
        variant_questions = []
        for variant in variants:
            vqs = list(
                variant.variantquestion_set.select_related("question").order_by("order")
            )
            variant_questions.append(vqs)

        # Apply simple anti-cheating optimization
        VariantGenerationService._simple_anti_cheating_optimization(
            variants, variant_questions
        )

    @staticmethod
    def _simple_anti_cheating_optimization(
        variants: List[Variant], variant_questions: List[List]
    ) -> None:
        """Enhanced anti-cheating optimization with mandatory question shuffling"""
        if len(variants) < 2:
            return

        # First, try to shuffle mandatory questions to different positions
        VariantGenerationService._shuffle_mandatory_questions(
            variants, variant_questions
        )

        # Then apply standard anti-cheating measures
        for position in range(len(variant_questions[0])):
            correct_answers = []
            for vqs in variant_questions:
                if position < len(vqs):
                    vq = vqs[position]
                    correct_answer = VariantGenerationService._get_correct_answer(vq)
                    correct_answers.append(correct_answer)

            # If all correct answers are the same, try to fix
            if len(set(correct_answers)) == 1 and len(correct_answers) > 1:
                VariantGenerationService._fix_position_simple(
                    variants, variant_questions, position
                )

    @staticmethod
    def _shuffle_mandatory_questions(
        variants: List[Variant], variant_questions: List[List]
    ) -> None:
        """Shuffle mandatory questions to different positions across variants"""
        if len(variants) < 2:
            return

        # Find mandatory questions (questions that appear in all variants)
        all_questions = set()
        for vqs in variant_questions:
            for vq in vqs:
                all_questions.add(vq.question.id)

        # Find questions that appear in all variants (likely mandatory)
        mandatory_questions = set()
        for question_id in all_questions:
            appears_in_all = True
            for vqs in variant_questions:
                if not any(vq.question.id == question_id for vq in vqs):
                    appears_in_all = False
                    break
            if appears_in_all:
                mandatory_questions.add(question_id)

        if not mandatory_questions:
            return

        # Try to shuffle mandatory questions to different positions
        for variant_idx in range(len(variants)):
            vqs = variant_questions[variant_idx]

            # Find mandatory questions in this variant
            mandatory_positions = []
            for pos, vq in enumerate(vqs):
                if vq.question.id in mandatory_questions:
                    mandatory_positions.append(pos)

            if len(mandatory_positions) > 1:
                # Shuffle mandatory question positions within this variant
                random.shuffle(mandatory_positions)

                # Create a new order for the questions
                # First, get all questions in their current order
                list(vqs)

                # Create new positions for mandatory questions
                mandatory_vqs = [vqs[pos] for pos in mandatory_positions]
                non_mandatory_vqs = [
                    vq for pos, vq in enumerate(vqs) if pos not in mandatory_positions
                ]

                # Create new order: place mandatory questions at shuffled positions
                new_order = []
                mandatory_idx = 0
                non_mandatory_idx = 0

                for i in range(len(vqs)):
                    if i in mandatory_positions:
                        # Place mandatory question
                        new_order.append(mandatory_vqs[mandatory_idx])
                        mandatory_idx += 1
                    else:
                        # Place non-mandatory question
                        new_order.append(non_mandatory_vqs[non_mandatory_idx])
                        non_mandatory_idx += 1

                # Update the database with new order
                for i, vq in enumerate(new_order):
                    vq.order = i
                    vq.save()

    @staticmethod
    def _fix_position_simple(
        variants: List[Variant], variant_questions: List[List], position: int
    ) -> None:
        """Try to fix same correct answers in same position"""
        if len(variants) < 2:
            return

        # Try swapping questions between variants
        for i in range(len(variants) - 1):
            for j in range(i + 1, len(variants)):
                if position < len(variant_questions[i]) and position < len(
                    variant_questions[j]
                ):
                    VariantGenerationService._try_swap_to_fix_vulnerability(
                        variants, variant_questions, position, i, j
                    )

    @staticmethod
    def _try_swap_to_fix_vulnerability(
        variants: List[Variant],
        variant_questions: List[List],
        position: int,
        v1_idx: int,
        v2_idx: int,
    ) -> None:
        """Try swapping questions to fix vulnerability"""
        if position >= len(variant_questions[v1_idx]) or position >= len(
            variant_questions[v2_idx]
        ):
            return

        vq1 = variant_questions[v1_idx][position]
        vq2 = variant_questions[v2_idx][position]

        # Check if swapping would help
        correct1 = VariantGenerationService._get_correct_answer(vq1)
        correct2 = VariantGenerationService._get_correct_answer(vq2)

        if correct1 != correct2:
            # Swap the questions
            VariantGenerationService._swap_variant_questions(
                variants[v1_idx], variants[v2_idx], position, position
            )

    @staticmethod
    def _get_correct_answer(vq) -> tuple:
        """Get correct answer for a variant question"""
        if vq.randomized_correct_answer:
            return tuple(sorted(vq.randomized_correct_answer))
        else:
            return tuple(sorted(vq.question.correct_answer))

    @staticmethod
    def _swap_variant_questions(
        variant1: Variant, variant2: Variant, pos1: int, pos2: int
    ) -> None:
        """Swap questions between variants"""
        vq1 = variant1.variantquestion_set.filter(order=pos1).first()
        vq2 = variant2.variantquestion_set.filter(order=pos2).first()

        if vq1 and vq2:
            # Swap the questions
            temp_question = vq1.question
            temp_randomized_choices = vq1.randomized_choices
            temp_randomized_correct = vq1.randomized_correct_answer

            vq1.question = vq2.question
            vq1.randomized_choices = vq2.randomized_choices
            vq1.randomized_correct_answer = vq2.randomized_correct_answer

            vq2.question = temp_question
            vq2.randomized_choices = temp_randomized_choices
            vq2.randomized_correct_answer = temp_randomized_correct

            vq1.save()
            vq2.save()

    @staticmethod
    def _assess_final_cheating_risk(
        variants: List[Variant], variant_questions: List[List]
    ) -> None:
        """Assess final cheating risk after all measures"""
        if len(variants) < 2:
            return

        risk_score = 0
        total_positions = len(variant_questions[0]) if variant_questions else 0

        for position in range(total_positions):
            correct_answers = []
            for vqs in variant_questions:
                if position < len(vqs):
                    vq = vqs[position]
                    correct_answer = VariantGenerationService._get_correct_answer(vq)
                    correct_answers.append(correct_answer)

            # Count unique correct answers at this position
            unique_answers = len(set(correct_answers))
            if unique_answers == 1:
                risk_score += 1

        # Calculate risk percentage
        if total_positions > 0:
            risk_percentage = (risk_score / total_positions) * 100
            logger.info(f"Final cheating risk: {risk_percentage:.1f}%")

    @staticmethod
    def calculate_cheating_risk(variants: List[Variant]) -> float:
        """Calculate cheating risk score for variants"""
        if len(variants) < 2:
            return 0.0

        risk_score = 0
        total_positions = 0

        for variant in variants:
            vqs = variant.variantquestion_set.select_related("question").order_by(
                "order"
            )
            total_positions = len(vqs)
            break

        for position in range(total_positions):
            correct_answers = []
            for variant in variants:
                vq = variant.variantquestion_set.filter(order=position).first()
                if vq:
                    correct_answer = VariantGenerationService._get_correct_answer(vq)
                    correct_answers.append(correct_answer)

            # Count unique correct answers at this position
            unique_answers = len(set(correct_answers))
            if unique_answers == 1:
                risk_score += 1

        # Calculate risk percentage
        if total_positions > 0:
            return (risk_score / total_positions) * 100
        return 0.0

    @staticmethod
    def generate_variants(exam: Exam) -> dict:
        """Generate variants for an exam"""
        import logging

        logger = logging.getLogger(__name__)

        def safe_int(val, default=0):
            try:
                return int(val) if val is not None else default
            except (ValueError, TypeError):
                return default

        # Don't delete any existing variants - just create new ones
        existing_variants = exam.variants.all()
        logger.info(
            f"Found {existing_variants.count()} existing variants for exam {exam.id}"
        )
        logger.info("Creating new variants without deleting existing ones")

        # Get exam parameters
        num_variants = min(safe_int(exam.num_variants, 1), 5)
        questions_per_variant = safe_int(exam.questions_per_variant, 1)
        reuse_mode = getattr(exam, "allow_reuse", False)

        logger.info("Exam parameters:")
        logger.info(f"  - num_variants: {num_variants}")
        logger.info(f"  - questions_per_variant: {questions_per_variant}")
        logger.info(f"  - reuse_mode: {reuse_mode}")

        # Get mandatory questions
        mandatory_questions = list(getattr(exam, "mandatory_questions").all())
        logger.info(f"Mandatory questions: {len(mandatory_questions)}")

        # Get questions from sections if they exist
        if hasattr(exam, "sections") and exam.sections.exists():
            logger.info("Using section-based variant generation")
            # Collect questions by section
            section_questions = {}
            for section in exam.sections.all():
                section_questions[section.id] = {
                    "section": section,
                    "questions": [],
                    "max_questions": section.configured_question_count or 5,
                }

                for question_bank in section.question_banks.all():
                    bank_questions = Question.objects.filter(bank=question_bank)
                    section_questions[section.id]["questions"].extend(bank_questions)

            # Create a mapping of questions to their original sections BEFORE removing mandatory questions
            question_to_section = {}
            for section_id, section_data in section_questions.items():
                section = section_data["section"]
                for question in section_data["questions"]:
                    question_to_section[question.id] = section

            # Remove mandatory questions from section questions
            for section_id, section_data in section_questions.items():
                section_question_ids = set(q.id for q in section_data["questions"])
                mandatory_question_ids = set(q.id for q in mandatory_questions)
                section_question_ids -= mandatory_question_ids
                section_data["questions"] = list(
                    Question.objects.filter(id__in=section_question_ids)
                )

            # Calculate proportional distribution
            questions_per_variant = safe_int(exam.questions_per_variant, 1)
            mandatory_count = len(mandatory_questions)
            non_mandatory_per_variant = max(0, questions_per_variant - mandatory_count)

            # Filter out sections without question banks
            sections_with_questions = {}
            for section_id, section_data in section_questions.items():
                if len(section_data["questions"]) > 0:
                    sections_with_questions[section_id] = section_data

            # Calculate section proportions only for sections with questions
            section_allocation = {}
            if sections_with_questions:
                total_max_questions = sum(
                    section_data["max_questions"]
                    for section_data in sections_with_questions.values()
                )

                if total_max_questions > 0:
                    for section_id, section_data in sections_with_questions.items():
                        proportion = section_data["max_questions"] / total_max_questions
                        section_allocation[section_id] = max(
                            1, round(non_mandatory_per_variant * proportion)
                        )
                else:
                    # Fallback: distribute evenly among sections with questions
                    num_sections = len(sections_with_questions)
                    if num_sections > 0:
                        questions_per_section = max(
                            1, non_mandatory_per_variant // num_sections
                        )
                        for section_id in sections_with_questions.keys():
                            section_allocation[section_id] = questions_per_section

                # Validate section allocation
                total_allocated = sum(section_allocation.values())
                if total_allocated != non_mandatory_per_variant:
                    # Adjust to match exactly
                    diff = non_mandatory_per_variant - total_allocated
                    if diff > 0:
                        # Add to largest section
                        largest_section = max(
                            section_allocation.items(), key=lambda x: x[1]
                        )[0]
                        section_allocation[largest_section] += diff
                    elif diff < 0:
                        # Subtract from largest section
                        largest_section = max(
                            section_allocation.items(), key=lambda x: x[1]
                        )[0]
                        section_allocation[largest_section] = max(
                            1, section_allocation[largest_section] + diff
                        )
            else:
                # No sections with questions - this should not happen in normal flow
                logger.warning("No sections with questions found for allocation")
                section_allocation = {}

            logger.info(f"Section allocation: {section_allocation}")
            logger.info(
                f"Non-mandatory questions per variant: {non_mandatory_per_variant}"
            )

            # Flatten available questions for backward compatibility
            available_questions = []
            for section_data in sections_with_questions.values():
                available_questions.extend(section_data["questions"])

            # Log the question distribution for debugging
            logger.info(f"Mandatory questions: {len(mandatory_questions)}")
            logger.info(
                f"Available non-mandatory questions: {len(available_questions)}"
            )
            logger.info(f"Questions per variant: {questions_per_variant}")
            logger.info(f"Mandatory questions per variant: {len(mandatory_questions)}")
            logger.info(
                f"Remaining questions to distribute: {questions_per_variant - len(mandatory_questions)}"
            )
            variants = []
            # Use local variables for all further logic
            n_variants = min(num_variants, 5)
            q_per_variant = questions_per_variant
            n_mandatory = len(mandatory_questions)
            n_non_mandatory = max(0, q_per_variant - n_mandatory)
            if exam.allow_reuse:
                logger.info("Using reuse mode for section-based generation")
                variants, warning, distribution_info = (
                    VariantGenerationService._generate_section_based_variants(
                        exam,
                        mandatory_questions,
                        section_questions,
                        section_allocation,
                        n_variants,
                        n_non_mandatory,
                        reuse_mode=True,
                        question_to_section=question_to_section,
                    )
                )
                if warning:
                    logger.warning(warning)
                if distribution_info.get("auto_balance_disabled", False):
                    logger.warning(
                        "Auto-balance disabled for reuse mode. "
                        "Some difficulties might not be distributed fairly."
                    )
                    # Note: Auto-balance disabled but variants already generated successfully
                    # No need to regenerate - just log the warning
                    setattr(exam, "easy_percentage", 0)
                    setattr(exam, "medium_percentage", 0)
                    setattr(exam, "hard_percentage", 0)
                    exam.save()
            else:
                logger.info("Using unique mode for section-based generation")
                variants, warning, distribution_info = (
                    VariantGenerationService._generate_section_based_variants(
                        exam,
                        mandatory_questions,
                        section_questions,
                        section_allocation,
                        n_variants,
                        n_non_mandatory,
                        reuse_mode=False,
                        question_to_section=question_to_section,
                    )
                )
                if warning:
                    logger.warning(warning)

        else:
            logger.info(
                "Using flat question selection (no sections) - converting to section-based approach"
            )
            # No sections - convert to section-based approach for compatibility
            # Get all available questions from exam's question banks
            available_questions = list(exam.questions.all())

            # Remove mandatory questions
            mandatory_question_ids = set(q.id for q in mandatory_questions)
            available_questions = [
                q for q in available_questions if q.id not in mandatory_question_ids
            ]

            # Create a virtual section for all questions
            from .models import ExamSection

            virtual_section = ExamSection(
                exam=exam,
                title="All Questions",
                instructions="",
                order=1,
                configured_question_count=len(available_questions),
            )

            # Create section-based data structure
            section_questions = {
                "virtual": {
                    "section": virtual_section,
                    "questions": available_questions,
                    "max_questions": len(available_questions),
                }
            }

            # Create section allocation
            non_mandatory_per_variant = max(
                0, questions_per_variant - len(mandatory_questions)
            )
            section_allocation = {"virtual": non_mandatory_per_variant}

            # Create question to section mapping
            question_to_section = {q.id: virtual_section for q in available_questions}

            # Log the question distribution for debugging
            logger.info(f"Mandatory questions: {len(mandatory_questions)}")
            logger.info(
                f"Available non-mandatory questions: {len(available_questions)}"
            )
            logger.info(f"Questions per variant: {questions_per_variant}")
            logger.info(f"Mandatory questions per variant: {len(mandatory_questions)}")
            logger.info(
                f"Remaining questions to distribute: {questions_per_variant - len(mandatory_questions)}"
            )

            if reuse_mode:
                logger.info(
                    "Using reuse mode for section-based generation (converted from flat)"
                )
                variants, warning, distribution_info = (
                    VariantGenerationService._generate_section_based_variants(
                        exam,
                        mandatory_questions,
                        section_questions,
                        section_allocation,
                        num_variants,
                        non_mandatory_per_variant,
                        reuse_mode=True,
                        question_to_section=question_to_section,
                    )
                )
                if warning:
                    logger.warning(warning)
            else:
                logger.info(
                    "Using unique mode for section-based generation (converted from flat)"
                )
                variants, warning, distribution_info = (
                    VariantGenerationService._generate_section_based_variants(
                        exam,
                        mandatory_questions,
                        section_questions,
                        section_allocation,
                        num_variants,
                        non_mandatory_per_variant,
                        reuse_mode=False,
                        question_to_section=question_to_section,
                    )
                )
                if warning:
                    logger.warning(warning)

        logger.info(f"Generated {len(variants)} variants")

        # Apply anti-cheating measures
        VariantGenerationService._apply_anti_cheating_measures(variants)

        # Always randomize choices (choice randomization is always enabled)
        VariantGenerationService._randomize_choices(variants)

        # Calculate final cheating risk
        final_risk = VariantGenerationService.calculate_cheating_risk(variants)
        logger.info(f"Final cheating risk: {final_risk:.1f}%")

        return {
            "variants_created": len(variants),
            "cheating_risk": final_risk,
            "distribution_info": distribution_info,
        }

    @staticmethod
    def _generate_section_based_variants(
        exam: Exam,
        mandatory_questions: List[Question],
        section_questions: Dict,
        section_allocation: Dict,
        num_variants: int,
        num_non_mandatory: int,
        reuse_mode: bool = False,
        question_to_section: Dict = None,
    ) -> Tuple[List[Variant], Optional[str], dict]:
        """Generate variants with section-based question allocation (delegated)"""
        if question_to_section is None:
            question_to_section = {}
            for section_id, section_data in section_questions.items():
                section = section_data["section"]
                for question in section_data["questions"]:
                    question_to_section[question.id] = section
        if reuse_mode:
            return generate_section_based_variants_reuse(
                exam,
                mandatory_questions,
                section_questions,
                section_allocation,
                num_variants,
                num_non_mandatory,
                question_to_section,
            )
        else:
            return generate_section_based_variants_unique(
                exam,
                mandatory_questions,
                section_questions,
                section_allocation,
                num_variants,
                num_non_mandatory,
                question_to_section,
            )

    @staticmethod
    def _randomize_choices(variants: List[Variant]) -> None:
        """Randomize choice order for all variants"""
        for variant in variants:
            vqs = variant.variantquestion_set.select_related("question").all()
            VariantGenerationService._randomize_choices_fallback(vqs)

    @staticmethod
    def _randomize_choices_fallback(vqs: List) -> None:
        """Randomize choices for variant questions"""
        for vq in vqs:
            if not vq.question.choices:
                continue

            choices = vq.question.choices
            if not isinstance(choices, dict):
                continue

            # Create randomized choices
            choice_items = list(choices.items())
            random.shuffle(choice_items)

            # Create new randomized choices dict
            randomized_choices = {}
            labels = ["A", "B", "C", "D", "E"]

            for i, (_, text) in enumerate(choice_items):
                if i < len(labels):
                    randomized_choices[labels[i]] = text

            # Update correct answer
            original_correct = vq.question.correct_answer
            if isinstance(original_correct, list) and len(original_correct) > 0:
                # Find the original correct answer text
                original_correct = vq.question.correct_answer
                new_correct = []
                for ans in original_correct:
                    orig_text = choices.get(ans, "")
                    for label, text in randomized_choices.items():
                        if text == orig_text:
                            new_correct.append(label)
                            break
                vq.randomized_correct_answer = new_correct

            vq.randomized_choices = randomized_choices
            vq.save()


class ExamExportService:
    """Service for exporting exam variants"""

    @staticmethod
    def generate_answer_key_csv(
        exam: Exam, variant_ids: Optional[List[int]] = None
    ) -> str:
        """Generate CSV answer key for variants"""
        output = io.StringIO()
        writer = csv.writer(output)

        # If variant_ids is None, export all variants
        if variant_ids is None:
            variant_ids = list(exam.variants.values_list("id", flat=True))

        # Write header
        writer.writerow(["Variant", "Question", "Correct Answer"])

        for vid in variant_ids:
            variant = exam.variants.get(id=vid)
            if not variant:
                continue

            vqs = variant.variantquestion_set.select_related("question").order_by(
                "order"
            )

            for vq in vqs:
                question = vq.question

                # Get correct answer
                if vq.randomized_correct_answer:
                    correct_answer = vq.randomized_correct_answer
                else:
                    correct_answer = question.correct_answer

                # Format correct answer
                if isinstance(correct_answer, list):
                    correct_answer_str = ", ".join(str(item) for item in correct_answer)
                else:
                    correct_answer_str = str(correct_answer)

                writer.writerow(
                    [variant.version_label, question.prompt, correct_answer_str]
                )

        return output.getvalue()

    @staticmethod
    def create_export_history(
        exam: Exam, user, export_format: str, variant_ids: Optional[List[int]] = None
    ) -> ExamExportHistory:
        """Create export history record"""
        if variant_ids is None:
            variant_ids = list(exam.variants.values_list("id", flat=True))

        variants_exported = exam.variants.filter(id__in=variant_ids)

        history = ExamExportHistory.objects.create(
            exam=exam,
            exported_by=user,
            export_format=export_format,
        )

        # Use .set() for many-to-many relationships
        history.variants_exported.set(variants_exported)

        return history

    @staticmethod
    def generate_answer_key_pdf(
        exam: Exam, variant_ids: Optional[List[int]] = None
    ) -> bytes:
        """Generate PDF answer key for variants"""
        # For now, return a simple text-based PDF
        from io import BytesIO

        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)

        # Add title
        c.setFont("Times-Bold", 16)
        c.drawString(100, 750, f"Answer Key: {exam.title}")

        # If variant_ids is None, export all variants
        if variant_ids is None:
            variant_ids = list(exam.variants.values_list("id", flat=True))

        y_position = 700
        for vid in variant_ids:
            variant = exam.variants.get(id=vid)
            if not variant:
                continue

            # Add variant header
            c.setFont("Times-Bold", 14)
            c.drawString(100, y_position, f"Variant {variant.version_label}")
            y_position -= 30

            # Add questions and answers
            c.setFont("Times-Roman", 12)
            vqs = variant.variantquestion_set.select_related("question").order_by(
                "order"
            )

            for i, vq in enumerate(vqs, 1):
                question = vq.question

                # Get correct answer
                if vq.randomized_correct_answer:
                    correct_answer = vq.randomized_correct_answer
                else:
                    correct_answer = question.correct_answer

                # Format correct answer
                if isinstance(correct_answer, list):
                    correct_answer_str = ", ".join(correct_answer)
                else:
                    correct_answer_str = str(correct_answer)

                # Add question and answer
                question_text = f"{i}. {question.prompt}"
                answer_text = f"Answer: {correct_answer_str}"

                # Wrap text if needed
                if len(question_text) > 80:
                    question_text = question_text[:77] + "..."

                c.drawString(120, y_position, question_text)
                y_position -= 20
                c.drawString(140, y_position, answer_text)
                y_position -= 30

                if y_position < 100:
                    c.showPage()
                    y_position = 750

        c.save()
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def generate_answer_key_docx(
        exam: Exam, variant_ids: Optional[List[int]] = None
    ) -> bytes:
        """Generate DOCX answer key for variants"""
        from io import BytesIO

        from docx import Document

        doc = Document()

        # If variant_ids is None, export all variants
        if variant_ids is None:
            variant_ids = list(exam.variants.values_list("id", flat=True))

        for vid in variant_ids:
            variant = exam.variants.get(id=vid)
            if not variant:
                continue

            # Add variant header
            doc.add_heading(f"Variant {variant.version_label}", level=1)

            # Add questions and answers
            vqs = variant.variantquestion_set.select_related("question").order_by(
                "order"
            )

            for i, vq in enumerate(vqs, 1):
                question = vq.question

                # Get correct answer
                if vq.randomized_correct_answer:
                    correct_answer = vq.randomized_correct_answer
                else:
                    correct_answer = question.correct_answer

                # Format correct answer
                if isinstance(correct_answer, list):
                    correct_answer_str = ", ".join(correct_answer)
                else:
                    correct_answer_str = str(correct_answer)

                # Add question and answer
                question_text = f"{i}. {question.prompt}"
                answer_text = f"Answer: {correct_answer_str}"

                doc.add_paragraph(question_text)
                doc.add_paragraph(answer_text)
                doc.add_paragraph()  # Empty line for spacing

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def generate_exam_csv(exam: Exam, variant_ids: Optional[List[int]] = None) -> str:
        """Generate CSV export of exam variants"""
        output = io.StringIO()
        writer = csv.writer(output)

        # If variant_ids is None, export all variants
        if variant_ids is None:
            variant_ids = list(exam.variants.values_list("id", flat=True))

        # Write header
        writer.writerow(["Variant", "Question", "Choices", "Correct Answer"])

        for vid in variant_ids:
            variant = exam.variants.get(id=vid)
            if not variant:
                continue

            vqs = variant.variantquestion_set.select_related("question").order_by(
                "order"
            )

            for vq in vqs:
                question = vq.question

                # Get choices
                if vq.randomized_choices:
                    choices = vq.randomized_choices
                else:
                    choices = question.choices

                # Format choices
                if isinstance(choices, dict):
                    choices_str = " | ".join([f"{k}: {v}" for k, v in choices.items()])
                else:
                    choices_str = str(choices)

                # Get correct answer
                if vq.randomized_correct_answer:
                    correct_answer = vq.randomized_correct_answer
                else:
                    correct_answer = question.correct_answer

                # Format correct answer
                if isinstance(correct_answer, list):
                    correct_answer_str = ", ".join(correct_answer)
                else:
                    correct_answer_str = str(correct_answer)

                writer.writerow(
                    [
                        variant.version_label,
                        question.prompt,
                        choices_str,
                        correct_answer_str,
                    ]
                )

        return output.getvalue()

    @staticmethod
    def export_variants_to_pdf_files(exam, variant_ids):
        """Export variants to PDF format with UBC professional formatting"""
        return PDFExporter.export_variants_to_pdf_files(exam, variant_ids)

    @staticmethod
    def export_variants_to_docx_files(exam, variant_ids):
        """Export variants to DOCX format with UBC professional formatting"""
        return DOCXExporter.export_variants_to_docx_files(exam, variant_ids)

    @staticmethod
    def export_variants_to_zip(exam, variant_ids, format_type):
        """Export variants to ZIP file"""
        # If variant_ids is None, export all variants
        if variant_ids is None:
            variant_ids = list(exam.variants.values_list("id", flat=True))

        if format_type == "pdf":
            files = PDFExporter.export_variants_to_pdf_files(exam, variant_ids)
        elif format_type == "docx":
            files = DOCXExporter.export_variants_to_docx_files(exam, variant_ids)
        else:
            raise ValueError(f"Unsupported format: {format_type}")

        # Create ZIP file
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w") as zip_file:
            for filename, content in files:
                zip_file.writestr(filename, content)

        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def export_answer_keys_to_zip(exam, variant_ids):
        """Export answer keys to ZIP file"""
        # If variant_ids is None, export all variants
        if variant_ids is None:
            variant_ids = list(exam.variants.values_list("id", flat=True))

        # Create ZIP file
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w") as zip_file:
            # Add CSV answer key
            csv_content = ExamExportService.generate_answer_key_csv(exam, variant_ids)
            zip_file.writestr(f"{exam.title}_answer_keys.csv", csv_content)

        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def export_variants_to_docx(exam, variant_ids):
        """Export variants to single DOCX file (legacy method)"""
        # This method is kept for backward compatibility
        # It now delegates to the new DOCX exporter
        files = DOCXExporter.export_variants_to_docx_files(exam, variant_ids)
        if files:
            return files[0][1]  # Return first file's content
        return b""


class ExamUpdateService:
    """Service for updating exam settings"""

    @staticmethod
    def update_exam_settings(exam: Exam, settings: Dict[str, Any]) -> Exam:
        """Update exam settings"""
        for key, value in settings.items():
            if hasattr(exam, key):
                setattr(exam, key, value)

        exam.save()
        return exam

    @staticmethod
    def update_mandatory_questions(exam: Exam, question_ids: List[int]) -> None:
        """Update mandatory questions for exam"""
        exam.mandatory_questions.set(question_ids)

    @staticmethod
    def update_variant_order(
        exam: Exam, variant_id: int, question_orders: Dict[str, int]
    ) -> None:
        """Update question order in a variant"""
        variant = exam.variants.get(id=variant_id)
        if not variant:
            return

        for question_id, order in question_orders.items():
            vq = variant.variantquestion_set.filter(question_id=question_id).first()
            if vq:
                vq.order = order
                vq.save()
