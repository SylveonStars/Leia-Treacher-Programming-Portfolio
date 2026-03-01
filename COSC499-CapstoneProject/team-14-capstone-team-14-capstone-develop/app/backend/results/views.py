import csv
from datetime import datetime
import io
import json
import zipfile

# results/views.py
from django.db import connection, transaction
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.utils import timezone
from rest_framework import status, viewsets
from rest_framework.decorators import action
from rest_framework.parsers import MultiPartParser
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView

from courses.models import CourseInstructor
from exams.models import Exam

from .models import ExamResult, OMRImportSession
from .serializers import (
    ExamResultSerializer,
    OMRImportSerializer,
    OMRImportSessionSerializer,
)
from .services import OMRImportService, OMRParser, ResultGradingService


class ResultsHealthCheckView(APIView):
    authentication_classes = []
    permission_classes = [AllowAny]

    def get(self, request):
        try:
            with connection.cursor() as cur:
                cur.execute("SELECT 1")
            db_status = "connected"
        except Exception:
            db_status = "unreachable"

        return Response(
            {
                "status": "ok",
                "service": "results",
                "database": db_status,
            }
        )


class UploadPlaceholderView(APIView):
    """
    Root endpoint for results app - provides service info
    """

    permission_classes = [AllowAny]

    def get(self, request):
        return Response(
            {
                "service": "results",
                "version": "1.0",
                "endpoints": {
                    "health": "/api/results/health/",
                    "omr_import": "/api/instructor/exams/{id}/omr/import/",
                    "omr_validate": "/api/instructor/exams/{id}/omr/validate/",
                    "omr_templates": "/api/instructor/exams/{id}/omr/templates/",
                    "exam_results": "/api/instructor/exams/{id}/results/",
                    "student_result": "/api/instructor/exams/{exam_id}/results/{student_id}/",
                },
            }
        )

    def post(self, request):
        return Response(
            {"message": "Please use specific endpoints for uploads"}, status=400
        )


class ExamResultViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing exam results.
    """

    serializer_class = ExamResultSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        # Filter based on user's accessible exams
        # Get courses where user is an instructor
        instructor_courses = CourseInstructor.objects.filter(
            user=self.request.user, accepted=True
        ).values_list("course_id", flat=True)

        # Get exams for those courses
        accessible_exams = Exam.objects.filter(
            course_id__in=instructor_courses
        ).values_list("id", flat=True)

        return ExamResult.objects.filter(exam_id__in=accessible_exams).select_related(
            "exam", "variant", "student"
        )

    def create(self, request, *args, **kwargs):
        """
        Create a single exam result (manual entry).
        """
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        # Grade the result before saving
        exam = serializer.validated_data["exam"]
        variant = serializer.validated_data["variant"]
        responses = serializer.validated_data["raw_responses"]

        grading_result = ResultGradingService.grade_exam_result(
            exam, variant, responses
        )

        # Save with grading results
        instance = serializer.save(
            imported_by=request.user, graded_at=timezone.now(), **grading_result
        )

        return Response(
            self.get_serializer(instance).data, status=status.HTTP_201_CREATED
        )


class OMRImportView(APIView):
    """
    Handle OMR imports for exam results.
    """

    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser]

    @transaction.atomic
    def post(self, request, exam_id):
        """
        POST /api/instructor/exams/:id/omr/import
        Import OMR scan results.
        """
        exam = get_object_or_404(Exam, id=exam_id)

        # Permission check
        if not CourseInstructor.objects.filter(
            course=exam.course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to import results for this exam"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Validate upload payload
        serializer = OMRImportSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        file_obj = serializer.validated_data["file"]
        file_format = serializer.validated_data["format"]
        field_mapping = serializer.validated_data.get("field_mapping")

        # Create import session record
        import_session = OMRImportSession.objects.create(
            exam=exam,
            imported_by=request.user,
            file_name=file_obj.name,
            file_format=file_format,
            status="processing",
        )

        try:
            # Read + parse the file
            content = file_obj.read().decode("utf-8")
            if file_format == "csv":
                parsed_data = OMRParser.parse_csv(content, field_mapping)
            else:
                parsed_data = OMRParser.parse_aiken_txt(content)

            import_session.total_records = len(parsed_data)
            import_session.save()

            # Validate parsed rows against your exam & roster
            validation_result = OMRImportService.validate_import_data(exam, parsed_data)

            if not validation_result["valid"]:
                # mark session failed
                import_session.status = "failed"
                import_session.validation_errors = validation_result["errors"]
                import_session.warnings = validation_result["warnings"]
                import_session.save()

                # build a JSON‑safe preview list
                serializable_preview = []
                for rec in validation_result.get("preview", []):
                    serializable_preview.append(
                        {
                            "student_id": rec["student"].display_id,
                            "variant_code": rec["variant"].version_label,
                            "responses": rec.get("responses", {}),
                            "raw_data": rec.get("raw_data", {}),
                            **(
                                {"row_number": rec["row_number"]}
                                if "row_number" in rec
                                else {}
                            ),
                        }
                    )

                payload = {
                    "valid": validation_result["valid"],
                    "total_records": validation_result["total_records"],
                    "valid_records": validation_result["valid_records"],
                    "errors": validation_result.get("errors", []),
                    "warnings": validation_result.get("warnings", []),
                    "preview": serializable_preview,
                }

                return Response(
                    {
                        "session_id": import_session.id,
                        "status": "failed",
                        "validation_result": payload,
                    },
                    status=status.HTTP_400_BAD_REQUEST,
                )

            # If valid, import them
            results, import_errors = OMRImportService.import_validated_data(
                exam, validation_result["parsed_records"], import_session, request.user
            )

            # finalize session
            import_session.successful_imports = len(results)
            import_session.failed_imports = len(import_errors)
            import_session.import_errors = import_errors
            import_session.warnings = validation_result["warnings"]
            import_session.status = "completed"
            import_session.completed_at = timezone.now()
            import_session.save()

            # serialize and return
            result_data = ExamResultSerializer(results, many=True).data
            return Response(
                {
                    "session_id": import_session.id,
                    "status": "completed",
                    "imported": len(results),
                    "failed": len(import_errors),
                    "warnings": validation_result["warnings"],
                    "results": result_data,
                },
                status=status.HTTP_201_CREATED,
            )

        except Exception as e:
            # on unexpected error, mark session failed
            import_session.status = "failed"
            import_session.import_errors = [{"error": str(e)}]
            import_session.save()
            return Response(
                {"session_id": import_session.id, "status": "failed", "error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )


class OMRValidateView(APIView):
    """
    Validate OMR data before import.
    """

    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser]

    def post(self, request, exam_id):
        """
        POST /api/instructor/exams/:id/omr/validate
        Validate imported data without actually importing.
        """
        exam = get_object_or_404(Exam, id=exam_id)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=exam.course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to access this exam"},
                status=status.HTTP_403_FORBIDDEN,
            )

        serializer = OMRImportSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        file_obj = serializer.validated_data["file"]
        file_format = serializer.validated_data["format"]
        field_mapping = serializer.validated_data.get("field_mapping")

        try:
            # Read and parse file
            file_content = file_obj.read().decode("utf-8")

            if file_format == "csv":
                parsed_data = OMRParser.parse_csv(file_content, field_mapping)
            else:  # txt/aiken
                parsed_data = OMRParser.parse_aiken_txt(file_content)

            # Validate data
            validation_result = OMRImportService.validate_import_data(exam, parsed_data)

            # Don't include full parsed records in response
            validation_result.pop("parsed_records", None)

            # Convert any model instances in `preview` into primitives
            serializable_preview = []
            for rec in validation_result.get("preview", []):
                serializable_preview.append(
                    {
                        # student.display_id is the string student_id
                        "student_id": getattr(rec["student"], "display_id", None),
                        # variant.version_label is the code letter
                        "variant_code": getattr(rec["variant"], "version_label", None),
                        "responses": rec.get("responses", {}),
                        "raw_data": rec.get("raw_data", {}),
                        # if you're interested you can also include row_number
                        **(
                            {"row_number": rec["row_number"]}
                            if "row_number" in rec
                            else {}
                        ),
                    }
                )
            validation_result["preview"] = serializable_preview

            return Response(validation_result)

        except Exception as e:
            return Response(
                {"valid": False, "error": str(e)}, status=status.HTTP_400_BAD_REQUEST
            )


class OMRTemplatesView(APIView):
    """
    Get OMR import format templates.
    """

    permission_classes = [IsAuthenticated]

    def get(self, request, exam_id):
        """
        GET /api/instructor/exams/:id/omr/templates
        Get import format templates and examples.
        """
        exam = get_object_or_404(Exam, id=exam_id)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=exam.course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to access this exam"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Get maximum questions from variants
        max_questions = (
            max(v.questions.count() for v in exam.variants.all())
            if exam.variants.exists()
            else 50
        )

        # Generate CSV template
        csv_template = io.StringIO()
        csv_writer = csv.writer(csv_template)

        # Header row
        headers = ["student_id", "variant_code"] + [
            f"q{i}" for i in range(1, max_questions + 1)
        ]
        csv_writer.writerow(headers)

        # Example rows
        csv_writer.writerow(
            ["S12345", "A"]
            + ["A", "B", "C", "D"] * (max_questions // 4)
            + ["A"] * (max_questions % 4)
        )
        csv_writer.writerow(
            ["S12346", "B"]
            + ["B", "C", "D", "A"] * (max_questions // 4)
            + ["B"] * (max_questions % 4)
        )

        # Aiken format template
        aiken_template = """STUDENT_ID: S12345
VARIANT: A
1. A
2. B
3. C
4. D
5. A
...

STUDENT_ID: S12346
VARIANT: B
1. B
2. C
3. D
4. A
5. B
..."""

        return Response(
            {
                "exam_id": exam.id,
                "exam_title": exam.title,
                "max_questions": max_questions,
                "variants": list(exam.variants.values_list("version_label", flat=True)),
                "templates": {
                    "csv": {
                        "description": "CSV format with student_id, variant_code, and question responses",
                        "example": csv_template.getvalue(),
                        "field_mapping": {
                            "student_id": "student_id",
                            "variant_code": "variant_code",
                            "questions": "q1, q2, q3, ...",
                        },
                    },
                    "txt": {
                        "description": "Aiken format text file with student ID, variant, and numbered responses",
                        "example": aiken_template,
                        "format_rules": [
                            "Each student record starts with STUDENT_ID: line",
                            "Followed by VARIANT: line",
                            "Then numbered responses (e.g., 1. A)",
                            "Blank line separates student records",
                        ],
                    },
                },
            }
        )


class ExamResultsView(APIView):
    """
    Get all results for a specific exam.
    """

    permission_classes = [IsAuthenticated]

    def get(self, request, exam_id):
        """
        GET /api/instructor/exams/:id/results
        Get all results for an exam with statistics.
        """
        exam = get_object_or_404(Exam, id=exam_id)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=exam.course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to view results for this exam"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Get all results
        results = (
            ExamResult.objects.filter(exam=exam)
            .select_related("student", "variant")
            .order_by("-submitted_at")
        )

        # Calculate statistics
        if results.exists():
            scores = [r.score for r in results if r.score is not None]
            if scores:
                import statistics

                stats = {
                    "total_results": results.count(),
                    "mean_score": round(statistics.mean(scores), 2),
                    "median_score": round(statistics.median(scores), 2),
                    "min_score": round(min(scores), 2),
                    "max_score": round(max(scores), 2),
                    "std_dev": (
                        round(statistics.stdev(scores), 2) if len(scores) > 1 else 0
                    ),
                }
            else:
                stats = {"total_results": 0}
        else:
            stats = {"total_results": 0}

        # Serialize results
        serializer = ExamResultSerializer(results, many=True)

        return Response(
            {
                "exam": {
                    "id": exam.id,
                    "title": exam.title,
                    "course": exam.course.code if exam.course else None,
                },
                "statistics": stats,
                "results": serializer.data,
            }
        )

    def delete(self, request, exam_id):
        """
        DELETE /api/results/instructor/exams/:id/results/
        Remove ALL results for this exam.
        """
        exam = get_object_or_404(Exam, id=exam_id)

        # permission check (same as GET)
        if not CourseInstructor.objects.filter(
            course=exam.course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to delete results for this exam"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # delete all results
        deleted_count, _ = ExamResult.objects.filter(exam=exam).delete()
        return Response({"deleted": deleted_count}, status=status.HTTP_204_NO_CONTENT)


class StudentResultView(APIView):
    """
    Get result for a specific student in an exam.
    """

    permission_classes = [IsAuthenticated]

    def get(self, request, exam_id, student_id):
        """
        GET /api/instructor/exams/:exam_id/results/:student_id
        Get detailed result for a specific student.
        """
        exam = get_object_or_404(Exam, id=exam_id)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=exam.course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to view results for this exam"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Get student result
        from courses.models import Student

        student = get_object_or_404(Student, id=student_id)

        try:
            result = ExamResult.objects.get(exam=exam, student=student)
            serializer = ExamResultSerializer(result)

            # Add question-by-question breakdown
            response_data = serializer.data
            response_data["question_breakdown"] = result.grading_details

            return Response(response_data)

        except ExamResult.DoesNotExist:
            return Response(
                {"error": "No result found for this student in this exam"},
                status=status.HTTP_404_NOT_FOUND,
            )


class ExamResultsSummaryView(APIView):
    """
    Get summary statistics for all exams in a course.
    """

    permission_classes = [IsAuthenticated]

    def get(self, request, course_id):
        """
        GET /api/instructor/courses/:course_id/results/summary
        Get summary of all exam results in a course.
        """
        from courses.models import Course

        course = get_object_or_404(Course, id=course_id)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to view results for this course"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Get all exams in the course
        exams = Exam.objects.filter(course=course)

        exam_summaries = []
        for exam in exams:
            results = ExamResult.objects.filter(exam=exam)
            if results.exists():
                scores = [r.score for r in results if r.score is not None]
                if scores:
                    import statistics

                    summary = {
                        "exam_id": exam.id,
                        "exam_title": exam.title,
                        "total_students": results.count(),
                        "mean_score": round(statistics.mean(scores), 2),
                        "median_score": round(statistics.median(scores), 2),
                        "pass_rate": round(
                            len([s for s in scores if s >= 60]) / len(scores) * 100, 2
                        ),  # Assuming 60% is passing
                    }
                    exam_summaries.append(summary)

        return Response(
            {
                "course": {"id": course.id, "code": course.code, "name": course.name},
                "exam_summaries": exam_summaries,
            }
        )


class OMRImportHistoryView(APIView):
    """
    Get import history for an exam.
    """

    permission_classes = [IsAuthenticated]

    def get(self, request, exam_id):
        """
        GET /api/instructor/exams/:id/omr/history
        Get history of all OMR imports for an exam.
        """
        exam = get_object_or_404(Exam, id=exam_id)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=exam.course, user=request.user, accepted=True
        ).exists():
            return Response(
                {
                    "error": "You don't have permission to view import history for this exam"
                },
                status=status.HTTP_403_FORBIDDEN,
            )

        # Get import sessions
        sessions = OMRImportSession.objects.filter(exam=exam).order_by("-created_at")
        serializer = OMRImportSessionSerializer(sessions, many=True)

        return Response(serializer.data)


class ResultExportView(APIView):
    """
    Export exam results in various formats.
    Matches the style of course export but focused only on results.
    """

    permission_classes = [IsAuthenticated]

    def get(self, request, exam_id):
        """
        GET /api/results/export/{exam_id}/?format=csv|pdf|docx|zip
        Export results in requested format.
        """
        exam = get_object_or_404(Exam, id=exam_id)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=exam.course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to export results for this exam"},
                status=status.HTTP_403_FORBIDDEN,
            )

        export_format = request.query_params.get("format", "csv")

        try:
            # Always create a ZIP file for consistency
            zip_buffer = io.BytesIO()

            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                # Add metadata
                metadata = {
                    "exam_id": exam.id,
                    "exam_title": exam.title,
                    "course_code": exam.course.code,
                    "export_date": datetime.now().isoformat(),
                    "exported_by": request.user.email,
                    "format": export_format,
                }
                zip_file.writestr(
                    "export_metadata.json", json.dumps(metadata, indent=2)
                )

                # Get results
                results = list(
                    ExamResult.objects.filter(exam=exam).select_related(
                        "student", "variant"
                    )
                )

                if not results:
                    # Add a note if no results
                    zip_file.writestr(
                        "NO_RESULTS.txt", "No results found for this exam."
                    )
                else:
                    if export_format == "csv":
                        self._add_csv_results_to_zip(zip_file, exam, results)
                    elif export_format == "pdf":
                        self._add_pdf_results_to_zip(zip_file, exam, results)
                    elif export_format == "docx":
                        self._add_docx_results_to_zip(zip_file, exam, results)
                    elif export_format == "zip":
                        # Raw data export - include all formats
                        self._add_csv_results_to_zip(zip_file, exam, results)
                        self._add_pdf_results_to_zip(zip_file, exam, results)
                        self._add_docx_results_to_zip(zip_file, exam, results)
                    else:
                        raise ValueError(f"Invalid format: {export_format}")

            zip_buffer.seek(0)

            # Return the file
            response = HttpResponse(
                zip_buffer.getvalue(), content_type="application/zip"
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{exam.title}_results_{export_format}_{datetime.now().strftime("%Y%m%d")}.zip"'
            )
            return response

        except Exception as e:
            return Response(
                {"error": f"Export failed: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

    @action(detail=False, methods=["get"], url_path="export/(?P<exam_id>[^/.]+)")
    def export(self, request, exam_id=None):
        """
        GET /api/results/exam-results/export/{exam_id}/?format=csv|pdf|docx|zip
        Export results in requested format.
        """
        exam = get_object_or_404(Exam, id=exam_id)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=exam.course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to export results for this exam"},
                status=status.HTTP_403_FORBIDDEN,
            )

        request.query_params.get("format", "csv")

        try:
            # Use the existing ResultExportView logic
            export_view = ResultExportView()
            export_view.format_kwarg = None
            export_view.request = request
            export_view.args = ()
            export_view.kwargs = {"exam_id": exam_id}
            return export_view.get(request, exam_id)
        except Exception as e:
            return Response(
                {"error": f"Export failed: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

    def _add_csv_results_to_zip(self, zip_file, exam, results):
        """Add CSV results files to ZIP with nested structure."""
        # Create a nested ZIP for results organized by variant
        results_zip_buffer = io.BytesIO()
        with zipfile.ZipFile(
            results_zip_buffer, "w", zipfile.ZIP_DEFLATED
        ) as results_zip:
            # Create overall summary CSV
            summary_csv = io.StringIO()
            writer = csv.writer(summary_csv)
            writer.writerow(["Metric", "Value"])

            scores = [r.score for r in results]
            avg_score = sum(scores) / len(scores) if scores else 0
            writer.writerow(["Total Students", str(len(results))])
            writer.writerow(["Average Score", f"{avg_score:.1f}%"])
            writer.writerow(
                ["Highest Score", f"{max(scores):.1f}%" if scores else "N/A"]
            )
            writer.writerow(
                ["Lowest Score", f"{min(scores):.1f}%" if scores else "N/A"]
            )

            # Variant breakdown
            writer.writerow([])  # Empty row
            writer.writerow(["Variant", "Students", "Average Score"])

            variants = {}
            for result in results:
                variant_label = (
                    result.variant.version_label if result.variant else "No Variant"
                )
                if variant_label not in variants:
                    variants[variant_label] = []
                variants[variant_label].append(result)

            for variant_label, variant_results in sorted(variants.items()):
                variant_scores = [r.score for r in variant_results]
                avg = sum(variant_scores) / len(variant_scores) if variant_scores else 0
                writer.writerow(
                    [variant_label, str(len(variant_results)), f"{avg:.1f}%"]
                )

            results_zip.writestr("results_summary.csv", summary_csv.getvalue())

            # Create detailed results CSV for each variant
            for variant_label, variant_results in variants.items():
                variant_csv = io.StringIO()
                writer = csv.writer(variant_csv)
                writer.writerow(
                    [
                        "Student ID",
                        "Display Name",
                        "Legal Name",
                        "Score (%)",
                        "Points Earned",
                        "Points Possible",
                        "Correct",
                        "Incorrect",
                        "Unanswered",
                        "Submitted At",
                    ]
                )

                for result in variant_results:
                    # Handle preferred names
                    if result.student.preferred_name:
                        display_name = (
                            result.student.effective_name or result.student.name
                        )
                    else:
                        display_name = result.student.name

                    writer.writerow(
                        [
                            result.student.student_id,
                            display_name,
                            result.student.name,
                            f"{result.score:.2f}",
                            (
                                f"{result.total_points_earned:.2f}"
                                if result.total_points_earned
                                else "0.00"
                            ),
                            (
                                f"{result.total_points_possible:.2f}"
                                if result.total_points_possible
                                else "0.00"
                            ),
                            str(result.correct_answers),
                            str(result.incorrect_answers),
                            str(result.unanswered),
                            (
                                result.submitted_at.isoformat()
                                if result.submitted_at
                                else ""
                            ),
                        ]
                    )

                results_zip.writestr(
                    f"variant_{variant_label}_results.csv", variant_csv.getvalue()
                )

            # Add question-by-question analysis CSV
            question_csv = io.StringIO()
            writer = csv.writer(question_csv)
            writer.writerow(
                [
                    "Student ID",
                    "Name",
                    "Variant",
                    "Question #",
                    "Response",
                    "Correct Answer",
                    "Status",
                    "Points Earned",
                    "Points Possible",
                ]
            )

            for result in results:
                student_id = result.student.student_id
                student_name = (
                    result.student.effective_name
                    if result.student.preferred_name
                    else result.student.name
                )
                variant_label = (
                    result.variant.version_label if result.variant else "N/A"
                )

                if result.grading_details:
                    for detail in result.grading_details:
                        writer.writerow(
                            [
                                student_id,
                                student_name,
                                variant_label,
                                detail.get("question_number", ""),
                                detail.get("student_answer", ""),
                                ", ".join(detail.get("correct_answers", [])),
                                detail.get("status", ""),
                                f"{detail.get('points_earned', 0):.2f}",
                                f"{detail.get('points_possible', 0):.2f}",
                            ]
                        )

            results_zip.writestr("question_analysis.csv", question_csv.getvalue())

        results_zip_buffer.seek(0)
        zip_file.writestr("exam_results.zip", results_zip_buffer.getvalue())

    def _add_pdf_results_to_zip(self, zip_file, exam, results):
        """Add PDF results files to ZIP."""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        styles = getSampleStyleSheet()

        # Create a nested ZIP for PDFs
        pdf_zip_buffer = io.BytesIO()
        with zipfile.ZipFile(pdf_zip_buffer, "w", zipfile.ZIP_DEFLATED) as pdf_zip:
            # Summary PDF
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            story = []

            story.append(
                Paragraph(f"{exam.course.code} - {exam.title}", styles["Title"])
            )
            story.append(Paragraph("Results Summary", styles["Heading1"]))
            story.append(Spacer(1, 12))

            # Overall statistics
            scores = [r.score for r in results]
            avg_score = sum(scores) / len(scores) if scores else 0
            story.append(Paragraph(f"Total Students: {len(results)}", styles["Normal"]))
            story.append(
                Paragraph(f"Average Score: {avg_score:.1f}%", styles["Normal"])
            )
            story.append(
                Paragraph(
                    f"Highest Score: {max(scores):.1f}%" if scores else "N/A",
                    styles["Normal"],
                )
            )
            story.append(
                Paragraph(
                    f"Lowest Score: {min(scores):.1f}%" if scores else "N/A",
                    styles["Normal"],
                )
            )
            story.append(Spacer(1, 12))

            # Results by variant
            story.append(Paragraph("Results by Variant:", styles["Heading2"]))
            variants = {}
            for result in results:
                variant_label = (
                    result.variant.version_label if result.variant else "No Variant"
                )
                if variant_label not in variants:
                    variants[variant_label] = []
                variants[variant_label].append(result)

            for variant_label, variant_results in sorted(variants.items()):
                variant_scores = [r.score for r in variant_results]
                avg = sum(variant_scores) / len(variant_scores) if variant_scores else 0
                story.append(
                    Paragraph(
                        f"Variant {variant_label}: {len(variant_results)} students, Average: {avg:.1f}%",
                        styles["Normal"],
                    )
                )

            doc.build(story)
            buffer.seek(0)
            pdf_zip.writestr("results_summary.pdf", buffer.getvalue())
            buffer.close()

            # Detailed results PDF for each variant
            for variant_label, variant_results in variants.items():
                buffer = io.BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=letter)
                story = []

                story.append(
                    Paragraph(f"{exam.course.code} - {exam.title}", styles["Title"])
                )
                story.append(
                    Paragraph(f"Variant {variant_label} Results", styles["Heading1"])
                )
                story.append(Spacer(1, 12))

                # Statistics for this variant
                variant_scores = [r.score for r in variant_results]
                avg = sum(variant_scores) / len(variant_scores) if variant_scores else 0
                story.append(
                    Paragraph(f"Students: {len(variant_results)}", styles["Normal"])
                )
                story.append(Paragraph(f"Average Score: {avg:.1f}%", styles["Normal"]))
                story.append(Spacer(1, 12))

                # Results table
                data = [
                    [
                        "Student ID",
                        "Display Name",
                        "Score",
                        "Correct",
                        "Incorrect",
                        "Unanswered",
                    ]
                ]
                for result in variant_results:
                    display_name = (
                        result.student.effective_name
                        if result.student.preferred_name
                        else result.student.name
                    )
                    data.append(
                        [
                            result.student.student_id,
                            display_name,
                            f"{result.score:.1f}%",
                            str(result.correct_answers),
                            str(result.incorrect_answers),
                            str(result.unanswered),
                        ]
                    )

                t = Table(data)
                t.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("FONTSIZE", (0, 0), (-1, 0), 10),
                            ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                            ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ]
                    )
                )
                story.append(t)

                doc.build(story)
                buffer.seek(0)
                pdf_zip.writestr(
                    f"variant_{variant_label}_results.pdf", buffer.getvalue()
                )
                buffer.close()

        pdf_zip_buffer.seek(0)
        zip_file.writestr("exam_results_pdf.zip", pdf_zip_buffer.getvalue())

    def _add_docx_results_to_zip(self, zip_file, exam, results):
        """Add DOCX results files to ZIP."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create a nested ZIP for DOCX files
        docx_zip_buffer = io.BytesIO()
        with zipfile.ZipFile(docx_zip_buffer, "w", zipfile.ZIP_DEFLATED) as docx_zip:
            # Summary DOCX
            document = Document()

            # Title
            title = document.add_heading(f"{exam.course.code} - {exam.title}", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_heading("Results Summary", level=1)

            # Statistics
            scores = [r.score for r in results]
            avg_score = sum(scores) / len(scores) if scores else 0
            document.add_paragraph(f"Total Students: {len(results)}")
            document.add_paragraph(f"Average Score: {avg_score:.1f}%")
            document.add_paragraph(
                f"Highest Score: {max(scores):.1f}%" if scores else "N/A"
            )
            document.add_paragraph(
                f"Lowest Score: {min(scores):.1f}%" if scores else "N/A"
            )

            # Results by variant
            document.add_heading("Results by Variant:", level=2)
            variants = {}
            for result in results:
                variant_label = (
                    result.variant.version_label if result.variant else "No Variant"
                )
                if variant_label not in variants:
                    variants[variant_label] = []
                variants[variant_label].append(result)

            for variant_label, variant_results in sorted(variants.items()):
                variant_scores = [r.score for r in variant_results]
                avg = sum(variant_scores) / len(variant_scores) if variant_scores else 0
                document.add_paragraph(
                    f"Variant {variant_label}: {len(variant_results)} students, Average: {avg:.1f}%",
                    style="List Bullet",
                )

            buffer = io.BytesIO()
            document.save(buffer)
            buffer.seek(0)
            docx_zip.writestr("results_summary.docx", buffer.getvalue())
            buffer.close()

            # Detailed results DOCX for each variant
            for variant_label, variant_results in variants.items():
                document = Document()

                # Header
                title = document.add_heading(f"{exam.course.code} - {exam.title}", 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                subtitle = document.add_heading(
                    f"Variant {variant_label} Results", level=1
                )
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Statistics
                variant_scores = [r.score for r in variant_results]
                avg = sum(variant_scores) / len(variant_scores) if variant_scores else 0
                document.add_paragraph(f"Students: {len(variant_results)}")
                document.add_paragraph(f"Average Score: {avg:.1f}%")

                # Results table
                table = document.add_table(rows=1, cols=7)
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Student ID"
                hdr_cells[1].text = "Display Name"
                hdr_cells[2].text = "Legal Name"
                hdr_cells[3].text = "Score"
                hdr_cells[4].text = "Correct"
                hdr_cells[5].text = "Incorrect"
                hdr_cells[6].text = "Unanswered"

                for result in variant_results:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(result.student.student_id)
                    if result.student.preferred_name:
                        display_name = (
                            result.student.effective_name or result.student.name
                        )
                    else:
                        display_name = result.student.name
                    row_cells[1].text = str(display_name)
                    row_cells[2].text = str(result.student.name)
                    row_cells[3].text = f"{result.score:.1f}%"
                    row_cells[4].text = str(result.correct_answers)
                    row_cells[5].text = str(result.incorrect_answers)
                    row_cells[6].text = str(result.unanswered)

                buffer = io.BytesIO()
                document.save(buffer)
                buffer.seek(0)
                docx_zip.writestr(
                    f"variant_{variant_label}_results.docx", buffer.getvalue()
                )
                buffer.close()

        docx_zip_buffer.seek(0)
        zip_file.writestr("exam_results_docx.zip", docx_zip_buffer.getvalue())


# results/views.py - Complete SimpleExportView
class SimpleExportView(APIView):
    """Export endpoint that supports multiple formats via query parameter."""

    permission_classes = [IsAuthenticated]

    def get(self, request, exam_id):
        """
        GET /api/results/export/{exam_id}/?format=csv|pdf|docx
        Export results in requested format.
        """
        exam = get_object_or_404(Exam, id=exam_id)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=exam.course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to export results"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Get format from query parameter, default to CSV
        export_format = request.query_params.get("format", "csv").lower()

        # Validate format
        if export_format not in ["csv", "pdf", "docx"]:
            return Response(
                {"error": f"Invalid format: {export_format}. Use csv, pdf, or docx."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            # Get results
            results = list(
                ExamResult.objects.filter(exam=exam).select_related(
                    "student", "variant"
                )
            )

            if not results:
                return Response(
                    {"error": "No results found for this exam."},
                    status=status.HTTP_404_NOT_FOUND,
                )

            # For CSV format, return a simple CSV file (not zipped)
            if export_format == "csv":
                response = HttpResponse(content_type="text/csv")
                response["Content-Disposition"] = (
                    f'attachment; filename="{exam.title}_results.csv"'
                )

                writer = csv.writer(response)
                # Header row
                writer.writerow(
                    [
                        "Student ID",
                        "Student Name",
                        "Variant",
                        "Score (%)",
                        "Points Earned",
                        "Points Possible",
                        "Correct",
                        "Incorrect",
                        "Unanswered",
                        "Submitted At",
                    ]
                )

                # Data rows
                for result in results:
                    writer.writerow(
                        [
                            result.student.student_id,
                            result.student.name,
                            result.variant.version_label if result.variant else "N/A",
                            f"{result.score:.2f}" if result.score else "0.00",
                            (
                                f"{result.total_points_earned:.2f}"
                                if result.total_points_earned
                                else "0.00"
                            ),
                            (
                                f"{result.total_points_possible:.2f}"
                                if result.total_points_possible
                                else "0.00"
                            ),
                            result.correct_answers or 0,
                            result.incorrect_answers or 0,
                            result.unanswered or 0,
                            (
                                result.submitted_at.isoformat()
                                if result.submitted_at
                                else ""
                            ),
                        ]
                    )

                return response

            # For PDF and DOCX, use the ResultExportView methods
            else:
                # Create an instance of ResultExportView to use its methods
                export_view = ResultExportView()

                # Create a ZIP file as ResultExportView expects
                zip_buffer = io.BytesIO()

                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                    # Add metadata
                    metadata = {
                        "exam_id": exam.id,
                        "exam_title": exam.title,
                        "course_code": exam.course.code,
                        "export_date": datetime.now().isoformat(),
                        "exported_by": request.user.email,
                        "format": export_format,
                    }
                    zip_file.writestr(
                        "export_metadata.json", json.dumps(metadata, indent=2)
                    )

                    # Use the appropriate method based on format
                    if export_format == "pdf":
                        export_view._add_pdf_results_to_zip(zip_file, exam, results)
                    elif export_format == "docx":
                        export_view._add_docx_results_to_zip(zip_file, exam, results)

                zip_buffer.seek(0)

                # Return the ZIP file
                response = HttpResponse(
                    zip_buffer.getvalue(), content_type="application/zip"
                )
                response["Content-Disposition"] = (
                    f'attachment; filename="{exam.title}_results_{export_format}.zip"'
                )
                return response

        except Exception as e:
            import traceback

            traceback.print_exc()  # This will help debug the issue
            return Response(
                {"error": f"Export failed: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )


# results/views.py - Add these three new view classes

# results/views.py - Enhanced export views that create ZIP files


class ExportCSVView(APIView):
    """Export exam results as CSV in a ZIP file."""

    permission_classes = [IsAuthenticated]

    def get(self, request, exam_id):
        """GET /api/results/export/{exam_id}/csv/"""
        exam = get_object_or_404(Exam, id=exam_id)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=exam.course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to export results"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Get results
        results = list(
            ExamResult.objects.filter(exam=exam).select_related("student", "variant")
        )

        if not results:
            return Response(
                {"error": "No results found for this exam."},
                status=status.HTTP_404_NOT_FOUND,
            )

        # Create ZIP file
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            # 1. Create overall results CSV
            overall_csv = io.StringIO()
            writer = csv.writer(overall_csv)
            writer.writerow(
                [
                    "Student ID",
                    "Student Name",
                    "Variant",
                    "Score (%)",
                    "Points Earned",
                    "Points Possible",
                    "Correct",
                    "Incorrect",
                    "Unanswered",
                    "Submitted At",
                ]
            )

            for result in results:
                writer.writerow(
                    [
                        result.student.student_id,
                        result.student.name,
                        result.variant.version_label if result.variant else "N/A",
                        f"{result.score:.2f}" if result.score else "0.00",
                        (
                            f"{result.total_points_earned:.2f}"
                            if result.total_points_earned
                            else "0.00"
                        ),
                        (
                            f"{result.total_points_possible:.2f}"
                            if result.total_points_possible
                            else "0.00"
                        ),
                        result.correct_answers or 0,
                        result.incorrect_answers or 0,
                        result.unanswered or 0,
                        result.submitted_at.isoformat() if result.submitted_at else "",
                    ]
                )

            zip_file.writestr("overall_results.csv", overall_csv.getvalue())

            # 2. Create summary CSV
            summary_csv = io.StringIO()
            writer = csv.writer(summary_csv)
            writer.writerow(["Metric", "Value"])

            scores = [r.score for r in results if r.score is not None]
            avg_score = sum(scores) / len(scores) if scores else 0
            writer.writerow(["Total Students", str(len(results))])
            writer.writerow(["Average Score", f"{avg_score:.1f}%"])
            writer.writerow(
                ["Highest Score", f"{max(scores):.1f}%" if scores else "N/A"]
            )
            writer.writerow(
                ["Lowest Score", f"{min(scores):.1f}%" if scores else "N/A"]
            )

            # Variant breakdown
            writer.writerow([])  # Empty row
            writer.writerow(["Variant", "Students", "Average Score"])

            variants = {}
            for result in results:
                variant_label = (
                    result.variant.version_label if result.variant else "No Variant"
                )
                if variant_label not in variants:
                    variants[variant_label] = []
                variants[variant_label].append(result)

            for variant_label, variant_results in sorted(variants.items()):
                variant_scores = [r.score for r in variant_results]
                avg = sum(variant_scores) / len(variant_scores) if variant_scores else 0
                writer.writerow(
                    [variant_label, str(len(variant_results)), f"{avg:.1f}%"]
                )

            zip_file.writestr("summary.csv", summary_csv.getvalue())

            # 3. Create CSV for each variant
            for variant_label, variant_results in variants.items():
                variant_csv = io.StringIO()
                writer = csv.writer(variant_csv)
                writer.writerow(
                    [
                        "Student ID",
                        "Student Name",
                        "Score (%)",
                        "Points Earned",
                        "Points Possible",
                        "Correct",
                        "Incorrect",
                        "Unanswered",
                        "Submitted At",
                    ]
                )

                for result in variant_results:
                    writer.writerow(
                        [
                            result.student.student_id,
                            result.student.name,
                            f"{result.score:.2f}" if result.score else "0.00",
                            (
                                f"{result.total_points_earned:.2f}"
                                if result.total_points_earned
                                else "0.00"
                            ),
                            (
                                f"{result.total_points_possible:.2f}"
                                if result.total_points_possible
                                else "0.00"
                            ),
                            result.correct_answers or 0,
                            result.incorrect_answers or 0,
                            result.unanswered or 0,
                            (
                                result.submitted_at.isoformat()
                                if result.submitted_at
                                else ""
                            ),
                        ]
                    )

                zip_file.writestr(
                    f"variant_{variant_label}_results.csv", variant_csv.getvalue()
                )

        zip_buffer.seek(0)
        response = HttpResponse(zip_buffer.getvalue(), content_type="application/zip")
        response["Content-Disposition"] = (
            f'attachment; filename="{exam.title}_results_csv.zip"'
        )
        return response


class ExportPDFView(APIView):
    """Export exam results as PDF in a ZIP file."""

    permission_classes = [IsAuthenticated]

    def get(self, request, exam_id):
        """GET /api/results/export/{exam_id}/pdf/"""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        exam = get_object_or_404(Exam, id=exam_id)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=exam.course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to export results"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Get results
        results = list(
            ExamResult.objects.filter(exam=exam).select_related("student", "variant")
        )

        if not results:
            return Response(
                {"error": "No results found for this exam."},
                status=status.HTTP_404_NOT_FOUND,
            )

        styles = getSampleStyleSheet()

        # Create ZIP file
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            # 1. Overall Results PDF
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            story = []

            story.append(
                Paragraph(f"{exam.course.code} - {exam.title}", styles["Title"])
            )
            story.append(Paragraph("Overall Exam Results", styles["Heading1"]))
            story.append(Spacer(1, 12))

            # Statistics
            scores = [r.score for r in results if r.score is not None]
            if scores:
                avg_score = sum(scores) / len(scores)
                story.append(
                    Paragraph(f"Total Students: {len(results)}", styles["Normal"])
                )
                story.append(
                    Paragraph(f"Average Score: {avg_score:.1f}%", styles["Normal"])
                )
                story.append(
                    Paragraph(f"Highest Score: {max(scores):.1f}%", styles["Normal"])
                )
                story.append(
                    Paragraph(f"Lowest Score: {min(scores):.1f}%", styles["Normal"])
                )
                story.append(Spacer(1, 20))

            # Results table
            data = [["Student ID", "Name", "Variant", "Score", "Points"]]
            for result in results:
                data.append(
                    [
                        result.student.student_id,
                        result.student.name,
                        result.variant.version_label if result.variant else "N/A",
                        f"{result.score:.1f}%" if result.score else "0%",
                        (
                            f"{result.total_points_earned:.1f}/{result.total_points_possible:.1f}"
                            if result.total_points_earned
                            and result.total_points_possible
                            else "N/A"
                        ),
                    ]
                )

            t = Table(data)
            t.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, 0), 10),
                        ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ]
                )
            )
            story.append(t)

            doc.build(story)
            buffer.seek(0)
            zip_file.writestr("overall_results.pdf", buffer.getvalue())
            buffer.close()

            # 2. Summary PDF
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            story = []

            story.append(
                Paragraph(f"{exam.course.code} - {exam.title}", styles["Title"])
            )
            story.append(Paragraph("Results Summary", styles["Heading1"]))
            story.append(Spacer(1, 12))

            # Group by variants
            variants = {}
            for result in results:
                variant_label = (
                    result.variant.version_label if result.variant else "No Variant"
                )
                if variant_label not in variants:
                    variants[variant_label] = []
                variants[variant_label].append(result)

            story.append(Paragraph("Results by Variant:", styles["Heading2"]))
            for variant_label, variant_results in sorted(variants.items()):
                variant_scores = [r.score for r in variant_results]
                avg = sum(variant_scores) / len(variant_scores) if variant_scores else 0
                story.append(
                    Paragraph(
                        f"Variant {variant_label}: {len(variant_results)} students, Average: {avg:.1f}%",
                        styles["Normal"],
                    )
                )

            doc.build(story)
            buffer.seek(0)
            zip_file.writestr("summary.pdf", buffer.getvalue())
            buffer.close()

            # 3. PDF for each variant
            for variant_label, variant_results in variants.items():
                buffer = io.BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=letter)
                story = []

                story.append(
                    Paragraph(f"{exam.course.code} - {exam.title}", styles["Title"])
                )
                story.append(
                    Paragraph(f"Variant {variant_label} Results", styles["Heading1"])
                )
                story.append(Spacer(1, 12))

                # Statistics for this variant
                variant_scores = [r.score for r in variant_results]
                avg = sum(variant_scores) / len(variant_scores) if variant_scores else 0
                story.append(
                    Paragraph(f"Students: {len(variant_results)}", styles["Normal"])
                )
                story.append(Paragraph(f"Average Score: {avg:.1f}%", styles["Normal"]))
                story.append(Spacer(1, 12))

                # Results table
                data = [
                    [
                        "Student ID",
                        "Name",
                        "Score",
                        "Correct",
                        "Incorrect",
                        "Unanswered",
                    ]
                ]
                for result in variant_results:
                    data.append(
                        [
                            result.student.student_id,
                            result.student.name,
                            f"{result.score:.1f}%",
                            str(result.correct_answers),
                            str(result.incorrect_answers),
                            str(result.unanswered),
                        ]
                    )

                t = Table(data)
                t.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("FONTSIZE", (0, 0), (-1, 0), 10),
                            ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                            ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ]
                    )
                )
                story.append(t)

                doc.build(story)
                buffer.seek(0)
                zip_file.writestr(
                    f"variant_{variant_label}_results.pdf", buffer.getvalue()
                )
                buffer.close()

        zip_buffer.seek(0)
        response = HttpResponse(zip_buffer.getvalue(), content_type="application/zip")
        response["Content-Disposition"] = (
            f'attachment; filename="{exam.title}_results_pdf.zip"'
        )
        return response


class ExportDOCXView(APIView):
    """Export exam results as DOCX in a ZIP file."""

    permission_classes = [IsAuthenticated]

    def get(self, request, exam_id):
        """GET /api/results/export/{exam_id}/docx/"""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        exam = get_object_or_404(Exam, id=exam_id)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=exam.course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to export results"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Get results
        results = list(
            ExamResult.objects.filter(exam=exam).select_related("student", "variant")
        )

        if not results:
            return Response(
                {"error": "No results found for this exam."},
                status=status.HTTP_404_NOT_FOUND,
            )

        # Create ZIP file
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            # 1. Overall Results DOCX
            document = Document()

            title = document.add_heading(f"{exam.course.code} - {exam.title}", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_heading("Overall Exam Results", level=1)

            # Statistics
            scores = [r.score for r in results if r.score is not None]
            if scores:
                avg_score = sum(scores) / len(scores)
                document.add_paragraph(f"Total Students: {len(results)}")
                document.add_paragraph(f"Average Score: {avg_score:.1f}%")
                document.add_paragraph(f"Highest Score: {max(scores):.1f}%")
                document.add_paragraph(f"Lowest Score: {min(scores):.1f}%")
                document.add_paragraph()

            # Results table
            table = document.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Student ID"
            hdr_cells[1].text = "Name"
            hdr_cells[2].text = "Variant"
            hdr_cells[3].text = "Score"
            hdr_cells[4].text = "Points"

            for result in results:
                row_cells = table.add_row().cells
                row_cells[0].text = str(result.student.student_id)
                row_cells[1].text = str(result.student.name)
                row_cells[2].text = str(
                    result.variant.version_label if result.variant else "N/A"
                )
                row_cells[3].text = f"{result.score:.1f}%" if result.score else "0%"
                row_cells[4].text = (
                    f"{result.total_points_earned:.1f}/{result.total_points_possible:.1f}"
                    if result.total_points_earned and result.total_points_possible
                    else "N/A"
                )

            buffer = io.BytesIO()
            document.save(buffer)
            buffer.seek(0)
            zip_file.writestr("overall_results.docx", buffer.getvalue())
            buffer.close()

            # 2. Summary DOCX
            document = Document()

            title = document.add_heading(f"{exam.course.code} - {exam.title}", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_heading("Results Summary", level=1)

            # Group by variants
            variants = {}
            for result in results:
                variant_label = (
                    result.variant.version_label if result.variant else "No Variant"
                )
                if variant_label not in variants:
                    variants[variant_label] = []
                variants[variant_label].append(result)

            document.add_heading("Results by Variant:", level=2)
            for variant_label, variant_results in sorted(variants.items()):
                variant_scores = [r.score for r in variant_results]
                avg = sum(variant_scores) / len(variant_scores) if variant_scores else 0
                document.add_paragraph(
                    f"Variant {variant_label}: {len(variant_results)} students, Average: {avg:.1f}%",
                    style="List Bullet",
                )

            buffer = io.BytesIO()
            document.save(buffer)
            buffer.seek(0)
            zip_file.writestr("summary.docx", buffer.getvalue())
            buffer.close()

            # 3. DOCX for each variant
            for variant_label, variant_results in variants.items():
                document = Document()

                title = document.add_heading(f"{exam.course.code} - {exam.title}", 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                subtitle = document.add_heading(
                    f"Variant {variant_label} Results", level=1
                )
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Statistics
                variant_scores = [r.score for r in variant_results]
                avg = sum(variant_scores) / len(variant_scores) if variant_scores else 0
                document.add_paragraph(f"Students: {len(variant_results)}")
                document.add_paragraph(f"Average Score: {avg:.1f}%")

                # Results table
                table = document.add_table(rows=1, cols=6)
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Student ID"
                hdr_cells[1].text = "Name"
                hdr_cells[2].text = "Score"
                hdr_cells[3].text = "Correct"
                hdr_cells[4].text = "Incorrect"
                hdr_cells[5].text = "Unanswered"

                for result in variant_results:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(result.student.student_id)
                    row_cells[1].text = str(result.student.name)
                    row_cells[2].text = f"{result.score:.1f}%"
                    row_cells[3].text = str(result.correct_answers)
                    row_cells[4].text = str(result.incorrect_answers)
                    row_cells[5].text = str(result.unanswered)

                buffer = io.BytesIO()
                document.save(buffer)
                buffer.seek(0)
                zip_file.writestr(
                    f"variant_{variant_label}_results.docx", buffer.getvalue()
                )
                buffer.close()

        zip_buffer.seek(0)
        response = HttpResponse(zip_buffer.getvalue(), content_type="application/zip")
        response["Content-Disposition"] = (
            f'attachment; filename="{exam.title}_results_docx.zip"'
        )
        return response
