import csv
from datetime import datetime
import io
import json
import uuid
import zipfile

from django.contrib.auth import get_user_model
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.db import transaction
from django.db.models import Q
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.utils import timezone
from rest_framework import status, viewsets
from rest_framework.decorators import action
from rest_framework.exceptions import PermissionDenied
from rest_framework.parsers import MultiPartParser
from rest_framework.permissions import BasePermission, IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView

from questions.models import Question, QuestionBank

from .models import (
    Course,
    CourseActivity,
    CourseExportHistory,
    CourseInstructor,
    Student,
)
from .permissions import IsInstructorOrReadOnly
from .serializers import (
    CourseDetailSerializer,
    CourseExportSerializer,
    CourseInstructorSerializer,
    CourseListSerializer,
    CourseSerializer,
    ImportedQuestionSerializer,
    StudentSerializer,
)

User = get_user_model()


# ─── helper ─────────────────────────────────────────────────────────
def _has_full_access(course: Course, user) -> bool:
    """
    True iff the caller is
      • the course creator, OR
      • has an accepted CourseInstructor row with access=FULL.

    For legacy fixtures created *before* CourseInstructor existed we
    fall back to the older `course.instructors` M2M **but only when no
    CourseInstructor row is present for that user**.
    """
    # 1 Creator always has full rights
    if course.creator_id == user.id:
        return True

    # 2 If a CourseInstructor row exists → respect its access flag
    ci_links = course.course_instructors.filter(user=user, accepted=True)
    if ci_links.exists():
        return ci_links.filter(access=CourseInstructor.Access.FULL).exists()

    # 3 Legacy fallback (old tests that still use course.instructors.add)
    return course.instructors.filter(id=user.id).exists()


class CoursesHealthCheckView(APIView):
    """Health check endpoint for courses service"""

    permission_classes = []  # No authentication required

    def get(self, request):
        try:
            # Test database connection
            course_count = Course.objects.count()
            return Response(
                {
                    "status": "ok",  # Changed from "healthy" to match your pattern
                    "service": "courses",
                    "database": "connected",
                    "course_count": course_count,
                }
            )
        except Exception as e:
            return Response(
                {
                    "status": "error",  # Changed from "unhealthy"
                    "service": "courses",
                    "database": "disconnected",
                    "error": str(e),
                },
                status=status.HTTP_503_SERVICE_UNAVAILABLE,
            )


class CourseViewSet(viewsets.ModelViewSet):
    """
    CRUD on courses + instructor roster management.
    """

    permission_classes = [IsAuthenticated, IsInstructorOrReadOnly]

    # ── queryset / serializers ───────────────────────────────────────
    def get_queryset(self):
        user = self.request.user

        unrestricted_actions = [
            "retrieve",
            "update",
            "partial_update",
            "destroy",
            "instructors",
            "add_instructor",
            "remove_instructor",
        ]

        if self.action in unrestricted_actions:
            # Let get_object() find the course first; we'll enforce
            # fine‑grained permissions inside the action methods.
            qs = Course.objects.all()
        else:
            # For list view etc. an instructor should only see their own courses
            qs = Course.objects.filter(
                course_instructors__user=user, course_instructors__accepted=True
            )

        # Optional filters …
        term = self.request.query_params.get("term")
        if term:
            qs = qs.filter(term=term)

        search = self.request.query_params.get("search")
        if search:
            qs = qs.filter(Q(code__icontains=search) | Q(name__icontains=search))

        return qs.distinct().order_by("-last_edited")

    def get_serializer_class(self):
        if self.action == "list":
            return CourseListSerializer
        if self.action == "retrieve":
            return CourseDetailSerializer
        return CourseSerializer

    # ── create ───────────────────────────────────────────────────────
    def perform_create(self, serializer):
        """
        • Set creator to current user (Course.save adds MAIN instructor row).
        • instructor text field → current user's name (for legacy display).
        """
        course = serializer.save(
            creator=self.request.user, instructor=self.request.user.name
        )

        # Log the activity
        CourseActivity.objects.create(
            course=course,  # Now course is defined
            user=self.request.user,
            activity_type=CourseActivity.ActivityType.COURSE_CREATED,
            description=f"{self.request.user.name} created course '{course.name}'",
            entity_type="course",
            entity_id=course.id,
        )

    @action(detail=True, methods=["get"], permission_classes=[IsAuthenticated])
    def activity(self, request, pk=None):
        """Get activity log for a course"""
        course = self.get_object()

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=course, user=request.user, accepted=True
        ).exists():
            return Response({"error": "No permission"}, status=403)

        # Get recent activities (last 20)
        activities = course.activities.select_related("user").all()[:20]

        data = []
        for activity in activities:
            data.append(
                {
                    "id": activity.id,
                    "user": activity.user.name if activity.user else "System",
                    "user_id": activity.user.id if activity.user else None,
                    "activity_type": activity.activity_type,
                    "description": activity.description,
                    "entity_type": activity.entity_type,
                    "entity_id": activity.entity_id,
                    "created_at": activity.created_at.isoformat(),
                }
            )

        return Response(data)

    # ── roster endpoints ─────────────────────────────────────────────
    @action(detail=True, methods=["get"], permission_classes=[IsAuthenticated])
    def instructors(self, request, pk=None):
        """
        GET /api/courses/{id}/instructors/
        Returns full roster for Course Settings page.
        """
        course = self.get_object()

        if not CourseInstructor.objects.filter(
            course=course, user=request.user, accepted=True
        ).exists():
            return Response({"error": "You don't have permission"}, status=403)

        data = CourseInstructorSerializer(
            course.course_instructors.select_related("user").all(),
            many=True,
        ).data
        return Response(data)

    @action(
        detail=True,
        methods=["post"],
        permission_classes=[IsAuthenticated],
        url_path="add_instructor",
    )
    def add_instructor(self, request, pk=None):
        """
        POST { "email": "...", "role": "SEC"[optional] }
        """
        course = self.get_object()

        if not _has_full_access(course, request.user):
            return Response(
                {"error": "No permission"}, status=status.HTTP_403_FORBIDDEN
            )

        email = request.data.get("email")
        role = request.data.get("role", CourseInstructor.Role.SEC)

        if not email:
            return Response(
                {"error": "Email is required"}, status=status.HTTP_400_BAD_REQUEST
            )

        from users.models import User

        try:
            user = User.objects.get(email=email)
        except User.DoesNotExist:
            return Response(
                {"error": "User not found"}, status=status.HTTP_404_NOT_FOUND
            )

        if role == CourseInstructor.Role.SEC:
            start_access = course.default_sec_access
        elif role == CourseInstructor.Role.TA:
            start_access = course.default_ta_access
        else:  # OTH or anything else
            start_access = course.default_oth_access

        link, created = CourseInstructor.objects.get_or_create(
            course=course,
            user=user,
            defaults={
                "role": role,
                "access": start_access,
                "accepted": False,
            },
        )
        ser = CourseInstructorSerializer(link)

        # Log the activity
        CourseActivity.objects.create(
            course=course,
            user=request.user,
            activity_type=CourseActivity.ActivityType.INSTRUCTOR_ADDED,
            description=f"{request.user.name} added {user.name} as {role}",
            entity_type="instructor",
            entity_id=user.id,
        )

        return Response(
            ser.data, status=status.HTTP_201_CREATED if created else status.HTTP_200_OK
        )

    @action(detail=True, methods=["post"], permission_classes=[IsAuthenticated])
    def leave(self, request, pk=None):
        """
        POST /api/courses/{id}/leave/
        Removes the calling user's CourseInstructor link.
        """
        course = self.get_object()

        try:
            link = CourseInstructor.objects.get(course=course, user=request.user)
        except CourseInstructor.DoesNotExist:
            return Response({"error": "Not an instructor"}, status=404)

        # Creators / MAIN instructors must transfer or delete the course explicitly
        if link.role == CourseInstructor.Role.MAIN:
            return Response(
                {"error": "Main instructor must delete or transfer ownership"},
                status=400,
            )

        link.delete()
        return Response({"status": "left"}, status=204)

    @action(
        detail=True,
        methods=["post"],
        permission_classes=[IsAuthenticated],
        url_path="remove_instructor",
    )
    def remove_instructor(self, request, pk=None):
        """
        POST { "email": "..." }
        Only MAIN with FULL access can remove others;
        cannot remove the last MAIN.
        """
        course = self.get_object()
        email = request.data.get("email")
        if not email:
            return Response(
                {"error": "Email is required"}, status=status.HTTP_400_BAD_REQUEST
            )

        if (
            not _has_full_access(course, request.user)
            or not course.course_instructors.filter(
                user=request.user, role=CourseInstructor.Role.MAIN
            ).exists()
        ):
            return Response(
                {"error": "No permission"}, status=status.HTTP_403_FORBIDDEN
            )

        try:
            link = course.course_instructors.select_related("user").get(
                user__email=email
            )
        except CourseInstructor.DoesNotExist:
            return Response(
                {"error": "Not an instructor on this course"},
                status=status.HTTP_404_NOT_FOUND,
            )

        if link.role == CourseInstructor.Role.MAIN:
            # Check if this is the last MAIN instructor
            main_count = course.course_instructors.filter(
                role=CourseInstructor.Role.MAIN
            ).count()
            if main_count <= 1:
                return Response(
                    {"error": "Cannot remove the last main instructor"},
                    status=status.HTTP_400_BAD_REQUEST,
                )

        # Store instructor info before deletion
        instructor_name = link.user.name
        instructor_id = link.user.id

        link.delete()

        # Log the activity
        CourseActivity.objects.create(
            course=course,
            user=request.user,
            activity_type=CourseActivity.ActivityType.INSTRUCTOR_REMOVED,
            description=f"{request.user.name} removed {instructor_name} from the course",
            entity_type="instructor",
            entity_id=instructor_id,
        )

        return Response(status=status.HTTP_204_NO_CONTENT)

    def update(self, request, *args, **kwargs):
        course = self.get_object()

        if not _has_full_access(course, request.user):
            return Response({"error": "No permission"}, status=403)

        response = super().update(request, *args, **kwargs)

        # Log the activity
        if response.status_code == 200:
            CourseActivity.objects.create(
                course=course,
                user=request.user,
                activity_type=CourseActivity.ActivityType.COURSE_UPDATED,
                description=f"{request.user.name} updated course details",
                entity_type="course",
                entity_id=course.id,
            )

        return response

    @action(
        detail=True,
        methods=["patch"],
        url_path=r"instructors/(?P<user_id>\d+)",
        permission_classes=[IsAuthenticated],
    )
    def update_instructor(self, request, pk=None, user_id=None):
        """PATCH body: { "access": "FULL" | "LIMITED" | "NONE" }"""
        course = self.get_object()

        # Only creator or anyone with FULL access may edit others
        if not _has_full_access(course, request.user):
            return Response({"error": "No permission"}, status=403)

        access = request.data.get("access")
        if access not in dict(CourseInstructor.Access.choices):
            return Response({"error": "Invalid access value"}, status=400)

        try:
            link = course.course_instructors.get(user_id=user_id)
        except CourseInstructor.DoesNotExist:
            return Response({"error": "Not an instructor"}, status=404)

        link.access = access
        link.save(update_fields=["access"])
        return Response(CourseInstructorSerializer(link).data)

    def partial_update(self, request, *args, **kwargs):
        course = self.get_object()

        if not _has_full_access(course, request.user):
            return Response({"error": "No permission"}, status=403)

        response = super().partial_update(request, *args, **kwargs)

        # Log the activity
        if response.status_code == 200:
            CourseActivity.objects.create(
                course=course,
                user=request.user,
                activity_type=CourseActivity.ActivityType.COURSE_UPDATED,
                description=f"{request.user.name} updated course settings",
                entity_type="course",
                entity_id=course.id,
            )

        return response

    @action(
        detail=True,
        methods=["put"],
        url_path=r"default_access/(?P<role>MAIN|SEC|TA|OTH)",
        url_name="default-access",
        permission_classes=[IsAuthenticated],
    )
    def default_access(self, request, pk=None, role=None):
        """
        Body: { "access": "FULL" | "LIMITED" | "NONE" }
        Only MAIN or anyone with FULL access can change this.
        """
        course = self.get_object()

        # permission check
        if not _has_full_access(course, request.user):
            return Response({"error": "No permission"}, status=403)

        access = request.data.get("access")
        if access not in dict(CourseInstructor.Access.choices):
            return Response(
                {"error": "Invalid access value"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # mutate the correct field
        field_map = {
            "SEC": "default_sec_access",
            "TA": "default_ta_access",
            "OTH": "default_oth_access",
        }
        if role == "MAIN":
            # we generally keep MAIN at FULL – reject silently
            return Response({"error": "MAIN access is always FULL"}, status=400)
        elif role in field_map:
            setattr(course, field_map[role], access)
            course.save(update_fields=[field_map[role]])
            return Response(
                {"status": "updated", "role": role, "access": access}, status=200
            )
        else:
            return Response({"error": "Invalid role"}, status=400)

    @action(detail=True, methods=["get", "post"], url_path="students")
    @transaction.atomic
    def students(self, request, pk=None):
        """
        List all students enrolled in a specific course or create a new student.

        GET /api/courses/{pk}/students/ - Returns a list of students with their details.
        POST /api/courses/{pk}/students/ - Creates a new student in the course.
        """
        course = self.get_object()

        # Check permissions - only instructors can access
        if request.user not in course.instructors.all():
            return Response(
                {"error": "Only instructors can view student list"},
                status=status.HTTP_403_FORBIDDEN,
            )

        if request.method == "GET":
            students = course.enrolled_students.filter(is_active=True)
            serializer = StudentSerializer(
                students, many=True, context={"request": request}
            )

            return Response(
                {
                    "count": students.count(),
                    "course_id": course.id,
                    "course_name": course.name,
                    "students": serializer.data,
                },
                status=status.HTTP_200_OK,
            )

        elif request.method == "POST":
            from django.db import IntegrityError

            serializer = StudentSerializer(
                data=request.data, context={"request": request}
            )
            if serializer.is_valid():
                try:
                    student = serializer.save(course=course)

                    # Log the activity
                    CourseActivity.objects.create(
                        course=course,
                        user=request.user,
                        activity_type=CourseActivity.ActivityType.STUDENT_ADDED,
                        description=f"{request.user.name} added student '{student.name}'",
                        entity_type="student",
                        entity_id=student.id,
                    )
                except IntegrityError:
                    transaction.set_rollback(True)
                    return Response(
                        {
                            "detail": "A student with this ID or user already exists in this course (unique constraint)."
                        },
                        status=status.HTTP_400_BAD_REQUEST,
                    )
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=["post"], url_path="students/add")
    def add_student(self, request, pk=None):
        """
        Add a student to a specific course.

        POST /api/courses/{pk}/students/add/
        Expected payload: {"user_id": 123} or {"email": "student@example.com"}
        """
        course = self.get_object()

        # Check permissions - only instructors can add students
        if request.user not in course.instructors.all():
            return Response(
                {"error": "Only course instructors can add students"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Get user_id or email from request
        user_id = request.data.get("user_id")
        email = request.data.get("email")

        if not user_id and not email:
            return Response(
                {"error": "Either user_id or email is required"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Find the user
        user = None
        if user_id:
            try:
                user = User.objects.get(id=user_id)
            except User.DoesNotExist:
                return Response(
                    {"error": "User not found"},
                    status=status.HTTP_404_NOT_FOUND,
                )
        elif email:
            try:
                user = User.objects.get(email=email)
            except User.DoesNotExist:
                return Response(
                    {"error": "User with this email not found"},
                    status=status.HTTP_404_NOT_FOUND,
                )

        # Check if student is already enrolled
        if Student.objects.filter(course=course, user=user).exists():
            return Response(
                {"error": "Student is already enrolled in this course"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Create the student enrollment
        from django.db import IntegrityError

        try:
            # Create the student enrollment using the serializer
            student = Student(
                course=course,
                user=user,
                student_id=(
                    user.student_id if hasattr(user, "student_id") else f"S{user.id}"
                ),
                name=user.name,
                email=user.email,
                is_active=True,
            )
            student.save()
        except IntegrityError:
            return Response(
                {
                    "error": "Student is already enrolled in this course (unique constraint)"
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        return Response(
            {
                "message": "Student added successfully",
                "student": {
                    "student_id": student.student_id,
                    "name": student.name,
                    "email": student.email,
                },
            },
            status=status.HTTP_201_CREATED,
        )

    def destroy(self, request, *args, **kwargs):
        course = self.get_object()

        # Only the *creator* may delete the entire course
        if request.user != course.creator:
            return Response(
                {"error": "Only the course creator can delete the course"},
                status=403,
            )

        # Use transaction to ensure atomic deletion
        with transaction.atomic():
            # Delete all related data in the correct order
            # 1. Delete questions first (they reference question banks)
            from questions.models import Question

            Question.objects.filter(bank__course=course).delete()
            # 2. Delete question banks
            from questions.models import QuestionBank

            QuestionBank.objects.filter(course=course).delete()
            # 3. Delete exams (they reference the course)
            from exams.models import Exam

            Exam.objects.filter(course=course).delete()
            # 4. Delete students
            course.enrolled_students.all().delete()
            # 5. Delete course instructors
            course.course_instructors.all().delete()
            # 6. Finally delete the course itself
            course.delete()

        return Response(status=204)

    @action(
        detail=False,
        methods=["get"],
        permission_classes=[IsAuthenticated],
        url_path="search",
    )
    def search(self, request):
        """
        GET /api/courses/search/?q=<query>&instructor_only=<true/false>
        Search courses by code or name
        """
        query = request.query_params.get("q", "").strip()
        instructor_only = (
            request.query_params.get("instructor_only", "false").lower() == "true"
        )

        if not query:
            return Response({"courses": []})

        # Start with base queryset
        if instructor_only:
            # Only courses where user is an instructor
            qs = Course.objects.filter(
                course_instructors__user=request.user, course_instructors__accepted=True
            ).distinct()
        else:
            # All courses user has created
            qs = Course.objects.filter(creator=request.user)

        # Apply search filter
        qs = qs.filter(Q(code__icontains=query) | Q(name__icontains=query)).order_by(
            "code", "name"
        )

        # Serialize results
        from .serializers import CourseListSerializer

        serializer = CourseListSerializer(qs, many=True)

        return Response({"courses": serializer.data, "count": len(serializer.data)})

    @action(detail=True, methods=["post"], permission_classes=[IsAuthenticated])
    def export(self, request, pk=None):
        """Export course data in various formats"""
        course = self.get_object()

        # Check permissions
        if not _has_full_access(course, request.user):
            return Response({"error": "No permission"}, status=403)

        serializer = CourseExportSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        options = serializer.validated_data

        try:
            format_type = options.get("format", "zip")

            # Always create a ZIP file, but contents depend on format
            zip_buffer = io.BytesIO()

            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                # Add metadata
                metadata = {
                    "course_code": course.code,
                    "course_name": course.name,
                    "export_date": datetime.now().isoformat(),
                    "exported_by": request.user.email,
                    "format": format_type,
                }
                zip_file.writestr(
                    "export_metadata.json", json.dumps(metadata, indent=2)
                )

                if format_type == "zip":
                    # Original raw data export
                    self._add_raw_data_to_zip(zip_file, course, options)
                elif format_type == "pdf":
                    # Generate individual PDF files for each data type
                    self._add_pdf_files_to_zip(zip_file, course, options)
                elif format_type == "docx":
                    # Generate individual DOCX files for each data type
                    self._add_docx_files_to_zip(zip_file, course, options)
                elif format_type == "csv":
                    # Generate individual CSV files for each data type
                    self._add_csv_files_to_zip(zip_file, course, options)
                else:
                    raise ValueError(f"Invalid format: {format_type}")

            zip_buffer.seek(0)

            # Generate unique filename
            job_id = str(uuid.uuid4())
            filename = f"exports/course_{course.id}/{job_id}_{course.code}_export_{format_type}_{datetime.now().strftime('%Y%m%d')}.zip"

            # Save file to storage
            file_content = ContentFile(zip_buffer.getvalue())
            file_path = default_storage.save(filename, file_content)

            # Create export history record with file path and format
            CourseExportHistory.objects.create(
                course=course,
                job_id=job_id,
                exported_by=request.user,
                file_size=len(zip_buffer.getvalue()),
                file_path=file_path,
                export_format=format_type,
                status="completed",
            )

            # Return the file for immediate download
            zip_buffer.seek(0)
            response = HttpResponse(
                zip_buffer.getvalue(), content_type="application/zip"
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{course.code}_export_{format_type}_{datetime.now().strftime("%Y%m%d")}.zip"'
            )
            return response

        except Exception as e:
            import traceback

            traceback.print_exc()
            return Response(
                {"error": f"Export failed: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

    @action(
        detail=True,
        methods=["get"],
        permission_classes=[IsAuthenticated],
        url_path="export/(?P<job_id>[^/]+)/download",
    )
    def download_export(self, request, pk=None, job_id=None):
        """Download a previously exported file"""
        course = self.get_object()

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=course, user=request.user, accepted=True
        ).exists():
            return Response({"error": "No permission"}, status=403)

        try:
            export_record = CourseExportHistory.objects.get(
                course=course, job_id=job_id, status="completed"
            )
        except CourseExportHistory.DoesNotExist:
            return Response({"error": "Export not found"}, status=404)

        # Check if file has expired
        if timezone.now() > export_record.expires_at:
            return Response({"error": "Export has expired"}, status=410)

        # Check if file exists
        if not default_storage.exists(export_record.file_path):
            return Response({"error": "Export file not found"}, status=404)

        # Serve the file
        file_content = default_storage.open(export_record.file_path, "rb").read()
        response = HttpResponse(file_content, content_type="application/zip")
        response["Content-Disposition"] = (
            f'attachment; filename="{course.code}_export_{export_record.export_format}_{export_record.created_at.strftime("%Y%m%d")}.zip"'
        )
        return response

    # Update the export_history action to include format info:
    @action(
        detail=True,
        methods=["get"],
        permission_classes=[IsAuthenticated],
        url_path="export/history",
    )
    def export_history(self, request, pk=None):
        """Get export history for a course"""
        try:
            course = Course.objects.get(pk=pk)
        except Course.DoesNotExist:
            return Response({"error": "Course not found"}, status=404)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=course, user=request.user, accepted=True
        ).exists():
            return Response({"error": "No permission"}, status=403)

        # Get export history for this course
        history = CourseExportHistory.objects.filter(
            course=course,
            expires_at__gt=timezone.now(),  # Only show non-expired exports
        ).order_by("-created_at")[:20]

        data = []
        for item in history:
            data.append(
                {
                    "id": item.job_id,
                    "createdAt": item.created_at.isoformat(),
                    "expiresAt": item.expires_at.isoformat(),
                    "size": item.file_size_display,
                    "format": item.export_format,
                    "formatDisplay": item.format_display,
                    "status": item.status,
                    "exportedBy": (
                        item.exported_by.name if item.exported_by else "Unknown"
                    ),
                }
            )

        return Response(data)

    def _add_raw_data_to_zip(self, zip_file, course, options):
        """Add raw JSON/CSV data to ZIP with nested structure"""
        import zipfile

        # Question banks - nested ZIP
        if options.get("question_banks") and course.question_banks.exists():
            qb_zip_buffer = io.BytesIO()
            with zipfile.ZipFile(qb_zip_buffer, "w", zipfile.ZIP_DEFLATED) as qb_zip:
                for bank in course.question_banks.all():
                    bank_data = {
                        "id": bank.id,
                        "title": bank.title,
                        "description": bank.description,
                        "total_questions": bank.questions.count(),
                        "questions": [],
                    }

                    for i, question in enumerate(bank.questions.all(), 1):
                        bank_data["questions"].append(
                            {
                                "number": i,
                                "id": question.id,
                                "prompt": question.prompt,
                                "choices": question.choices,
                                "correct_answer": question.correct_answer,
                                "difficulty": question.difficulty,
                                "tags": question.tags,
                                "explanation": question.explanation,
                            }
                        )

                    safe_title = "".join(
                        c for c in bank.title if c.isalnum() or c in (" ", "-", "_")
                    ).rstrip()
                    qb_zip.writestr(
                        f"{safe_title}.json", json.dumps(bank_data, indent=2)
                    )

            qb_zip_buffer.seek(0)
            zip_file.writestr("question_banks.zip", qb_zip_buffer.getvalue())

        # Exams - nested ZIP
        if options.get("exams") and course.exams.exists():
            exam_zip_buffer = io.BytesIO()
            with zipfile.ZipFile(
                exam_zip_buffer, "w", zipfile.ZIP_DEFLATED
            ) as exam_zip:
                for exam in course.exams.all():
                    safe_exam_title = "".join(
                        c for c in exam.title if c.isalnum() or c in (" ", "-", "_")
                    ).rstrip()

                    # Exam overview
                    exam_data = {
                        "id": exam.id,
                        "title": exam.title,
                        "description": exam.description,
                        "time_limit": exam.time_limit,
                        "questions_per_variant": exam.questions_per_variant,
                        "exam_type": exam.exam_type,
                        "randomize_questions": exam.randomize_questions,
                        "randomize_choices": exam.randomize_choices,
                        "total_variants": exam.variants.count(),
                    }
                    exam_zip.writestr(
                        f"{safe_exam_title}/exam_config.json",
                        json.dumps(exam_data, indent=2),
                    )

                    # Each variant as separate file
                    for variant in exam.variants.all():
                        variant_data = {
                            "exam_title": exam.title,
                            "variant_version": variant.version_label,
                            "time_limit": exam.time_limit,
                            "total_questions": variant.variantquestion_set.count(),
                            "questions": [],
                        }

                        for vq in variant.variantquestion_set.select_related(
                            "question"
                        ).order_by("order"):
                            variant_data["questions"].append(
                                {
                                    "order": vq.order,
                                    "question_id": vq.question.id,
                                    "prompt": vq.question.prompt,
                                    "choices": (
                                        vq.randomized_choices
                                        if vq.randomized_choices
                                        else vq.question.choices
                                    ),
                                    "correct_answer": (
                                        vq.randomized_correct_answer
                                        if vq.randomized_correct_answer
                                        else vq.question.correct_answer
                                    ),
                                    "original_choices": vq.question.choices,
                                    "original_correct_answer": vq.question.correct_answer,
                                }
                            )

                        exam_zip.writestr(
                            f"{safe_exam_title}/variant_{variant.version_label}.json",
                            json.dumps(variant_data, indent=2),
                        )

            exam_zip_buffer.seek(0)
            zip_file.writestr("exams.zip", exam_zip_buffer.getvalue())

        # Students
        if options.get("students"):
            students_csv = io.StringIO()
            writer = csv.writer(students_csv)
            writer.writerow(
                [
                    "Student ID",
                    "Name",
                    "Preferred Name",
                    "Email",
                    "Section",
                    "Status",
                    "Enrolled At",
                ]
            )

            for student in course.enrolled_students.all():
                if options.get("anonymize_students"):
                    writer.writerow(
                        [
                            f"ANON_{student.id}",
                            f"Student {student.id}",
                            "",
                            f"student{student.id}@anonymous.edu",
                            student.section or "",
                            "Active" if student.is_active else "Inactive",
                            student.enrolled_at.isoformat(),
                        ]
                    )
                else:
                    writer.writerow(
                        [
                            student.student_id,
                            student.name,
                            student.preferred_name or "",
                            student.email or "",
                            student.section or "",
                            "Active" if student.is_active else "Inactive",
                            student.enrolled_at.isoformat(),
                        ]
                    )

            zip_file.writestr("student_roster.csv", students_csv.getvalue())

        # Results - UPDATED TO INCLUDE PREFERRED NAMES
        # Results - UPDATED TO MATCH CSV STRUCTURE WITH NESTED ZIP
        if options.get("results"):
            from results.models import ExamResult

            # Create a nested ZIP for results
            results_zip_buffer = io.BytesIO()
            with zipfile.ZipFile(
                results_zip_buffer, "w", zipfile.ZIP_DEFLATED
            ) as results_zip:
                for exam in course.exams.all():
                    results = list(
                        ExamResult.objects.filter(exam=exam).select_related("student")
                    )
                    if results:
                        safe_exam_title = "".join(
                            c for c in exam.title if c.isalnum() or c in (" ", "-", "_")
                        ).rstrip()

                        # Create results overview CSV (same as in CSV export)
                        overview_csv = io.StringIO()
                        writer = csv.writer(overview_csv)
                        writer.writerow(["Metric", "Value"])

                        scores = [r.score for r in results]
                        avg_score = sum(scores) / len(scores) if scores else 0
                        writer.writerow(["Total Students", str(len(results))])
                        writer.writerow(["Average Score", f"{avg_score:.1f}%"])
                        writer.writerow(
                            [
                                "Highest Score",
                                f"{max(scores):.1f}%" if scores else "N/A",
                            ]
                        )
                        writer.writerow(
                            ["Lowest Score", f"{min(scores):.1f}%" if scores else "N/A"]
                        )

                        # Add variant breakdown
                        writer.writerow([])  # Empty row
                        writer.writerow(["Variant", "Students", "Average Score"])

                        variants = {}
                        for result in results:
                            variant_label = (
                                result.variant.version_label
                                if result.variant
                                else "No Variant"
                            )
                            if variant_label not in variants:
                                variants[variant_label] = []
                            variants[variant_label].append(result.score)

                        for variant_label, variant_scores in sorted(variants.items()):
                            avg = (
                                sum(variant_scores) / len(variant_scores)
                                if variant_scores
                                else 0
                            )
                            writer.writerow(
                                [variant_label, str(len(variant_scores)), f"{avg:.1f}%"]
                            )

                        results_zip.writestr(
                            f"{safe_exam_title}/results_overview.csv",
                            overview_csv.getvalue(),
                        )

                        # Create detailed results CSV for each variant
                        for variant_label in variants.keys():
                            variant_result_objects = [
                                r
                                for r in results
                                if (
                                    r.variant.version_label
                                    if r.variant
                                    else "No Variant"
                                )
                                == variant_label
                            ]

                            variant_csv = io.StringIO()
                            writer = csv.writer(variant_csv)
                            writer.writerow(
                                [
                                    "Student ID",
                                    "Display Name",
                                    "Legal Name",
                                    "Score",
                                    "Correct",
                                    "Incorrect",
                                    "Unanswered",
                                    "Date",
                                ]
                            )

                            for result in variant_result_objects:
                                if options.get("anonymize_students"):
                                    student_id = f"ANON_{result.student.id}"
                                    display_name = f"Student {result.student.id}"
                                    legal_name = f"Student {result.student.id}"
                                else:
                                    student_id = result.student.student_id
                                    if result.student.preferred_name:
                                        display_name = (
                                            result.student.effective_name
                                            or result.student.name
                                        )
                                    else:
                                        display_name = result.student.name
                                    legal_name = result.student.name

                                writer.writerow(
                                    [
                                        student_id,
                                        display_name,
                                        legal_name,
                                        str(result.score),
                                        str(result.correct_answers),
                                        str(result.incorrect_answers),
                                        str(result.unanswered),
                                        (
                                            result.submitted_at.isoformat()
                                            if hasattr(result, "submitted_at")
                                            else ""
                                        ),
                                    ]
                                )

                            results_zip.writestr(
                                f"{safe_exam_title}/variant_{variant_label}_results.csv",
                                variant_csv.getvalue(),
                            )

            # Add the nested ZIP to main ZIP
            results_zip_buffer.seek(0)
            zip_file.writestr("exam_results.zip", results_zip_buffer.getvalue())

    def _add_pdf_files_to_zip(self, zip_file, course, options):
        """Generate separate PDF files for each data type with nested structure"""
        import zipfile

        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import (
            PageBreak,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        styles = getSampleStyleSheet()

        # Question Banks - Create a nested ZIP with individual PDFs for each bank
        if options.get("question_banks") and course.question_banks.exists():
            # Create a nested ZIP for question banks
            qb_zip_buffer = io.BytesIO()
            with zipfile.ZipFile(qb_zip_buffer, "w", zipfile.ZIP_DEFLATED) as qb_zip:
                for bank in course.question_banks.all():
                    buffer = io.BytesIO()
                    doc = SimpleDocTemplate(buffer, pagesize=letter)
                    story = []

                    # Title page for this question bank
                    story.append(
                        Paragraph(f"{course.code} - Question Bank", styles["Title"])
                    )
                    story.append(Spacer(1, 12))
                    story.append(Paragraph(bank.title, styles["Heading1"]))
                    if bank.description:
                        story.append(Paragraph(bank.description, styles["Normal"]))
                    story.append(Spacer(1, 24))
                    story.append(
                        Paragraph(
                            f"Total Questions: {bank.questions.count()}",
                            styles["Normal"],
                        )
                    )
                    story.append(PageBreak())

                    # Add each question
                    for i, question in enumerate(bank.questions.all(), 1):
                        story.append(Paragraph(f"Question {i}", styles["Heading2"]))
                        story.append(Spacer(1, 12))

                        # Question prompt
                        story.append(Paragraph("Prompt:", styles["Heading3"]))
                        story.append(Paragraph(question.prompt, styles["Normal"]))
                        story.append(Spacer(1, 12))

                        # Choices
                        story.append(Paragraph("Answer Choices:", styles["Heading3"]))
                        for key, value in sorted(question.choices.items()):
                            story.append(Paragraph(f"{key}) {value}", styles["Normal"]))
                        story.append(Spacer(1, 12))

                        # Correct answer - FIX: Ensure all elements are strings
                        correct_answer_str = ", ".join(
                            str(ans) for ans in question.correct_answer
                        )
                        story.append(
                            Paragraph(
                                f"Correct Answer(s): {correct_answer_str}",
                                styles["Normal"],
                            )
                        )
                        story.append(Spacer(1, 6))

                        # Additional info
                        if question.difficulty:
                            story.append(
                                Paragraph(
                                    f"Difficulty: {question.difficulty}",
                                    styles["Normal"],
                                )
                            )

                        if question.tags:
                            # FIX: Ensure tags are strings
                            tags_str = ", ".join(str(tag) for tag in question.tags)
                            story.append(
                                Paragraph(f"Tags: {tags_str}", styles["Normal"])
                            )

                        if question.explanation:
                            story.append(Spacer(1, 6))
                            story.append(Paragraph("Explanation:", styles["Heading3"]))
                            story.append(
                                Paragraph(question.explanation, styles["Italic"])
                            )

                        # Add separator between questions
                        if i < bank.questions.count():
                            story.append(Spacer(1, 12))
                            story.append(Paragraph("_" * 80, styles["Normal"]))
                            story.append(Spacer(1, 12))

                    # Build the PDF
                    doc.build(story)

                    # Flush and seek
                    buffer.flush()
                    buffer.seek(0)

                    # Add to nested ZIP with sanitized filename
                    safe_title = "".join(
                        c for c in bank.title if c.isalnum() or c in (" ", "-", "_")
                    ).rstrip()
                    qb_zip.writestr(f"{safe_title}.pdf", buffer.getvalue())
                    buffer.close()

            # Close and flush the question banks ZIP
            qb_zip.close()
            qb_zip_buffer.flush()
            qb_zip_buffer.seek(0)
            zip_file.writestr("question_banks.zip", qb_zip_buffer.getvalue())
            qb_zip_buffer.close()

        # Exams - Create a nested ZIP with PDFs for each exam and its variants
        if options.get("exams") and course.exams.exists():
            # Create a nested ZIP for exams
            exam_zip_buffer = io.BytesIO()
            with zipfile.ZipFile(
                exam_zip_buffer, "w", zipfile.ZIP_DEFLATED
            ) as exam_zip:
                for exam in course.exams.all():
                    safe_exam_title = "".join(
                        c for c in exam.title if c.isalnum() or c in (" ", "-", "_")
                    ).rstrip()

                    # Create exam overview PDF
                    buffer = io.BytesIO()
                    doc = SimpleDocTemplate(buffer, pagesize=letter)
                    story = []

                    story.append(
                        Paragraph(f"{course.code} - Exam Overview", styles["Title"])
                    )
                    story.append(Spacer(1, 12))
                    story.append(Paragraph(exam.title, styles["Heading1"]))
                    if exam.description:
                        story.append(Paragraph(exam.description, styles["Normal"]))
                    story.append(Spacer(1, 12))

                    # Exam details
                    story.append(
                        Paragraph(
                            f"Time Limit: {exam.time_limit} minutes", styles["Normal"]
                        )
                    )
                    story.append(
                        Paragraph(
                            f"Questions per Variant: {exam.questions_per_variant}",
                            styles["Normal"],
                        )
                    )
                    story.append(
                        Paragraph(f"Exam Type: {exam.exam_type}", styles["Normal"])
                    )
                    story.append(
                        Paragraph(
                            f"Randomize Questions: {'Yes' if exam.randomize_questions else 'No'}",
                            styles["Normal"],
                        )
                    )
                    story.append(
                        Paragraph(
                            f"Randomize Choices: {'Yes' if exam.randomize_choices else 'No'}",
                            styles["Normal"],
                        )
                    )
                    story.append(Spacer(1, 12))

                    # Variant summary
                    story.append(Paragraph("Variants:", styles["Heading2"]))
                    for variant in exam.variants.all():
                        story.append(
                            Paragraph(
                                f"• Version {variant.version_label}: {variant.variantquestion_set.count()} questions",
                                styles["Normal"],
                            )
                        )

                    # Build the PDF
                    doc.build(story)

                    # IMPORTANT: Flush and seek to beginning
                    buffer.flush()
                    buffer.seek(0)

                    # Write to exam ZIP
                    exam_zip.writestr(
                        f"{safe_exam_title}/exam_overview.pdf", buffer.getvalue()
                    )
                    buffer.close()  # Close the buffer

                    # Create a PDF for each variant
                    for variant in exam.variants.all():
                        buffer = io.BytesIO()
                        doc = SimpleDocTemplate(buffer, pagesize=letter)
                        story = []

                        # Variant header
                        story.append(
                            Paragraph(f"{course.code} - {exam.title}", styles["Title"])
                        )
                        story.append(Spacer(1, 12))
                        story.append(
                            Paragraph(
                                f"Variant {variant.version_label}", styles["Heading1"]
                            )
                        )
                        story.append(Spacer(1, 6))
                        story.append(
                            Paragraph(
                                f"Time Limit: {exam.time_limit} minutes",
                                styles["Normal"],
                            )
                        )
                        story.append(
                            Paragraph(
                                f"Total Questions: {variant.variantquestion_set.count()}",
                                styles["Normal"],
                            )
                        )
                        story.append(PageBreak())

                        # Add each question in the variant
                        for vq in variant.variantquestion_set.select_related(
                            "question"
                        ).order_by("order"):
                            question = vq.question
                            story.append(
                                Paragraph(f"Question {vq.order}", styles["Heading2"])
                            )
                            story.append(Spacer(1, 12))

                            # Question prompt
                            story.append(Paragraph(question.prompt, styles["Normal"]))
                            story.append(Spacer(1, 12))

                            # Choices (using randomized order if available)
                            if vq.randomized_choices:
                                choices = vq.randomized_choices
                            else:
                                choices = question.choices

                            for key, value in sorted(choices.items()):
                                story.append(
                                    Paragraph(f"{key}) {value}", styles["Normal"])
                                )
                            story.append(Spacer(1, 12))

                            # For instructor reference - FIX: Ensure answer elements are strings
                            if vq.randomized_correct_answer:
                                answer_str = ", ".join(
                                    str(ans) for ans in vq.randomized_correct_answer
                                )
                                story.append(
                                    Paragraph(
                                        f"[Answer Key: {answer_str}]", styles["Italic"]
                                    )
                                )
                            else:
                                answer_str = ", ".join(
                                    str(ans) for ans in question.correct_answer
                                )
                                story.append(
                                    Paragraph(
                                        f"[Answer Key: {answer_str}]", styles["Italic"]
                                    )
                                )

                            # Add separator between questions
                            if vq.order < variant.variantquestion_set.count():
                                story.append(Spacer(1, 12))
                                story.append(Paragraph("_" * 80, styles["Normal"]))
                                story.append(Spacer(1, 12))

                        # Build the variant PDF
                        doc.build(story)

                        # IMPORTANT: Flush and seek
                        buffer.flush()
                        buffer.seek(0)

                        # Write to exam ZIP
                        exam_zip.writestr(
                            f"{safe_exam_title}/variant_{variant.version_label}.pdf",
                            buffer.getvalue(),
                        )
                        buffer.close()  # Close the buffer

            # IMPORTANT: Properly close the exam ZIP
            exam_zip.close()

            # Flush and seek the exam ZIP buffer
            exam_zip_buffer.flush()
            exam_zip_buffer.seek(0)

            # Add the nested ZIP to main ZIP
            zip_file.writestr("exams.zip", exam_zip_buffer.getvalue())
            exam_zip_buffer.close()  # Close the buffer

        # Students PDF - UPDATED
        if options.get("students") and course.enrolled_students.exists():
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            story = []

            story.append(Paragraph(f"{course.code} - Student Roster", styles["Title"]))
            story.append(Spacer(1, 12))
            story.append(
                Paragraph(
                    f"Total Students: {course.enrolled_students.count()}",
                    styles["Normal"],
                )
            )
            story.append(Spacer(1, 12))

            # Create table
            data = [
                [
                    "Student ID",
                    "Display Name",
                    "Legal Name",
                    "Email",
                    "Section",
                    "Status",
                ]
            ]
            for student in course.enrolled_students.all():
                if options.get("anonymize_students"):
                    data.append(
                        [
                            f"ANON_{student.id}",
                            f"Student {student.id}",
                            f"Student {student.id}",
                            f"student{student.id}@anonymous.edu",
                            student.section or "",
                            "Active" if student.is_active else "Inactive",
                        ]
                    )
                else:
                    display_name = (
                        student.effective_name
                        if student.preferred_name
                        else student.name
                    )
                    data.append(
                        [
                            student.student_id,
                            display_name,
                            student.name,
                            student.email or "",
                            student.section or "",
                            "Active" if student.is_active else "Inactive",
                        ]
                    )

            if len(data) > 1:  # Only create table if there are students
                t = Table(data)
                t.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("FONTSIZE", (0, 0), (-1, 0), 10),
                            ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                            ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ]
                    )
                )
                story.append(t)

            # Build the PDF
            if story:  # Only build if there's content
                doc.build(story)
                buffer.flush()
                buffer.seek(0)
                zip_file.writestr("student_roster.pdf", buffer.getvalue())
                buffer.close()

        # Results PDF - NEW NESTED STRUCTURE
        if options.get("results"):
            from results.models import ExamResult

            # Create a nested ZIP for results
            results_zip_buffer = io.BytesIO()
            with zipfile.ZipFile(
                results_zip_buffer, "w", zipfile.ZIP_DEFLATED
            ) as results_zip:
                for exam in course.exams.all():
                    results = list(
                        ExamResult.objects.filter(exam=exam).select_related("student")
                    )
                    if results:  # Only process if there are results
                        safe_exam_title = "".join(
                            c for c in exam.title if c.isalnum() or c in (" ", "-", "_")
                        ).rstrip()

                        # Create results overview PDF
                        buffer = io.BytesIO()
                        doc = SimpleDocTemplate(buffer, pagesize=letter)
                        story = []

                        story.append(
                            Paragraph(
                                f"{course.code} - {exam.title} Results Overview",
                                styles["Title"],
                            )
                        )
                        story.append(Spacer(1, 12))

                        # Overall statistics
                        scores = [r.score for r in results]
                        avg_score = sum(scores) / len(scores) if scores else 0
                        story.append(
                            Paragraph(
                                f"Total Students: {len(results)}", styles["Normal"]
                            )
                        )
                        story.append(
                            Paragraph(
                                f"Average Score: {avg_score:.1f}%", styles["Normal"]
                            )
                        )
                        story.append(
                            Paragraph(
                                (
                                    f"Highest Score: {max(scores):.1f}%"
                                    if scores
                                    else "N/A"
                                ),
                                styles["Normal"],
                            )
                        )
                        story.append(
                            Paragraph(
                                (
                                    f"Lowest Score: {min(scores):.1f}%"
                                    if scores
                                    else "N/A"
                                ),
                                styles["Normal"],
                            )
                        )
                        story.append(Spacer(1, 12))

                        # Results by variant summary
                        story.append(
                            Paragraph("Results by Variant:", styles["Heading2"])
                        )
                        variants = {}
                        for result in results:
                            variant_label = (
                                result.variant.version_label
                                if result.variant
                                else "No Variant"
                            )
                            if variant_label not in variants:
                                variants[variant_label] = []
                            variants[variant_label].append(result.score)

                        for variant_label, variant_scores in sorted(variants.items()):
                            avg = (
                                sum(variant_scores) / len(variant_scores)
                                if variant_scores
                                else 0
                            )
                            story.append(
                                Paragraph(
                                    f"Variant {variant_label}: {len(variant_scores)} students, Average: {avg:.1f}%",
                                    styles["Normal"],
                                )
                            )

                        doc.build(story)
                        buffer.flush()
                        buffer.seek(0)
                        results_zip.writestr(
                            f"{safe_exam_title}/results_overview.pdf", buffer.getvalue()
                        )
                        buffer.close()

                        # Create detailed results PDF for each variant
                        for variant_label, variant_results in variants.items():
                            variant_result_objects = [
                                r
                                for r in results
                                if (
                                    r.variant.version_label
                                    if r.variant
                                    else "No Variant"
                                )
                                == variant_label
                            ]

                            buffer = io.BytesIO()
                            doc = SimpleDocTemplate(buffer, pagesize=letter)
                            story = []

                            story.append(
                                Paragraph(
                                    f"{course.code} - {exam.title}", styles["Title"]
                                )
                            )
                            story.append(
                                Paragraph(
                                    f"Variant {variant_label} Results",
                                    styles["Heading1"],
                                )
                            )
                            story.append(Spacer(1, 12))

                            # Variant statistics
                            avg = (
                                sum(variant_results) / len(variant_results)
                                if variant_results
                                else 0
                            )
                            story.append(
                                Paragraph(
                                    f"Students: {len(variant_results)}",
                                    styles["Normal"],
                                )
                            )
                            story.append(
                                Paragraph(
                                    f"Average Score: {avg:.1f}%", styles["Normal"]
                                )
                            )
                            story.append(Spacer(1, 12))

                            # Detailed results table
                            data = [
                                [
                                    "Student ID",
                                    "Display Name",
                                    "Legal Name",
                                    "Score",
                                    "Correct",
                                    "Incorrect",
                                    "Unanswered",
                                ]
                            ]
                            for result in variant_result_objects:
                                if options.get("anonymize_students"):
                                    student_id = f"ANON_{result.student.id}"
                                    display_name = f"Student {result.student.id}"
                                    legal_name = f"Student {result.student.id}"
                                else:
                                    student_id = result.student.student_id
                                    if result.student.preferred_name:
                                        display_name = (
                                            result.student.effective_name
                                            or result.student.name
                                        )
                                    else:
                                        display_name = result.student.name
                                    legal_name = result.student.name

                                data.append(
                                    [
                                        student_id,
                                        display_name,
                                        legal_name,
                                        f"{result.score}%",
                                        str(result.correct_answers),
                                        str(result.incorrect_answers),
                                        str(result.unanswered),
                                    ]
                                )

                            t = Table(data)
                            t.setStyle(
                                TableStyle(
                                    [
                                        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                                        (
                                            "TEXTCOLOR",
                                            (0, 0),
                                            (-1, 0),
                                            colors.whitesmoke,
                                        ),
                                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                                    ]
                                )
                            )
                            story.append(t)

                            doc.build(story)
                            buffer.flush()
                            buffer.seek(0)
                            results_zip.writestr(
                                f"{safe_exam_title}/variant_{variant_label}_results.pdf",
                                buffer.getvalue(),
                            )
                            buffer.close()

            # Close and flush the results ZIP
            results_zip.close()
            results_zip_buffer.flush()
            results_zip_buffer.seek(0)

            # Add the nested ZIP to main ZIP
            zip_file.writestr("exam_results.zip", results_zip_buffer.getvalue())
            results_zip_buffer.close()

    def _add_docx_files_to_zip(self, zip_file, course, options):
        """Generate separate DOCX files with nested structure"""
        import zipfile

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Question Banks - nested ZIP
        if options.get("question_banks") and course.question_banks.exists():
            qb_zip_buffer = io.BytesIO()
            with zipfile.ZipFile(qb_zip_buffer, "w", zipfile.ZIP_DEFLATED) as qb_zip:
                for bank in course.question_banks.all():
                    document = Document()

                    # Title
                    title = document.add_heading(f"{course.code} - Question Bank", 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    document.add_heading(bank.title, level=1)
                    if bank.description:
                        document.add_paragraph(bank.description)

                    document.add_paragraph(f"Total Questions: {bank.questions.count()}")
                    document.add_page_break()

                    # Add each question
                    for i, question in enumerate(bank.questions.all(), 1):
                        document.add_heading(f"Question {i}", level=2)

                        # Prompt
                        document.add_heading("Prompt:", level=3)
                        document.add_paragraph(question.prompt)

                        # Choices
                        document.add_heading("Answer Choices:", level=3)
                        for key, value in sorted(question.choices.items()):
                            document.add_paragraph(
                                f"{key}) {value}", style="List Bullet"
                            )

                        # Correct Answer - FIX: Ensure it's a list and all elements are strings
                        p = document.add_paragraph()
                        p.add_run("Correct Answer(s): ").bold = True
                        # Handle if correct_answer is not a list
                        if isinstance(question.correct_answer, list):
                            correct_answer_str = ", ".join(
                                str(ans) for ans in question.correct_answer
                            )
                        else:
                            correct_answer_str = str(question.correct_answer)
                        p.add_run(correct_answer_str)

                        # Additional info
                        if question.difficulty:
                            p = document.add_paragraph()
                            p.add_run("Difficulty: ").bold = True
                            p.add_run(str(question.difficulty))

                        if question.tags:
                            p = document.add_paragraph()
                            p.add_run("Tags: ").bold = True
                            # Handle if tags is not a list
                            if isinstance(question.tags, list):
                                tags_str = ", ".join(str(tag) for tag in question.tags)
                            else:
                                tags_str = str(question.tags)
                            p.add_run(tags_str)

                        if question.explanation:
                            document.add_heading("Explanation:", level=3)
                            document.add_paragraph(str(question.explanation))

                        if i < bank.questions.count():
                            document.add_paragraph("_" * 80)

                    buffer = io.BytesIO()
                    document.save(buffer)
                    buffer.flush()
                    buffer.seek(0)

                    safe_title = "".join(
                        c for c in bank.title if c.isalnum() or c in (" ", "-", "_")
                    ).rstrip()
                    qb_zip.writestr(f"{safe_title}.docx", buffer.getvalue())
                    buffer.close()

            qb_zip.close()
            qb_zip_buffer.flush()
            qb_zip_buffer.seek(0)
            zip_file.writestr("question_banks.zip", qb_zip_buffer.getvalue())
            qb_zip_buffer.close()

        # Exams - nested ZIP
        if options.get("exams") and course.exams.exists():
            exam_zip_buffer = io.BytesIO()
            with zipfile.ZipFile(
                exam_zip_buffer, "w", zipfile.ZIP_DEFLATED
            ) as exam_zip:
                for exam in course.exams.all():
                    safe_exam_title = "".join(
                        c for c in exam.title if c.isalnum() or c in (" ", "-", "_")
                    ).rstrip()

                    # Exam overview
                    document = Document()
                    document.add_heading(f"{course.code} - Exam Overview", 0)
                    document.add_heading(exam.title, level=1)
                    if exam.description:
                        document.add_paragraph(exam.description)

                    document.add_paragraph(f"Time Limit: {exam.time_limit} minutes")
                    document.add_paragraph(
                        f"Questions per Variant: {exam.questions_per_variant}"
                    )
                    document.add_paragraph(f"Exam Type: {exam.exam_type}")
                    document.add_paragraph(
                        f'Randomize Questions: {"Yes" if exam.randomize_questions else "No"}'
                    )
                    document.add_paragraph(
                        f'Randomize Choices: {"Yes" if exam.randomize_choices else "No"}'
                    )

                    document.add_heading("Variants:", level=2)
                    for variant in exam.variants.all():
                        document.add_paragraph(
                            f"Version {variant.version_label}: {variant.variantquestion_set.count()} questions",
                            style="List Bullet",
                        )

                    buffer = io.BytesIO()
                    document.save(buffer)
                    buffer.flush()
                    buffer.seek(0)
                    exam_zip.writestr(
                        f"{safe_exam_title}/exam_overview.docx", buffer.getvalue()
                    )
                    buffer.close()

                    # Each variant
                    for variant in exam.variants.all():
                        document = Document()

                        # Header
                        title = document.add_heading(f"{course.code} - {exam.title}", 0)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        subtitle = document.add_heading(
                            f"Variant {variant.version_label}", level=1
                        )
                        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        document.add_paragraph(f"Time Limit: {exam.time_limit} minutes")
                        document.add_paragraph(
                            f"Total Questions: {variant.variantquestion_set.count()}"
                        )
                        document.add_page_break()

                        # Questions
                        for vq in variant.variantquestion_set.select_related(
                            "question"
                        ).order_by("order"):
                            question = vq.question
                            document.add_heading(f"Question {vq.order}", level=2)
                            document.add_paragraph(question.prompt)

                            # Choices
                            choices = (
                                vq.randomized_choices
                                if vq.randomized_choices
                                else question.choices
                            )
                            for key, value in sorted(choices.items()):
                                document.add_paragraph(
                                    f"{key}) {value}", style="List Bullet"
                                )

                            # Answer key - FIX: Ensure all elements are strings and handle non-list cases
                            p = document.add_paragraph()
                            p.add_run("[Answer Key: ").italic = True

                            if vq.randomized_correct_answer:
                                if isinstance(vq.randomized_correct_answer, list):
                                    answer_str = ", ".join(
                                        str(ans) for ans in vq.randomized_correct_answer
                                    )
                                else:
                                    answer_str = str(vq.randomized_correct_answer)
                            else:
                                if isinstance(question.correct_answer, list):
                                    answer_str = ", ".join(
                                        str(ans) for ans in question.correct_answer
                                    )
                                else:
                                    answer_str = str(question.correct_answer)

                            p.add_run(answer_str).italic = True
                            p.add_run("]").italic = True

                            if vq.order < variant.variantquestion_set.count():
                                document.add_paragraph("_" * 80)

                        buffer = io.BytesIO()
                        document.save(buffer)
                        buffer.flush()
                        buffer.seek(0)
                        exam_zip.writestr(
                            f"{safe_exam_title}/variant_{variant.version_label}.docx",
                            buffer.getvalue(),
                        )
                        buffer.close()

            exam_zip.close()
            exam_zip_buffer.flush()
            exam_zip_buffer.seek(0)
            zip_file.writestr("exams.zip", exam_zip_buffer.getvalue())
            exam_zip_buffer.close()

        # Students (single file)
        if options.get("students"):
            document = Document()
            document.add_heading(f"{course.code} - Student Roster", 0)
            document.add_paragraph(
                f"Total Students: {course.enrolled_students.count()}"
            )

            # Create table
            table = document.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Student ID"
            hdr_cells[1].text = "Display Name"
            hdr_cells[2].text = "Legal Name"
            hdr_cells[3].text = "Email"
            hdr_cells[4].text = "Section"
            hdr_cells[5].text = "Status"

            for student in course.enrolled_students.all():
                row_cells = table.add_row().cells
                if options.get("anonymize_students"):
                    row_cells[0].text = f"ANON_{student.id}"
                    row_cells[1].text = f"Student {student.id}"
                    row_cells[2].text = f"Student {student.id}"
                    row_cells[3].text = f"student{student.id}@anonymous.edu"
                else:
                    row_cells[0].text = str(student.student_id)
                    display_name = (
                        student.effective_name
                        if student.preferred_name
                        else student.name
                    )
                    row_cells[1].text = str(display_name)
                    row_cells[2].text = str(student.name)
                    row_cells[3].text = str(student.email or "")
                row_cells[4].text = str(student.section or "")
                row_cells[5].text = "Active" if student.is_active else "Inactive"

            buffer = io.BytesIO()
            document.save(buffer)
            buffer.flush()
            buffer.seek(0)
            zip_file.writestr("student_roster.docx", buffer.getvalue())
            buffer.close()

        # Results - NEW NESTED STRUCTURE
        if options.get("results"):
            from results.models import ExamResult

            # Create a nested ZIP for results
            results_zip_buffer = io.BytesIO()
            with zipfile.ZipFile(
                results_zip_buffer, "w", zipfile.ZIP_DEFLATED
            ) as results_zip:
                for exam in course.exams.all():
                    results = list(
                        ExamResult.objects.filter(exam=exam).select_related("student")
                    )
                    if results:
                        safe_exam_title = "".join(
                            c for c in exam.title if c.isalnum() or c in (" ", "-", "_")
                        ).rstrip()

                        # Create results overview DOCX
                        document = Document()
                        document.add_heading(
                            f"{course.code} - {exam.title} Results Overview", 0
                        )

                        # Overall statistics
                        scores = [r.score for r in results]
                        avg_score = sum(scores) / len(scores) if scores else 0
                        document.add_paragraph(f"Total Students: {len(results)}")
                        document.add_paragraph(f"Average Score: {avg_score:.1f}%")
                        document.add_paragraph(
                            f"Highest Score: {max(scores):.1f}%" if scores else "N/A"
                        )
                        document.add_paragraph(
                            f"Lowest Score: {min(scores):.1f}%" if scores else "N/A"
                        )

                        # Results by variant
                        document.add_heading("Results by Variant:", level=2)
                        variants = {}
                        for result in results:
                            variant_label = (
                                result.variant.version_label
                                if result.variant
                                else "No Variant"
                            )
                            if variant_label not in variants:
                                variants[variant_label] = []
                            variants[variant_label].append(result.score)

                        for variant_label, variant_scores in sorted(variants.items()):
                            avg = (
                                sum(variant_scores) / len(variant_scores)
                                if variant_scores
                                else 0
                            )
                            document.add_paragraph(
                                f"Variant {variant_label}: {len(variant_scores)} students, Average: {avg:.1f}%",
                                style="List Bullet",
                            )

                        buffer = io.BytesIO()
                        document.save(buffer)
                        buffer.flush()
                        buffer.seek(0)
                        results_zip.writestr(
                            f"{safe_exam_title}/results_overview.docx",
                            buffer.getvalue(),
                        )
                        buffer.close()

                        # Create detailed results DOCX for each variant
                        for variant_label, variant_results in variants.items():
                            variant_result_objects = [
                                r
                                for r in results
                                if (
                                    r.variant.version_label
                                    if r.variant
                                    else "No Variant"
                                )
                                == variant_label
                            ]

                            document = Document()
                            document.add_heading(f"{course.code} - {exam.title}", 0)
                            document.add_heading(
                                f"Variant {variant_label} Results", level=1
                            )

                            # Variant statistics
                            avg = (
                                sum(variant_results) / len(variant_results)
                                if variant_results
                                else 0
                            )
                            document.add_paragraph(f"Students: {len(variant_results)}")
                            document.add_paragraph(f"Average Score: {avg:.1f}%")

                            # Results table
                            table = document.add_table(rows=1, cols=7)
                            table.style = "Table Grid"
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = "Student ID"
                            hdr_cells[1].text = "Display Name"
                            hdr_cells[2].text = "Legal Name"
                            hdr_cells[3].text = "Score"
                            hdr_cells[4].text = "Correct"
                            hdr_cells[5].text = "Incorrect"
                            hdr_cells[6].text = "Unanswered"

                            for result in variant_result_objects:
                                row_cells = table.add_row().cells
                                if options.get("anonymize_students"):
                                    row_cells[0].text = f"ANON_{result.student.id}"
                                    row_cells[1].text = f"Student {result.student.id}"
                                    row_cells[2].text = f"Student {result.student.id}"
                                else:
                                    row_cells[0].text = str(result.student.student_id)
                                    if result.student.preferred_name:
                                        display_name = (
                                            result.student.effective_name
                                            or result.student.name
                                        )
                                    else:
                                        display_name = result.student.name
                                    row_cells[1].text = str(display_name)
                                    row_cells[2].text = str(result.student.name)
                                row_cells[3].text = f"{result.score}%"
                                row_cells[4].text = str(result.correct_answers)
                                row_cells[5].text = str(result.incorrect_answers)
                                row_cells[6].text = str(result.unanswered)

                            buffer = io.BytesIO()
                            document.save(buffer)
                            buffer.flush()
                            buffer.seek(0)
                            results_zip.writestr(
                                f"{safe_exam_title}/variant_{variant_label}_results.docx",
                                buffer.getvalue(),
                            )
                            buffer.close()

            # Close and flush results ZIP
            results_zip.close()
            results_zip_buffer.flush()
            results_zip_buffer.seek(0)

            # Add the nested ZIP to main ZIP
            zip_file.writestr("exam_results.zip", results_zip_buffer.getvalue())
            results_zip_buffer.close()

    def _add_csv_files_to_zip(self, zip_file, course, options):
        """Generate separate CSV files with nested structure"""
        import zipfile

        # Question Banks - nested ZIP
        if options.get("question_banks") and course.question_banks.exists():
            qb_zip_buffer = io.BytesIO()
            with zipfile.ZipFile(qb_zip_buffer, "w", zipfile.ZIP_DEFLATED) as qb_zip:
                for bank in course.question_banks.all():
                    string_buffer = io.StringIO()
                    writer = csv.writer(string_buffer)
                    writer.writerow(
                        [
                            "Question #",
                            "Prompt",
                            "Choice A",
                            "Choice B",
                            "Choice C",
                            "Choice D",
                            "Correct Answer",
                            "Difficulty",
                            "Tags",
                            "Explanation",
                        ]
                    )

                    for i, question in enumerate(bank.questions.all(), 1):
                        # FIX: Convert all to strings
                        writer.writerow(
                            [
                                str(i),
                                question.prompt,
                                question.choices.get("A", ""),
                                question.choices.get("B", ""),
                                question.choices.get("C", ""),
                                question.choices.get("D", ""),
                                ", ".join(str(ans) for ans in question.correct_answer),
                                question.difficulty or "",
                                (
                                    ", ".join(str(tag) for tag in question.tags)
                                    if question.tags
                                    else ""
                                ),
                                question.explanation or "",
                            ]
                        )

                    safe_title = "".join(
                        c for c in bank.title if c.isalnum() or c in (" ", "-", "_")
                    ).rstrip()
                    qb_zip.writestr(
                        f"{safe_title}.csv", string_buffer.getvalue().encode("utf-8")
                    )

            qb_zip_buffer.seek(0)
            zip_file.writestr("question_banks.zip", qb_zip_buffer.getvalue())

        # Exams - nested ZIP
        if options.get("exams") and course.exams.exists():
            exam_zip_buffer = io.BytesIO()
            with zipfile.ZipFile(
                exam_zip_buffer, "w", zipfile.ZIP_DEFLATED
            ) as exam_zip:
                for exam in course.exams.all():
                    safe_exam_title = "".join(
                        c for c in exam.title if c.isalnum() or c in (" ", "-", "_")
                    ).rstrip()

                    # Exam overview
                    string_buffer = io.StringIO()
                    writer = csv.writer(string_buffer)
                    writer.writerow(["Property", "Value"])
                    writer.writerow(["Title", exam.title])
                    writer.writerow(["Description", exam.description or ""])
                    writer.writerow(["Time Limit", f"{exam.time_limit} minutes"])
                    writer.writerow(
                        ["Questions per Variant", str(exam.questions_per_variant)]
                    )
                    writer.writerow(["Exam Type", exam.exam_type])
                    writer.writerow(
                        [
                            "Randomize Questions",
                            "Yes" if exam.randomize_questions else "No",
                        ]
                    )
                    writer.writerow(
                        ["Randomize Choices", "Yes" if exam.randomize_choices else "No"]
                    )
                    writer.writerow(["Total Variants", str(exam.variants.count())])
                    exam_zip.writestr(
                        f"{safe_exam_title}/exam_config.csv",
                        string_buffer.getvalue().encode("utf-8"),
                    )

                    # Each variant
                    for variant in exam.variants.all():
                        string_buffer = io.StringIO()
                        writer = csv.writer(string_buffer)
                        writer.writerow(
                            [
                                "Question #",
                                "Prompt",
                                "Choice A",
                                "Choice B",
                                "Choice C",
                                "Choice D",
                                "Correct Answer",
                            ]
                        )

                        for vq in variant.variantquestion_set.select_related(
                            "question"
                        ).order_by("order"):
                            question = vq.question
                            choices = (
                                vq.randomized_choices
                                if vq.randomized_choices
                                else question.choices
                            )
                            correct = (
                                vq.randomized_correct_answer
                                if vq.randomized_correct_answer
                                else question.correct_answer
                            )

                            # FIX: Convert all to strings
                            writer.writerow(
                                [
                                    str(vq.order),
                                    question.prompt,
                                    choices.get("A", ""),
                                    choices.get("B", ""),
                                    choices.get("C", ""),
                                    choices.get("D", ""),
                                    ", ".join(str(ans) for ans in correct),
                                ]
                            )

                        exam_zip.writestr(
                            f"{safe_exam_title}/variant_{variant.version_label}.csv",
                            string_buffer.getvalue().encode("utf-8"),
                        )

            exam_zip_buffer.seek(0)
            zip_file.writestr("exams.zip", exam_zip_buffer.getvalue())

        # Students (single file)
        if options.get("students"):
            string_buffer = io.StringIO()
            writer = csv.writer(string_buffer)
            writer.writerow(
                [
                    "Student ID",
                    "Name",
                    "Preferred Name",
                    "Email",
                    "Section",
                    "Status",
                    "Enrolled At",
                ]
            )

            for student in course.enrolled_students.all():
                if options.get("anonymize_students"):
                    writer.writerow(
                        [
                            f"ANON_{student.id}",
                            f"Student {student.id}",
                            "",
                            f"student{student.id}@anonymous.edu",
                            student.section or "",
                            "Active" if student.is_active else "Inactive",
                            student.enrolled_at.isoformat(),
                        ]
                    )
                else:
                    writer.writerow(
                        [
                            student.student_id,
                            student.name,
                            student.preferred_name or "",
                            student.email or "",
                            student.section or "",
                            "Active" if student.is_active else "Inactive",
                            student.enrolled_at.isoformat(),
                        ]
                    )

            zip_file.writestr(
                "student_roster.csv", string_buffer.getvalue().encode("utf-8")
            )

        # Results - NEW NESTED STRUCTURE
        if options.get("results"):
            from results.models import ExamResult

            # Create a nested ZIP for results
            results_zip_buffer = io.BytesIO()
            with zipfile.ZipFile(
                results_zip_buffer, "w", zipfile.ZIP_DEFLATED
            ) as results_zip:
                for exam in course.exams.all():
                    results = list(
                        ExamResult.objects.filter(exam=exam).select_related("student")
                    )
                    if results:
                        safe_exam_title = "".join(
                            c for c in exam.title if c.isalnum() or c in (" ", "-", "_")
                        ).rstrip()

                        # Create results overview CSV
                        string_buffer = io.StringIO()
                        writer = csv.writer(string_buffer)
                        writer.writerow(["Metric", "Value"])

                        scores = [r.score for r in results]
                        avg_score = sum(scores) / len(scores) if scores else 0
                        writer.writerow(["Total Students", str(len(results))])
                        writer.writerow(["Average Score", f"{avg_score:.1f}%"])
                        writer.writerow(
                            [
                                "Highest Score",
                                f"{max(scores):.1f}%" if scores else "N/A",
                            ]
                        )
                        writer.writerow(
                            ["Lowest Score", f"{min(scores):.1f}%" if scores else "N/A"]
                        )

                        # Add variant breakdown
                        writer.writerow([])  # Empty row
                        writer.writerow(["Variant", "Students", "Average Score"])

                        variants = {}
                        for result in results:
                            variant_label = (
                                result.variant.version_label
                                if result.variant
                                else "No Variant"
                            )
                            if variant_label not in variants:
                                variants[variant_label] = []
                            variants[variant_label].append(result.score)

                        for variant_label, variant_scores in sorted(variants.items()):
                            avg = (
                                sum(variant_scores) / len(variant_scores)
                                if variant_scores
                                else 0
                            )
                            writer.writerow(
                                [variant_label, str(len(variant_scores)), f"{avg:.1f}%"]
                            )

                        results_zip.writestr(
                            f"{safe_exam_title}/results_overview.csv",
                            string_buffer.getvalue().encode("utf-8"),
                        )

                        # Create detailed results CSV for each variant
                        for variant_label in variants.keys():
                            variant_result_objects = [
                                r
                                for r in results
                                if (
                                    r.variant.version_label
                                    if r.variant
                                    else "No Variant"
                                )
                                == variant_label
                            ]

                            string_buffer = io.StringIO()
                            writer = csv.writer(string_buffer)
                            writer.writerow(
                                [
                                    "Student ID",
                                    "Display Name",
                                    "Legal Name",
                                    "Score",
                                    "Correct",
                                    "Incorrect",
                                    "Unanswered",
                                    "Date",
                                ]
                            )

                            for result in variant_result_objects:
                                if options.get("anonymize_students"):
                                    student_id = f"ANON_{result.student.id}"
                                    display_name = f"Student {result.student.id}"
                                    legal_name = f"Student {result.student.id}"
                                else:
                                    student_id = result.student.student_id
                                    if result.student.preferred_name:
                                        display_name = (
                                            result.student.effective_name
                                            or result.student.name
                                        )
                                    else:
                                        display_name = result.student.name
                                    legal_name = result.student.name

                                writer.writerow(
                                    [
                                        student_id,
                                        display_name,
                                        legal_name,
                                        str(result.score),
                                        str(result.correct_answers),
                                        str(result.incorrect_answers),
                                        str(result.unanswered),
                                        (
                                            result.submitted_at.isoformat()
                                            if hasattr(result, "submitted_at")
                                            else ""
                                        ),
                                    ]
                                )

                            results_zip.writestr(
                                f"{safe_exam_title}/variant_{variant_label}_results.csv",
                                string_buffer.getvalue().encode("utf-8"),
                            )

            # Add the nested ZIP to main ZIP
            results_zip_buffer.seek(0)
            zip_file.writestr("exam_results.zip", results_zip_buffer.getvalue())


class IsCourseInstructor(BasePermission):
    def has_permission(self, request, view):
        course_id = view.kwargs.get("course_pk")
        if not course_id:
            return False
        from .models import Course, CourseInstructor

        try:
            course = Course.objects.get(pk=course_id)
        except Course.DoesNotExist:
            return False
        is_instructor = CourseInstructor.objects.filter(
            course=course, user=request.user, accepted=True
        ).exists()
        if not is_instructor:
            raise PermissionDenied(
                "Only instructors can view or manage students in this course."
            )
        return True


class StudentViewSet(viewsets.ModelViewSet):
    """
    ViewSet for Student model
    Provides CRUD operations for managing students in courses,
    plus CSV import/export.
    """

    serializer_class = StudentSerializer
    permission_classes = [IsAuthenticated, IsCourseInstructor]

    def get_queryset(self):
        """Filter students by course and only allow instructors to see them."""
        course_id = self.kwargs.get("course_pk")
        if not course_id:
            return Student.objects.none()

        course = get_object_or_404(Course, pk=course_id)

        # Check if user is an instructor using the CourseInstructor model
        is_instructor = CourseInstructor.objects.filter(
            course=course, user=self.request.user, accepted=True
        ).exists()

        if not is_instructor:
            raise PermissionDenied(
                "Only instructors can view or manage students in this course."
            )

        return Student.objects.filter(course=course)

    def list(self, request, *args, **kwargs):
        """List students with proper permission checking and structured response."""
        from rest_framework.exceptions import PermissionDenied

        course_id = self.kwargs.get("course_pk")
        if not course_id:
            return Response(
                {"error": "Course not found"}, status=status.HTTP_404_NOT_FOUND
            )

        course = get_object_or_404(Course, pk=course_id)
        is_instructor = CourseInstructor.objects.filter(
            course=course, user=request.user, accepted=True
        ).exists()
        if not is_instructor:
            raise PermissionDenied(
                "Only instructors can view or manage students in this course."
            )

        # Get students for this course
        students = Student.objects.filter(course=course)
        serializer = self.get_serializer(students, many=True)
        # Return structured response with metadata
        return Response(
            {
                "students": serializer.data,
                "count": students.count(),
                "course_id": course.id,
                "course_name": course.name,
            }
        )

    @action(
        detail=False,
        methods=["post"],
        parser_classes=[MultiPartParser],
        permission_classes=[IsAuthenticated],
    )
    def import_csv(self, request, course_pk=None):
        """
        POST /courses/{course_pk}/students/import_csv/
        Accepts a form-file field named 'file' containing CSV with headers:
        student_id, name, preferred_name, email, section, is_active
        """
        course = get_object_or_404(Course, pk=course_pk)

        # Check using CourseInstructor model
        if not CourseInstructor.objects.filter(
            course=course, user=request.user, accepted=True
        ).exists():
            return Response(status=status.HTTP_403_FORBIDDEN)

        uploaded = request.FILES.get("file")
        if not uploaded:
            return Response(
                {"detail": "No file provided."}, status=status.HTTP_400_BAD_REQUEST
            )

        text = io.TextIOWrapper(uploaded.file, encoding="utf-8")
        reader = csv.DictReader(text)
        created = []

        for row in reader:
            # Map CSV headers to model fields
            # Handle both lowercase and capitalized headers
            student_id_value = (
                row.get("student_id", "")
                or row.get("Student ID", "")
                or row.get("studentid", "")
            ).strip()

            name_value = (row.get("name", "") or row.get("Name", "")).strip()

            # NEW: Handle preferred_name with multiple possible header formats
            preferred_name_value = (
                row.get("preferred_name", "")
                or row.get("Preferred Name", "")
                or row.get("PreferredName", "")
                or row.get("preferred name", "")
            ).strip() or None  # Set to None if empty string

            email_value = (row.get("email", "") or row.get("Email", "")).strip()

            section_value = (
                row.get("section", "") or row.get("Section", "")
            ).strip() or None

            is_active_value = (
                row.get("is_active", "true") or row.get("Is Active", "true")
            ).lower() in ("true", "1", "yes")

            student_obj, _ = Student.objects.update_or_create(
                course=course,
                student_id=student_id_value,  # Use student_id, not display_id
                defaults={
                    "name": name_value,  # Use name, not display_name
                    "preferred_name": preferred_name_value,  # NEW: Add preferred_name
                    "email": email_value,
                    "section": section_value,
                    "is_active": is_active_value,
                },
            )
            created.append(StudentSerializer(student_obj).data)

            # Log the bulk import activity
            if created:
                CourseActivity.objects.create(
                    course=course,
                    user=request.user,
                    activity_type=CourseActivity.ActivityType.RESULTS_IMPORTED,  # or create STUDENTS_IMPORTED
                    description=f"{request.user.name} imported {len(created)} students via CSV",
                    entity_type="bulk_import",
                    entity_id=None,
                )

        return Response(created, status=status.HTTP_201_CREATED)

    @action(detail=False, methods=["get"], permission_classes=[IsAuthenticated])
    def export_csv(self, request, course_pk=None):
        """
        GET /courses/{course_pk}/students/export_csv/
        Returns a CSV download of all students in the course.
        Now includes preferred_name in the export.
        """
        course = get_object_or_404(Course, pk=course_pk)

        # Check using CourseInstructor model
        if not CourseInstructor.objects.filter(
            course=course, user=request.user, accepted=True
        ).exists():
            return Response(status=status.HTTP_403_FORBIDDEN)

        queryset = Student.objects.filter(course=course)

        # Updated header to include Preferred Name
        header = [
            "Student ID",
            "Name",
            "Preferred Name",
            "Email",
            "Section",
            "Is Active",
            "Enrolled At",
        ]

        rows = [
            [
                s.display_id,
                s.name,  # Always export the legal name
                s.preferred_name
                or "",  # NEW: Export preferred_name, empty string if None
                s.email or "",
                s.section or "",
                "true" if s.is_active else "false",
                s.enrolled_at.isoformat(),
            ]
            for s in queryset
        ]

        response = HttpResponse(content_type="text/csv")
        response["Content-Disposition"] = (
            f"attachment; filename=students_{course_pk}.csv"
        )
        writer = csv.writer(response)
        writer.writerow(header)
        writer.writerows(rows)

        return response

    @action(detail=False, methods=["delete"], permission_classes=[IsAuthenticated])
    def delete_all(self, request, course_pk=None):
        """
        DELETE /courses/{course_pk}/students/delete_all/
        Deletes all students in the course.
        """
        course = get_object_or_404(Course, pk=course_pk)

        # Check permissions using CourseInstructor
        if not CourseInstructor.objects.filter(
            course=course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "Only course instructors can delete all students"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Delete all students in the course
        deleted_count = Student.objects.filter(course=course).delete()[0]

        return Response(
            {
                "message": (
                    f"Successfully deleted {deleted_count} students from "
                    f"course {course.code}"
                ),
                "deleted_count": deleted_count,
            },
            status=status.HTTP_200_OK,
        )

    def destroy(self, request, *args, **kwargs):
        """Delete a student with activity tracking"""
        student = self.get_object()
        course_id = self.kwargs.get("course_pk")
        course = get_object_or_404(Course, pk=course_id)

        # Store student info before deletion
        student_name = student.name
        student_id = student.id

        response = super().destroy(request, *args, **kwargs)

        # Log the activity
        CourseActivity.objects.create(
            course=course,
            user=request.user,
            activity_type=CourseActivity.ActivityType.STUDENT_REMOVED,
            description=f"{request.user.name} removed student '{student_name}'",
            entity_type="student",
            entity_id=student_id,
        )

        return response

    @action(detail=False, methods=["post"], permission_classes=[IsAuthenticated])
    def anonymize_all(self, request, course_pk=None):
        """
        POST /courses/{course_pk}/students/anonymize_all/
        Sets is_anonymous=True for all students in the course.
        """
        course = get_object_or_404(Course, pk=course_pk)

        # Check permissions using CourseInstructor
        if not CourseInstructor.objects.filter(
            course=course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "Only course instructors can anonymize all students"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Update all students to be anonymous
        updated_count = Student.objects.filter(course=course).update(is_anonymous=True)

        return Response(
            {
                "message": (
                    f"Successfully anonymized {updated_count} students in "
                    f"course {course.code}"
                ),
                "anonymized_count": updated_count,
            },
            status=status.HTTP_200_OK,
        )

    @action(detail=False, methods=["post"], permission_classes=[IsAuthenticated])
    def deanonymize_all(self, request, course_pk=None):
        """
        POST /courses/{course_pk}/students/deanonymize_all/
        Sets is_anonymous=False for all students in the course.
        """
        course = get_object_or_404(Course, pk=course_pk)

        # Check permissions using CourseInstructor
        if not CourseInstructor.objects.filter(
            course=course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "Only course instructors can deanonymize all students"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Update all students to not be anonymous
        updated_count = Student.objects.filter(course=course).update(is_anonymous=False)

        return Response(
            {
                "message": (
                    f"Successfully deanonymized {updated_count} students in "
                    f"course {course.code}"
                ),
                "deanonymized_count": updated_count,
            },
            status=status.HTTP_200_OK,
        )

    @transaction.atomic
    def create(self, request, *args, **kwargs):
        """Create a new student with proper error handling for duplicates."""
        course_id = self.kwargs.get("course_pk")
        course = get_object_or_404(Course, pk=course_id)
        from django.db import IntegrityError

        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        try:
            student = serializer.save(course=course)

            # Log the activity
            CourseActivity.objects.create(
                course=course,
                user=request.user,
                activity_type=CourseActivity.ActivityType.STUDENT_ADDED,
                description=f"{request.user.name} added student '{student.name}'",
                entity_type="student",
                entity_id=student.id,
            )
        except IntegrityError:
            transaction.set_rollback(True)
            return Response(
                {
                    "detail": "A student with this ID or user already exists in this course (unique constraint)."
                },
                status=status.HTTP_400_BAD_REQUEST,
            )
        headers = self.get_success_headers(serializer.data)
        return Response(
            serializer.data, status=status.HTTP_201_CREATED, headers=headers
        )

    def perform_create(self, serializer):
        """Assign the course when creating a student."""
        course_id = self.kwargs.get("course_pk")
        course = get_object_or_404(Course, pk=course_id)
        serializer.save(course=course)


class InstructorCoursesAPIView(APIView):
    """Get all courses for the authenticated instructor"""

    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Get courses where user is an instructor (via CourseInstructor)
        course_ids = CourseInstructor.objects.filter(
            user=request.user, accepted=True
        ).values_list("course_id", flat=True)

        courses = Course.objects.filter(id__in=course_ids)
        serializer = CourseSerializer(courses, many=True)
        return Response(serializer.data)


class CourseDetailAPIView(APIView):
    """Get detailed information about a specific course"""

    permission_classes = [IsAuthenticated]

    def get(self, request, pk):
        course = get_object_or_404(Course, pk=pk)
        serializer = CourseDetailSerializer(course)
        return Response(serializer.data)

    def post(self, request):
        # need to add the user pk to the data
        data = request.data  # the current data from the frontend

        serializer = CourseDetailSerializer(data=data)

        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        else:
            return Response(serializer.errors, status=400)

    def put(self, request, pk):
        data = Course.objects.get(pk=pk)
        serializer = CourseDetailSerializer(data, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk):
        data = Course.objects.get(pk=pk)
        data.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


class CourseStudentsAPIView(APIView):
    """Get students enrolled in a course"""

    permission_classes = [IsAuthenticated]

    def get(self, request, pk):
        course = get_object_or_404(Course, pk=pk)

        # Check using CourseInstructor model
        if not CourseInstructor.objects.filter(
            course=course, user=request.user, accepted=True
        ).exists():
            return Response(
                {"error": "Only instructors can view student list"},
                status=status.HTTP_403_FORBIDDEN,
            )

        students = Student.objects.filter(course=course, is_active=True)
        serializer = StudentSerializer(students, many=True)

        return Response(
            {
                "course": CourseSerializer(course).data,
                "student_count": len(serializer.data),
                "students": serializer.data,
            }
        )


class PendingInvitesView(APIView):
    """Get all pending course invites for the authenticated user"""

    permission_classes = [IsAuthenticated]

    def get(self, request):
        pending_invites = CourseInstructor.objects.filter(
            user=request.user, accepted=False
        ).select_related("course", "course__creator")

        data = []
        for invite in pending_invites:
            # Find who invited them (the MAIN instructor)
            main_instructor = CourseInstructor.objects.filter(
                course=invite.course, role=CourseInstructor.Role.MAIN
            ).first()

            data.append(
                {
                    "id": invite.id,
                    "course_id": invite.course.id,
                    "course_code": invite.course.code,
                    "course_title": invite.course.name,
                    "inviter_name": (
                        main_instructor.user.name if main_instructor else "Unknown"
                    ),
                    "inviter_email": (
                        main_instructor.user.email
                        if main_instructor
                        else "unknown@example.com"
                    ),
                    "role": invite.role,
                    "permissions": invite.access,
                    "created_at": invite.invited_at.isoformat(),  # Use invited_at, not created_at
                }
            )

        return Response(data)


# Fix 3: Update AcceptInviteView to handle already accepted case
class AcceptInviteView(APIView):
    """Accept a course invitation"""

    permission_classes = [IsAuthenticated]

    def post(self, request, pk):
        try:
            invite = CourseInstructor.objects.get(id=pk, user=request.user)
        except CourseInstructor.DoesNotExist:
            return Response(
                {"error": "Invite not found"}, status=status.HTTP_404_NOT_FOUND
            )

        # Check if already accepted
        if invite.accepted:
            return Response(
                {"error": "Invite already accepted"}, status=status.HTTP_400_BAD_REQUEST
            )

        invite.accepted = True
        invite.save()

        return Response({"status": "Invite accepted"})


class DeclineInviteView(APIView):
    """Decline a course invitation"""

    permission_classes = [IsAuthenticated]

    def post(self, request, pk):
        invite = get_object_or_404(
            CourseInstructor, id=pk, user=request.user, accepted=False
        )

        # Remove the invite
        invite.delete()

        return Response({"status": "Invite declined"})


class PreviewImportQuestions(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, course_id):
        try:
            course = Course.objects.get(pk=course_id)
        except Course.DoesNotExist:
            return Response({"error": "Course not found"}, status=404)

        serializer = ImportedQuestionSerializer(
            data=request.data.get("questions", []), many=True
        )
        serializer.is_valid(raise_exception=True)

        banks = QuestionBank.objects.filter(course=course)

        # ✅ Get prompts as strings and lowercase them
        existing_prompts = {
            p.lower()
            for p in Question.objects.filter(bank__in=banks).values_list(
                "prompt", flat=True
            )
        }

        result = []
        for q in serializer.validated_data:
            result.append(
                {
                    **q,
                    "is_duplicate": q.get("prompt", "").strip().lower()
                    in existing_prompts,
                }
            )

        return Response(result, status=200)
