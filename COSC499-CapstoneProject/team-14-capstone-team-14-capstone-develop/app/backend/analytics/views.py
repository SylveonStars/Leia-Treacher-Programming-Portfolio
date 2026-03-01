from datetime import timedelta

# analytics/views.py
from django.db.models import Avg
from django.http import Http404, HttpResponse
from django.shortcuts import get_object_or_404
from django.utils import timezone
import numpy as np
from rest_framework import status
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView
from scipy import stats

from analytics.models import SimilarityFlag
from courses.models import Course, CourseInstructor, Student
from exams.models import Exam, Variant, VariantQuestion
from questions.models import Question, QuestionBank
from results.models import ExamResult


class VariantSetAnalyticsView(APIView):
    """Get comprehensive analytics for a specific variant set"""

    permission_classes = [IsAuthenticated]

    def get(self, request, exam_id, variant_ids):
        try:
            # Parse variant IDs
            variant_id_list = [int(vid) for vid in variant_ids.split(",")]

            # Get the exam
            exam = get_object_or_404(Exam, id=exam_id)

            # Check permissions
            if not CourseInstructor.objects.filter(
                course=exam.course, user=request.user, accepted=True
            ).exists():
                return Response(
                    {"error": "You don't have permission to view this exam"},
                    status=status.HTTP_403_FORBIDDEN,
                )

            # Get variants
            variants = Variant.objects.filter(id__in=variant_id_list, exam=exam)

            if not variants.exists():
                return Response(
                    {"error": "No valid variants found"},
                    status=status.HTTP_404_NOT_FOUND,
                )

            # Calculate analytics
            analytics_data = self._calculate_variant_set_analytics(exam, variants)

            return Response(analytics_data)

        except Exception as e:
            return Response(
                {"error": f"Failed to calculate analytics: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

    def _calculate_variant_set_analytics(self, exam, variants):
        """Calculate comprehensive analytics for a variant set"""

        # Get all questions from all variants
        all_questions = []
        question_counts = {}

        for variant in variants:
            variant_questions = variant.questions.all()
            all_questions.extend(variant_questions)

            # Count questions per variant
            question_counts[variant.id] = variant_questions.count()

        total_questions = len(all_questions)
        unique_questions = len(set(q.id for q in all_questions))

        # Calculate question diversity
        question_diversity = {
            "total_questions": total_questions,
            "unique_questions": unique_questions,
            "diversity_ratio": (
                unique_questions / total_questions if total_questions > 0 else 0
            ),
            "diversity_percentage": (
                (unique_questions / total_questions * 100) if total_questions > 0 else 0
            ),
            "reuse_rate": (
                ((total_questions - unique_questions) / total_questions * 100)
                if total_questions > 0
                else 0
            ),
        }

        # Calculate difficulty distribution
        difficulty_counts = {}
        for question in all_questions:
            difficulty = getattr(question, "difficulty", "medium")
            difficulty_counts[difficulty] = difficulty_counts.get(difficulty, 0) + 1

        total_difficulty_questions = sum(difficulty_counts.values())
        difficulty_percentages = {
            difficulty: (
                (count / total_difficulty_questions * 100)
                if total_difficulty_questions > 0
                else 0
            )
            for difficulty, count in difficulty_counts.items()
        }

        # Calculate balance score (how evenly distributed difficulties are)
        if len(difficulty_counts) > 1:
            balance_score = (
                100
                - max(difficulty_percentages.values())
                + min(difficulty_percentages.values())
            )
        else:
            balance_score = 100  # Perfect balance if only one difficulty level

        difficulty_distribution = {
            "counts": difficulty_counts,
            "percentages": difficulty_percentages,
            "balance_score": balance_score,
            "total_questions": total_difficulty_questions,
        }

        # Calculate answer pattern analysis (Hamming distance simulation)
        variant_count = variants.count()
        if variant_count > 1:
            # Simulate Hamming distances between variants
            hamming_distances = [85, 78, 92, 88, 75]  # Mock data
            hamming_percentage = np.mean(hamming_distances)
            pattern_diversity = 100 - (
                hamming_percentage / 100 * 50
            )  # Higher diversity = lower average distance
        else:
            hamming_distances = []
            hamming_percentage = 0
            pattern_diversity = 100

        answer_pattern_analysis = {
            "hamming_distance": hamming_percentage,
            "hamming_percentage": hamming_percentage,
            "pattern_diversity": pattern_diversity,
            "risk_score": max(0, 100 - pattern_diversity),
            "unique_patterns": variant_count,
            "total_variants": variant_count,
        }

        # Calculate integrity score
        diversity_score = question_diversity["diversity_percentage"]
        balance_score = difficulty_distribution["balance_score"]
        pattern_score = answer_pattern_analysis["pattern_diversity"]

        # Calculate penalties
        variant_count_penalty = variant_count < 3  # Penalty for too few variants
        reuse_penalty = question_diversity["reuse_rate"] > 30  # Penalty for high reuse

        # Calculate final integrity score
        base_score = (diversity_score + balance_score + pattern_score) / 3

        # Apply penalties
        if variant_count_penalty:
            base_score -= 15
        if reuse_penalty:
            base_score -= 10

        integrity_score = max(0, min(100, base_score))

        # Determine grade and color
        if integrity_score >= 90:
            grade = "Excellent"
            color = "green"
        elif integrity_score >= 75:
            grade = "Good"
            color = "blue"
        elif integrity_score >= 60:
            grade = "Moderate"
            color = "yellow"
        else:
            grade = "Poor"
            color = "red"

        integrity_data = {
            "score": round(integrity_score, 1),
            "grade": grade,
            "color": color,
            "components": {
                "diversity_score": round(diversity_score, 1),
                "balance_score": round(balance_score, 1),
                "pattern_score": round(pattern_score, 1),
            },
            "penalties": {
                "variant_count_penalty": variant_count_penalty,
                "reuse_penalty": reuse_penalty,
            },
        }

        return {
            "integrity_score": integrity_data,
            "question_diversity": question_diversity,
            "difficulty_distribution": difficulty_distribution,
            "answer_pattern_analysis": answer_pattern_analysis,
            "question_reuse_rate": question_diversity["reuse_rate"],
            "mandatory_overlap": 0,  # Placeholder
            "hamming_distances": hamming_distances,
            "variant_count": variant_count,
            "total_questions": total_questions,
            "unique_questions": unique_questions,
            "metadata": {
                "exam_id": exam.id,
                "exam_title": exam.title,
                "variant_ids": [v.id for v in variants],
                "variant_labels": [f"Variant {i+1}" for i in range(variant_count)],
                "calculated_at": timezone.now().isoformat(),
            },
        }


class AnalyticsHealthView(APIView):
    """Health check endpoint for analytics"""

    permission_classes = []

    def get(self, request):
        return Response(
            {
                "status": "healthy",
                "message": "Analytics service is running",
                "timestamp": timezone.now().isoformat(),
            }
        )


class InstructorOverviewView(APIView):
    """Get instructor overview analytics"""

    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            user = request.user

            # Get instructor's courses
            instructor_courses = CourseInstructor.objects.filter(
                user=user, accepted=True
            ).values_list("course_id", flat=True)

            # Calculate overview statistics
            total_courses = Course.objects.filter(id__in=instructor_courses).count()
            total_students = Student.objects.filter(
                course_id__in=instructor_courses
            ).count()
            total_exams = Exam.objects.filter(course_id__in=instructor_courses).count()

            # Calculate average grade
            exam_results = ExamResult.objects.filter(
                exam__course_id__in=instructor_courses
            )
            average_grade = exam_results.aggregate(avg=Avg("score"))["avg"] or 0

            overview = {
                "total_courses": total_courses,
                "total_students": total_students,
                "total_exams": total_exams,
                "average_grade": round(average_grade, 1),
            }

            # Get top performing courses
            courses = Course.objects.filter(id__in=instructor_courses)
            top_courses = []

            for course in courses[:5]:  # Top 5 courses
                course_results = ExamResult.objects.filter(exam__course=course)
                if course_results.exists():
                    avg_score = course_results.aggregate(avg=Avg("score"))["avg"] or 0
                    student_count = Student.objects.filter(course=course).count()
                    exam_count = Exam.objects.filter(course=course).count()

                    top_courses.append(
                        {
                            "id": course.id,
                            "code": course.code,
                            "name": course.name,
                            "avg_score": round(avg_score, 1),
                            "student_count": student_count,
                            "exam_count": exam_count,
                        }
                    )

            # Mock recent activity
            recent_activity = [
                {
                    "type": "exam",
                    "action": "created",
                    "course": "COSC 499",
                    "date": "2025-01-15",
                    "relative_date": "2 days ago",
                },
                {
                    "type": "course",
                    "action": "updated",
                    "course": "COSC 301",
                    "date": "2025-01-14",
                    "relative_date": "3 days ago",
                },
            ]

            # Mock grade trends
            grade_trends = [
                {"month": "Jan", "average": 78.5},
                {"month": "Feb", "average": 82.1},
                {"month": "Mar", "average": 79.8},
                {"month": "Apr", "average": 85.2},
            ]

            return Response(
                {
                    "overview": overview,
                    "top_performing_courses": top_courses,
                    "recent_activity": recent_activity,
                    "grade_trends": grade_trends,
                }
            )

        except Exception as e:
            return Response(
                {"error": f"Failed to get instructor overview: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )


class QuestionAnalyticsView(APIView):
    """Get question-level analytics (FR7.1)"""

    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user

        # Get all questions from instructor's courses
        instructor_courses = CourseInstructor.objects.filter(
            user=user, accepted=True
        ).values_list("course_id", flat=True)

        # Get the most recent 10 exams with results
        recent_exams = (
            Exam.objects.filter(course__in=instructor_courses, results__isnull=False)
            .distinct()
            .order_by("-created_at")[:10]
        )

        most_missed_per_exam = []

        # Also collect data for difficulty vs discrimination
        question_performance_data = {}

        for exam in recent_exams:
            # Get all results for this exam
            exam_results = ExamResult.objects.filter(exam=exam)

            if not exam_results.exists():
                continue

            # Track wrong answers per question
            question_stats = {}

            # INITIALIZE ALL QUESTIONS FIRST
            # Get all questions for this exam through its variants
            for variant in exam.variants.all():
                variant_questions = (
                    VariantQuestion.objects.filter(variant=variant)
                    .select_related("question")
                    .order_by("order")
                )

                for vq in variant_questions:
                    # Use the same question_id format as in the grading details
                    question_id = f"Q{vq.order + 1}"

                    if question_id not in question_stats:
                        question_stats[question_id] = {
                            "total_attempts": 0,
                            "incorrect": 0,
                            "question_text": (
                                vq.question.prompt[:100] + "..."
                                if len(vq.question.prompt) > 100
                                else vq.question.prompt
                            ),
                            "question_number": vq.order + 1,
                            "exam_id": exam.id,
                            "exam_title": exam.title,
                        }

            # Also track performance by student score for discrimination
            student_scores = []
            student_answers = {}  # {student_id: {question_id: correct}}

            for result in exam_results:
                student_id = result.student.id
                student_score = float(result.score) if result.score else 0
                student_scores.append((student_id, student_score))
                student_answers[student_id] = {}

                if result.grading_details and isinstance(result.grading_details, list):
                    for detail in result.grading_details:
                        if not isinstance(detail, dict):
                            continue

                        question_id = detail.get("question_id")
                        if not question_id:
                            question_num = detail.get("question_number")
                            if question_num:
                                question_id = f"Q{question_num}"

                        if question_id:
                            if question_id not in question_stats:
                                question_stats[question_id] = {
                                    "total_attempts": 0,
                                    "incorrect": 0,
                                    "question_text": detail.get("question_text", ""),
                                    "question_number": detail.get("question_number", 0),
                                    "exam_id": exam.id,
                                    "exam_title": exam.title,
                                }

                            question_stats[question_id]["total_attempts"] += 1

                            # Check if answer was incorrect
                            status = detail.get("status", "").lower()
                            is_correct = status == "correct"

                            if not is_correct:
                                question_stats[question_id]["incorrect"] += 1

                            # Track for discrimination calculation
                            student_answers[student_id][question_id] = is_correct

            # Calculate discrimination for each question
            if (
                len(student_scores) >= 4
            ):  # Need at least 4 students for meaningful discrimination
                # Sort students by total score
                sorted_students = sorted(
                    student_scores, key=lambda x: x[1], reverse=True
                )

                # Get top 27% and bottom 27% (standard practice in item analysis)
                n_group = max(1, int(len(sorted_students) * 0.27))
                top_students = [s[0] for s in sorted_students[:n_group]]
                bottom_students = [s[0] for s in sorted_students[-n_group:]]

                # Calculate discrimination for each question
                for question_id, stats in question_stats.items():
                    # Count correct answers in top and bottom groups
                    top_correct = sum(
                        1
                        for sid in top_students
                        if student_answers.get(sid, {}).get(question_id, False)
                    )
                    bottom_correct = sum(
                        1
                        for sid in bottom_students
                        if student_answers.get(sid, {}).get(question_id, False)
                    )

                    # Discrimination index = (top_correct - bottom_correct) / n_group
                    discrimination = (
                        (top_correct - bottom_correct) / n_group if n_group > 0 else 0
                    )

                    # Difficulty = proportion who got it wrong
                    difficulty = (
                        stats["incorrect"] / stats["total_attempts"]
                        if stats["total_attempts"] > 0
                        else 0
                    )

                    # Store for scatter plot
                    if question_id not in question_performance_data:
                        question_performance_data[question_id] = {
                            "difficulty": difficulty,
                            "discrimination": discrimination,
                            "question_text": stats["question_text"],
                            "question_number": stats["question_number"],
                            "exam_title": stats["exam_title"],
                            "examId": stats["exam_id"],
                            "total_attempts": stats["total_attempts"],
                        }

            # Find the most missed question for this exam (existing logic)
            if question_stats:
                most_missed = max(
                    question_stats.items(), key=lambda x: x[1]["incorrect"]
                )
                question_id, stats = most_missed

                question_text = stats["question_text"]
                if not question_text and isinstance(question_id, int):
                    try:
                        question = Question.objects.get(id=question_id)
                        question_text = (
                            question.prompt[:100] + "..."
                            if len(question.prompt) > 100
                            else question.prompt
                        )
                    except Question.DoesNotExist:
                        question_text = f"Question {stats['question_number']}"
                elif not question_text:
                    question_text = (
                        f"Question {stats['question_number'] or question_id}"
                    )

                most_missed_per_exam.append(
                    {
                        "examId": exam.id,
                        "examTitle": (
                            exam.title[:30] + "..."
                            if len(exam.title) > 30
                            else exam.title
                        ),
                        "courseCode": exam.course.code,
                        "questionId": question_id,
                        "questionNumber": stats["question_number"] or question_id,
                        "questionText": question_text,
                        "missedCount": stats["incorrect"],
                        "totalAttempts": stats["total_attempts"],
                        "missRate": round(
                            (
                                (stats["incorrect"] / stats["total_attempts"] * 100)
                                if stats["total_attempts"] > 0
                                else 0
                            ),
                            1,
                        ),
                    }
                )

        # Sort by most recent exam first
        most_missed_per_exam.reverse()

        # Prepare question statistics for difficulty vs discrimination scatter plot
        question_stats = []
        for question_id, perf_data in question_performance_data.items():
            # Get question's labeled difficulty if available
            labeled_difficulty = None
            if isinstance(question_id, int):
                try:
                    question = Question.objects.get(id=question_id)
                    # Map integer difficulty values (1=Easy, 2=Medium, 3=Hard)
                    if (
                        hasattr(question, "difficulty")
                        and question.difficulty is not None
                    ):
                        difficulty_map = {
                            1: 0.0,  # Easy
                            2: 0.5,  # Medium
                            3: 1.0,  # Hard
                        }
                        labeled_difficulty = difficulty_map.get(
                            question.difficulty, None
                        )
                except Question.DoesNotExist:
                    pass

            question_stats.append(
                {
                    "questionNumber": (
                        perf_data["question_number"]
                        if perf_data["question_number"]
                        else question_id
                    ),
                    "questionText": (
                        perf_data["question_text"][:50] + "..."
                        if len(perf_data["question_text"]) > 50
                        else perf_data["question_text"]
                    ),
                    "examTitle": perf_data["exam_title"],
                    "examId": perf_data["examId"],
                    "difficulty": round(
                        perf_data["difficulty"], 3
                    ),  # Actual difficulty (% wrong)
                    "labeledDifficulty": labeled_difficulty,  # What we labeled it as
                    "discrimination": round(perf_data["discrimination"], 3),
                    "pointBiserial": 0,  # Not calculated in this simplified version
                    "missedCount": int(
                        perf_data["difficulty"] * perf_data["total_attempts"]
                    ),
                    "correctPercent": round((1 - perf_data["difficulty"]) * 100, 1),
                    "averageTime": 0,  # Not tracked in this version
                }
            )

        # Get total questions count through question banks
        try:
            question_banks = QuestionBank.objects.filter(course__in=instructor_courses)
            total_questions = Question.objects.filter(bank__in=question_banks).count()
        except:
            total_questions = 0

        return Response(
            {
                "questionStatistics": question_stats,
                "totalQuestions": total_questions,
                "examCount": Exam.objects.filter(course__in=instructor_courses).count(),
                "mostMissedPerCourse": most_missed_per_exam,
            }
        )


class GradeDistributionView(APIView):
    """Get grade distribution analytics"""

    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            user = request.user

            # Get instructor's courses
            instructor_courses = CourseInstructor.objects.filter(
                user=user, accepted=True
            ).values_list("course_id", flat=True)

            # Get exam results
            results = ExamResult.objects.filter(exam__course_id__in=instructor_courses)

            # Calculate grade distribution
            distribution = []
            total_results = results.count()

            if total_results > 0:
                # Define grade ranges
                ranges = [
                    ("A (90-100%)", 90, 100),
                    ("B (80-89%)", 80, 89),
                    ("C (70-79%)", 70, 79),
                    ("D (60-69%)", 60, 69),
                    ("F (0-59%)", 0, 59),
                ]

                for range_name, min_score, max_score in ranges:
                    count = results.filter(
                        score__gte=min_score, score__lte=max_score
                    ).count()
                    percentage = (
                        (count / total_results * 100) if total_results > 0 else 0
                    )

                    distribution.append(
                        {
                            "range": range_name,
                            "count": count,
                            "percentage": round(percentage, 1),
                        }
                    )

            return Response(
                {"distribution": distribution, "totalResults": total_results}
            )

        except Exception as e:
            return Response(
                {"error": f"Failed to get grade distribution: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )


class PerformanceMetricsView(APIView):
    """Get performance metrics"""

    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            user = request.user

            # Get instructor's courses
            instructor_courses = CourseInstructor.objects.filter(
                user=user, accepted=True
            ).values_list("course_id", flat=True)

            # Get exam results
            results = ExamResult.objects.filter(exam__course_id__in=instructor_courses)

            scores = list(results.values_list("score", flat=True))

            if scores:
                # Convert Decimal to float for numpy/scipy operations
                scores = np.array(
                    [float(score) for score in scores if score is not None]
                )

                if len(scores) > 0:
                    # Calculate metrics
                    mean = np.mean(scores)
                    median = np.median(scores)
                    std_dev = np.std(scores)

                    # Calculate skewness only if we have enough data points
                    if len(scores) > 2:
                        skewness = stats.skew(scores)
                    else:
                        skewness = 0.0

                    # Calculate reliability (Cronbach's alpha approximation)
                    # This is simplified - real implementation would analyze item responses
                    reliability = 0.85 if std_dev < 15 else 0.75
                else:
                    mean = median = std_dev = skewness = reliability = 0
            else:
                mean = median = std_dev = skewness = reliability = 0

            return Response(
                {
                    "mean": round(mean, 2),
                    "median": round(median, 2),
                    "standardDeviation": round(std_dev, 2),
                    "skewness": round(skewness, 3),
                    "reliability": round(reliability, 3),
                    "totalResults": len(scores),
                }
            )

        except Exception as e:
            return Response(
                {"error": f"Failed to get performance metrics: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )


class SimilarityFlagsView(APIView):
    """Get similarity flags for instructor's exams"""

    permission_classes = [IsAuthenticated]

    def _detect_similarities(self, exam):
        """
        Detect similarities between student answers, considering correctness.
        Only flags suspicious patterns, not legitimate correct answers.
        """
        from exams.models import Variant
        from results.models import ExamResult

        flags = []

        # Get all variants for this exam
        variants = Variant.objects.filter(exam=exam)

        for variant in variants:
            # Get all results for this variant
            results = ExamResult.objects.filter(variant=variant).select_related(
                "student"
            )

            if len(results) < 2:
                continue  # Need at least 2 students to compare

            # Compare each pair of students
            for i, result1 in enumerate(results):
                for j, result2 in enumerate(results[i + 1 :], i + 1):
                    similarity_data = self._calculate_similarity(result1, result2)

                    # Only flag if similarity is suspicious (high similarity + incorrect answers)
                    if similarity_data["is_suspicious"]:
                        flag, created = SimilarityFlag.objects.get_or_create(
                            exam=exam,
                            variant=variant,
                            student1=result1.student,
                            student2=result2.student,
                            defaults={
                                "similarity_score": similarity_data["similarity_score"],
                                "flagged_questions": similarity_data[
                                    "flagged_questions"
                                ],
                                "status": "pending",
                                "severity": similarity_data["severity"],
                            },
                        )

                        if not created:
                            # Update existing flag
                            flag.similarity_score = similarity_data["similarity_score"]
                            flag.flagged_questions = similarity_data[
                                "flagged_questions"
                            ]
                            flag.severity = similarity_data["severity"]
                            flag.save()

                        flags.append(flag)

        return flags

    def _calculate_similarity(self, result1, result2):
        """
        Calculate similarity between two student results.
        Returns similarity data with suspicion analysis.
        """
        import json

        # Parse raw responses
        responses1 = (
            result1.raw_responses
            if isinstance(result1.raw_responses, dict)
            else json.loads(result1.raw_responses or "{}")
        )
        responses2 = (
            result2.raw_responses
            if isinstance(result2.raw_responses, dict)
            else json.loads(result2.raw_responses or "{}")
        )

        # Get grading details to know correct answers
        grading1 = (
            result1.grading_details
            if isinstance(result1.grading_details, dict)
            else json.loads(result1.grading_details or "{}")
        )
        grading2 = (
            result2.grading_details
            if isinstance(result2.grading_details, dict)
            else json.loads(result2.grading_details or "{}")
        )

        # Get variant questions to know correct answers
        variant_questions = result1.variant.variantquestion_set.all().select_related(
            "question"
        )

        total_questions = len(variant_questions)
        matching_answers = 0
        suspicious_matches = 0
        flagged_questions = []

        for vq in variant_questions:
            question_id = str(vq.order + 1)  # Questions are 1-indexed in responses

            answer1 = responses1.get(question_id, "")
            answer2 = responses2.get(question_id, "")

            # Check if answers match
            if answer1 == answer2 and answer1 != "":
                matching_answers += 1

                # Check if this is a suspicious match (both wrong)
                is_correct1 = grading1.get(question_id, {}).get("is_correct", False)
                is_correct2 = grading2.get(question_id, {}).get("is_correct", False)

                # If both students got the same wrong answer, this is suspicious
                if not is_correct1 and not is_correct2:
                    suspicious_matches += 1
                    flagged_questions.append(int(question_id))

        # Calculate similarity score
        similarity_score = (
            (matching_answers / total_questions * 100) if total_questions > 0 else 0
        )

        # Determine if this is suspicious based on:
        # 1. High overall similarity (>80%)
        # 2. High proportion of matching wrong answers
        suspicious_ratio = (
            (suspicious_matches / matching_answers * 100) if matching_answers > 0 else 0
        )

        is_suspicious = (
            similarity_score >= 80  # High overall similarity
            and suspicious_ratio >= 60  # At least 60% of matches are wrong answers
            and matching_answers >= 3  # At least 3 matching answers
        )

        # Determine severity
        if similarity_score >= 90 and suspicious_ratio >= 80:
            severity = "high"
        elif similarity_score >= 80 and suspicious_ratio >= 60:
            severity = "medium"
        else:
            severity = "low"

        return {
            "similarity_score": similarity_score,
            "suspicious_ratio": suspicious_ratio,
            "matching_answers": matching_answers,
            "suspicious_matches": suspicious_matches,
            "total_questions": total_questions,
            "is_suspicious": is_suspicious,
            "severity": severity,
            "flagged_questions": flagged_questions,
        }

    def get(self, request):
        try:
            # Get instructor's courses
            instructor_courses = CourseInstructor.objects.filter(
                user=request.user, accepted=True
            ).values_list("course_id", flat=True)

            # Get all exams for instructor's courses
            exams = Exam.objects.filter(course_id__in=instructor_courses)

            # Run similarity detection for each exam
            all_flags = []
            for exam in exams:
                flags = self._detect_similarities(exam)
                all_flags.extend(flags)

            # Get existing flags and combine with newly detected ones
            existing_flags = SimilarityFlag.objects.filter(
                exam__course_id__in=instructor_courses
            ).select_related("exam", "student1", "student2")

            # Format flags data
            formatted_flags = []
            for flag in existing_flags:
                # Determine severity type based on similarity score
                if flag.similarity_score >= 90:
                    severity_type = "high"
                elif flag.similarity_score >= 75:
                    severity_type = "medium"
                else:
                    severity_type = "low"

                formatted_flags.append(
                    {
                        "id": flag.id,
                        "type": severity_type,
                        "course": flag.exam.course.code,
                        "exam": flag.exam.title,
                        "studentPair": {
                            "student1": {
                                "name": flag.student1.name,
                                "id": str(flag.student1.id),
                            },
                            "student2": {
                                "name": flag.student2.name,
                                "id": str(flag.student2.id),
                            },
                        },
                        "similarityScore": float(flag.similarity_score),
                        "flaggedQuestions": flag.flagged_questions or [],
                        "dateDetected": flag.created_at.isoformat(),
                        "status": flag.status,
                        "reviewer": (
                            flag.reviewer.get_full_name() if flag.reviewer else None
                        ),
                        "notes": flag.notes,
                    }
                )

            # Get summary statistics
            total_flags = len(formatted_flags)
            active_flags = len([f for f in formatted_flags if f["status"] == "pending"])
            high_risk_flags = len(
                [
                    f
                    for f in formatted_flags
                    if f["type"] == "high" and f["status"] == "pending"
                ]
            )

            return Response(
                {
                    "flags": formatted_flags,
                    "totalFlags": total_flags,
                    "activeFlags": active_flags,
                    "highRiskFlags": high_risk_flags,
                }
            )

        except Exception as e:
            print(f"Error in SimilarityFlagsView: {str(e)}")
            import traceback

            traceback.print_exc()

            # Return empty data instead of error to prevent frontend crashes
            return Response(
                {"flags": [], "totalFlags": 0, "activeFlags": 0, "highRiskFlags": 0}
            )

    def post(self, request):
        """Update a similarity flag status"""
        flag_id = request.data.get("flag_id")
        new_status = request.data.get("status")
        notes = request.data.get("notes", "")

        if not flag_id:
            return Response(
                {"error": "flag_id is required"}, status=status.HTTP_400_BAD_REQUEST
            )

        if new_status not in ["reviewed", "dismissed", "confirmed"]:
            return Response(
                {"error": "Invalid status"}, status=status.HTTP_400_BAD_REQUEST
            )

        try:
            flag = SimilarityFlag.objects.get(id=flag_id)

            # Verify user has permission
            if not CourseInstructor.objects.filter(
                course=flag.exam.course, user=request.user, accepted=True
            ).exists():
                return Response(
                    {"error": "You don't have permission to update this flag"},
                    status=status.HTTP_403_FORBIDDEN,
                )

            # Update flag
            flag.status = new_status
            flag.notes = notes
            flag.reviewer = request.user
            flag.save()

            return Response(
                {
                    "success": True,
                    "message": f"Flag updated to {new_status}",
                    "flag": {
                        "id": flag.id,
                        "status": flag.status,
                        "reviewer": flag.reviewer.get_full_name(),
                        "notes": flag.notes,
                    },
                }
            )

        except SimilarityFlag.DoesNotExist:
            return Response(
                {"error": "Flag not found"}, status=status.HTTP_404_NOT_FOUND
            )


class YearOverYearTrendsView(APIView):
    """Get year-over-year performance trends (FR7.2)"""

    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user
        timeframe = request.query_params.get("timeframe", "1year")

        # Get instructor's courses
        instructor_courses = CourseInstructor.objects.filter(
            user=user, accepted=True
        ).values_list("course_id", flat=True)

        # Determine date range based on timeframe
        end_date = timezone.now()

        if timeframe == "1week":
            # Show the last 7 days including today
            start_date = end_date - timedelta(days=6)  # 6 days ago + today = 7 days
            interval_days = 1  # Daily intervals for week view
        elif timeframe == "1month":
            start_date = end_date - timedelta(days=30)
            interval_days = 1  # Daily intervals for month view to show all days
        elif timeframe == "1year":
            start_date = end_date - timedelta(days=365)
            interval_days = 30  # Monthly intervals for year view
        else:  # 'all'
            start_date = end_date - timedelta(days=1825)  # 5 years
            interval_days = 120  # Quarterly intervals for all time

        # Get results within timeframe
        results = ExamResult.objects.filter(
            exam__course__in=instructor_courses,
            submitted_at__gte=start_date,
            submitted_at__lte=end_date,
        )

        # Group by appropriate intervals
        trends = []
        current_date = start_date

        # Process data for each interval
        while current_date <= end_date:
            next_date = current_date + timedelta(days=interval_days)

            # Get results for this specific interval
            interval_results = results.filter(
                submitted_at__gte=current_date, submitted_at__lt=next_date
            )

            # Calculate average for this interval
            if interval_results.exists():
                avg_score = interval_results.aggregate(avg=Avg("score"))["avg"] or 0
                count = interval_results.count()
            else:
                # For shorter timeframes, include days with no data as 0 or null
                if timeframe in ["1week", "1month"]:
                    avg_score = None  # Show as null to indicate no data for this day
                    count = 0
                else:
                    # For longer timeframes, skip periods with no data
                    current_date = next_date
                    continue

            # Format label based on timeframe
            if timeframe == "1week":
                label = current_date.strftime("%a %m/%d")  # Mon 08/04
            elif timeframe == "1month":
                label = current_date.strftime("%m/%d")  # 08/04
            elif timeframe == "1year":
                label = current_date.strftime("%b %Y")  # Dec 2024
            else:  # all
                # Determine term for quarterly view
                month = current_date.month
                if month in [1, 2, 3, 4]:
                    term = "W1"
                elif month in [5, 6, 7, 8]:
                    term = "S1"
                else:
                    term = "W2"
                label = f"{term} {current_date.year}"

            trends.append(
                {
                    "year": current_date.year,
                    "term": label,
                    "average": (
                        round(float(avg_score), 1) if avg_score is not None else None
                    ),
                    "count": count,
                    "date": current_date.isoformat(),  # Include actual date for debugging
                }
            )

            current_date = next_date

        # For month view, if we have data, ensure we show the full month
        if timeframe == "1month" and trends:
            # Get the earliest date with data
            first_data_date = None
            for trend in trends:
                if trend["count"] > 0:
                    first_data_date = timezone.datetime.fromisoformat(
                        trend["date"].replace("+00:00", "")
                    ).replace(tzinfo=timezone.utc)
                    break

            if first_data_date:
                # Show from beginning of month if we have any data
                month_start = first_data_date.replace(day=1)
                month_end = (month_start + timedelta(days=32)).replace(
                    day=1
                ) - timedelta(days=1)

                # Regenerate trends for the full month
                trends = []
                current_date = month_start

                while current_date <= min(month_end, end_date):
                    interval_results = results.filter(
                        submitted_at__date=current_date.date()
                    )

                    if interval_results.exists():
                        avg_score = (
                            interval_results.aggregate(avg=Avg("score"))["avg"] or 0
                        )
                        count = interval_results.count()
                    else:
                        avg_score = None
                        count = 0

                    trends.append(
                        {
                            "year": current_date.year,
                            "term": current_date.strftime("%b %d"),
                            "average": (
                                round(float(avg_score), 1)
                                if avg_score is not None
                                else None
                            ),
                            "count": count,
                        }
                    )

                    current_date += timedelta(days=1)

        # Return the response
        return Response(
            {"trends": trends, "timeframe": timeframe, "totalDataPoints": len(trends)}
        )


class CourseStatisticsView(APIView):
    """Get detailed statistics for a specific course (UR2.18)"""

    permission_classes = [IsAuthenticated]

    def get(self, request, course_code):
        user = request.user

        # Get course
        course = get_object_or_404(Course, code=course_code)

        # Check permissions
        if not CourseInstructor.objects.filter(
            course=course, user=user, accepted=True
        ).exists():
            return Response(
                {"error": "You don't have permission to view this course"},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Get all exam results for this course
        results = ExamResult.objects.filter(exam__course=course).select_related(
            "exam", "student"
        )

        # Calculate historical data by term
        historical_data = []

        # Group results by exam/term
        exams = Exam.objects.filter(course=course).order_by("created_at")

        for exam in exams:
            exam_results = results.filter(exam=exam)

            if exam_results.exists():
                avg_score = exam_results.aggregate(avg=Avg("score"))["avg"] or 0
                student_count = exam_results.count()

                # Parse term from course term or exam creation date
                term_parts = course.term.split()
                if len(term_parts) >= 2:
                    term = term_parts[0]
                    year = int(term_parts[1])
                else:
                    # Fallback to exam creation date
                    year = exam.created_at.year
                    month = exam.created_at.month
                    if month <= 4:
                        term = "W1"
                    elif month <= 8:
                        term = "S1"
                    else:
                        term = "W2"

                # Determine semester
                semester = "Winter" if term.startswith("W") else "Summer"

                historical_data.append(
                    {
                        "year": year,
                        "term": term,
                        "semester": semester,
                        "average": round(avg_score, 1),
                        "studentCount": student_count,
                        "professor": course.instructor or "Unknown",
                        "ta": "N/A",  # Would need to track TAs separately
                    }
                )

        # Calculate aggregates
        all_scores = list(results.values_list("score", flat=True))
        if all_scores:
            average_all_years = np.mean(all_scores)

            # Last 5 years
            recent_results = results.filter(
                submitted_at__gte=timezone.now() - timedelta(days=1825)
            )
            recent_scores = list(recent_results.values_list("score", flat=True))
            average_last_5_years = (
                np.mean(recent_scores) if recent_scores else average_all_years
            )

            max_section_average = (
                max(h["average"] for h in historical_data) if historical_data else 0
            )
            min_section_average = (
                min(h["average"] for h in historical_data) if historical_data else 0
            )
        else:
            average_all_years = 0
            average_last_5_years = 0
            max_section_average = 0
            min_section_average = 0

        return Response(
            {
                "courseCode": course.code,
                "courseName": course.name,
                "totalSections": len(historical_data),
                "averageAllYears": round(average_all_years, 1),
                "averageLast5Years": round(average_last_5_years, 1),
                "maxSectionAverage": round(max_section_average, 1),
                "minSectionAverage": round(min_section_average, 1),
                "historicalData": historical_data,
            }
        )


class AnalyticsCoursesView(APIView):
    """Get all courses for analytics with proper formatting"""

    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user

        # Get courses where user is an instructor
        instructor_courses = CourseInstructor.objects.filter(
            user=user, accepted=True
        ).values_list("course_id", flat=True)

        courses = Course.objects.filter(id__in=instructor_courses)

        # Format for analytics
        course_data = []
        for course in courses:
            course_data.append(
                {
                    "id": course.id,
                    "code": course.code,
                    "title": course.name,
                    "term": course.term,
                    "description": course.description,
                }
            )

        return Response({"courses": course_data})


class CompareCourseView(APIView):
    """Compare performance across courses (UR2.18)"""

    permission_classes = [IsAuthenticated]

    def post(self, request):
        course_ids = request.data.get("course_ids", [])

        if len(course_ids) < 2:
            return Response(
                {"error": "At least two course IDs required"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Get courses
        courses = Course.objects.filter(id__in=course_ids)

        # Check permissions
        for course in courses:
            if not CourseInstructor.objects.filter(
                course=course, user=request.user, accepted=True
            ).exists():
                return Response(
                    {
                        "error": f"You don't have permission to view course {course.code}"
                    },
                    status=status.HTTP_403_FORBIDDEN,
                )

        comparison_data = []

        for course in courses:
            # Get metrics for each course
            results = ExamResult.objects.filter(exam__course=course)

            avg_score = results.aggregate(avg=Avg("score"))["avg"] or 0
            student_count = Student.objects.filter(
                course=course, is_active=True
            ).count()
            exam_count = Exam.objects.filter(course=course).count()

            # Calculate pass rate (assuming 60% is passing)
            pass_count = results.filter(score__gte=60).count()
            total_count = results.count()
            pass_rate = (pass_count / total_count * 100) if total_count > 0 else 0

            comparison_data.append(
                {
                    "course": {
                        "id": course.id,
                        "code": course.code,
                        "title": course.name,
                    },
                    "metrics": {
                        "averageScore": round(avg_score, 1),
                        "studentCount": student_count,
                        "examCount": exam_count,
                        "passRate": round(pass_rate, 1),
                    },
                }
            )

        return Response(comparison_data)


class StudentReportView(APIView):
    """Generate detailed student report"""

    permission_classes = [IsAuthenticated]

    def get(self, request, course_id, student_id):
        print(f"=== STUDENT REPORT VIEW ENTRY ===")
        print(f"Method: {request.method}")
        print(f"Path: {request.path}")
        print(f"User: {request.user}")
        print(f"Course ID: {course_id}, Student ID: {student_id}")

        try:
            from .helpers import DataFormatter, StatisticsCalculator

            print(f"=== STUDENT REPORT REQUEST ===")
            print(f"User: {request.user}")
            print(f"Authenticated: {request.user.is_authenticated}")
            print(f"Course ID: {course_id}")
            print(f"Student ID: {student_id}")

            course = get_object_or_404(Course, id=course_id)
            student = get_object_or_404(Student, id=student_id, course=course)

            print(f"Found course: {course.code} - {course.name}")
            print(f"Found student: {student.name}")

            # Check permissions
            instructor_perms = CourseInstructor.objects.filter(
                course=course, user=request.user, accepted=True
            ).exists()

            print(f"Instructor permissions: {instructor_perms}")

            if not instructor_perms:
                print("Permission denied: Not an instructor for this course")
                return Response(
                    {"error": "You don't have permission to view this course"},
                    status=status.HTTP_403_FORBIDDEN,
                )

            # Use helper classes to get performance data
            student_performance = StatisticsCalculator.get_student_performance(
                student_id, course
            )
            course_performance = StatisticsCalculator.get_course_performance(course)

            print(f"Student performance: {student_performance['stats']}")
            print(f"Class statistics: {course_performance['stats']}")

            # Format the report data using helper
            report_data = DataFormatter.format_student_report_data(
                student, course, student_performance, course_performance
            )

            print(f"=== RETURNING REPORT DATA ===")
            print(f"Report data keys: {list(report_data.keys())}")
            print(f"Exam results count: {len(report_data['examResults'])}")
            print(f"Recommendations count: {len(report_data['recommendations'])}")

            return Response(report_data)

        except Http404 as e:
            print(f"=== 404 NOT FOUND ===")
            print(f"Error: {str(e)}")
            return Response(
                {"error": "Student not found"}, status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            print(f"=== EXCEPTION IN STUDENT REPORT ===")
            print(f"Error: {str(e)}")
            import traceback

            traceback.print_exc()
            return Response(
                {"error": f"Server error: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )


# Alias for the courses view to match the URL pattern
CoursesView = AnalyticsCoursesView
TrendsView = YearOverYearTrendsView
SearchCoursesView = AnalyticsCoursesView  # Use the same view for search


class StudentReportExportView(APIView):
    """Export student report as PDF or DOCX"""

    permission_classes = [IsAuthenticated]

    def get(self, request, course_id=None, student_id=None, format=None):
        return self._export_report(request, course_id, student_id, format)

    def post(self, request, course_id=None, student_id=None, format=None):
        return self._export_report(request, course_id, student_id, format)

    def _export_report(self, request, course_id=None, student_id=None, format=None):
        try:
            print(f"=== STUDENT REPORT EXPORT REQUEST ===")
            print(f"User: {request.user}")
            print(f"Course ID: {course_id}")
            print(f"Student ID: {student_id}")

            # Check URL parameter first, then query parameter, then request body, default to pdf
            export_format = (
                format
                or request.query_params.get("format")
                or request.data.get("format", "pdf")
            ).lower()

            print(f"Export format: {export_format}")

            if export_format not in ["pdf", "docx", "csv"]:
                return Response(
                    {"error": "Invalid format. Use 'pdf', 'docx', or 'csv'."},
                    status=status.HTTP_400_BAD_REQUEST,
                )

            # Get and validate objects
            try:
                course = get_object_or_404(Course, id=course_id)
                print(f"Found course: {course.code} - {course.name}")
            except Exception as e:
                print(f"Course not found: {e}")
                return Response(
                    {"error": "Course not found"}, status=status.HTTP_404_NOT_FOUND
                )

            try:
                student = get_object_or_404(Student, id=student_id, course=course)
                print(f"Found student: {student.name}")
            except Exception as e:
                print(f"Student not found: {e}")
                return Response(
                    {"error": "Student not found in this course"},
                    status=status.HTTP_404_NOT_FOUND,
                )

            # Check permissions
            instructor_perms = CourseInstructor.objects.filter(
                course=course, user=request.user, accepted=True
            )

            if not instructor_perms.exists():
                return Response(
                    {
                        "error": "You don't have permission to export reports for this course"
                    },
                    status=status.HTTP_403_FORBIDDEN,
                )

            # Get exam results for this student in this course
            results = (
                ExamResult.objects.filter(student=student, exam__course=course)
                .select_related("exam", "variant")
                .order_by("exam__created_at")
            )

            if not results.exists():
                return Response(
                    {"error": "No exam results found for this student in this course"},
                    status=status.HTTP_404_NOT_FOUND,
                )

            # Import the export generators
            from .helpers import (
                CSVExportGenerator,
                DOCXReportGenerator,
                PDFReportGenerator,
            )

            # Generate the appropriate export
            if export_format == "csv":
                generator = CSVExportGenerator()
                response = generator.generate_student_report_csv(
                    student, course, results
                )
            elif export_format == "docx":
                generator = DOCXReportGenerator()
                response = generator.generate_student_report(student, course, results)
            else:  # pdf
                generator = PDFReportGenerator()
                response = generator.generate_student_report(student, course, results)

            print(f"Export generated successfully: {export_format}")
            return response

        except Exception as e:
            print(f"Export error: {str(e)}")
            import traceback

            traceback.print_exc()
            return Response(
                {"error": f"Failed to generate export: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )


class StudentReportExportDOCXView(StudentReportExportView):
    """Export student report as DOCX specifically"""

    def get(self, request, course_id=None, student_id=None):
        # Override to force DOCX format
        return self._export_report(request, course_id, student_id, format="docx")


class StudentReportExportCSVView(StudentReportExportView):
    """Export student report as CSV specifically"""

    def get(self, request, course_id=None, student_id=None):
        # Override to force CSV format
        return self._export_report(request, course_id, student_id, format="csv")


class DebugExportView(APIView):
    """Debug view to test URL routing"""

    permission_classes = []  # No authentication required for debug

    def get(self, request, course_id=None, student_id=None):
        try:
            # Test the actual export logic without authentication
            from courses.models import Course, Student
            from results.models import ExamResult

            from .helpers import (
                CSVExportGenerator,
                DOCXReportGenerator,
                PDFReportGenerator,
            )

            # Get the first course and student for testing
            course = Course.objects.first()
            student = Student.objects.first() if course else None

            if not course or not student:
                return Response(
                    {
                        "error": "No test data available",
                        "courses": Course.objects.count(),
                        "students": Student.objects.count(),
                    }
                )

            # Get exam results
            results = ExamResult.objects.filter(
                student=student, exam__course=course
            ).select_related("exam", "variant")

            # Test CSV generation
            try:
                csv_generator = CSVExportGenerator()
                csv_response = csv_generator.generate_student_report_csv(
                    student, course, results
                )
                csv_status = "OK"
            except Exception as e:
                csv_status = f"Error: {str(e)}"

            # Test PDF generation
            try:
                pdf_generator = PDFReportGenerator()
                pdf_response = pdf_generator.generate_student_report(
                    student, course, results
                )
                pdf_status = "OK"
            except Exception as e:
                pdf_status = f"Error: {str(e)}"

            # Test DOCX generation
            try:
                docx_generator = DOCXReportGenerator()
                docx_response = docx_generator.generate_student_report(
                    student, course, results
                )
                docx_status = "OK"
            except Exception as e:
                docx_status = f"Error: {str(e)}"

            return Response(
                {
                    "status": "Debug test completed",
                    "csv": csv_status,
                    "pdf": pdf_status,
                    "docx": docx_status,
                    "course": f"{course.code} - {course.name}" if course else "None",
                    "student": student.name if student else "None",
                    "results_count": results.count(),
                }
            )

        except Exception as e:
            return Response({"error": f"Debug test failed: {str(e)}"}, status=500)


class DebugCSVExportView(APIView):
    """Debug endpoint for CSV export testing"""

    permission_classes = []

    def get(self, request):
        try:
            print("=== DEBUG CSV EXPORT TEST ===")
            from .helpers import CSVExportGenerator

            CSVExportGenerator()

            # Create a simple test CSV response
            import csv
            import io

            output = io.StringIO()
            writer = csv.writer(output)
            writer.writerow(["Test", "CSV", "Export"])
            writer.writerow(["Value1", "Value2", "Value3"])

            response = HttpResponse(output.getvalue(), content_type="text/csv")
            response["Content-Disposition"] = 'attachment; filename="test.csv"'
            return response
        except Exception as e:
            print(f"Debug CSV error: {str(e)}")
            import traceback

            traceback.print_exc()
            return Response({"error": str(e)}, status=500)


class DebugPDFExportView(APIView):
    """Debug endpoint for PDF export testing"""

    permission_classes = []

    def get(self, request):
        try:
            print("=== DEBUG PDF EXPORT TEST ===")

            # Test basic ReportLab functionality first
            import io

            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate

            print("PDF Generator: Creating simple test PDF")

            # Create a simple PDF buffer
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            story = []

            styles = getSampleStyleSheet()
            title = Paragraph("Test PDF Export", styles["Title"])
            story.append(title)

            content = Paragraph(
                "This is a test PDF generated successfully!", styles["Normal"]
            )
            story.append(content)

            doc.build(story)
            buffer.seek(0)

            # Return the PDF response
            response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
            response["Content-Disposition"] = 'attachment; filename="debug_test.pdf"'
            print("PDF Generator: Simple PDF created successfully")
            return response

        except Exception as e:
            print(f"Debug PDF error: {str(e)}")
            import traceback

            traceback.print_exc()
            return Response({"error": str(e)}, status=500)


class DebugDOCXExportView(APIView):
    """Debug endpoint for DOCX export testing"""

    permission_classes = []

    def get(self, request):
        try:
            print("=== DEBUG DOCX EXPORT TEST ===")

            # Test basic python-docx functionality first
            import io

            from docx import Document

            print("DOCX Generator: Creating simple test DOCX")

            # Create a simple DOCX document
            doc = Document()
            doc.add_heading("Test DOCX Export", 0)
            doc.add_paragraph("This is a test DOCX document generated successfully!")

            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Return the DOCX response
            response = HttpResponse(
                buffer.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = 'attachment; filename="debug_test.docx"'
            print("DOCX Generator: Simple DOCX created successfully")
            return response

        except Exception as e:
            print(f"Debug DOCX error: {str(e)}")
            import traceback

            traceback.print_exc()
            return Response({"error": str(e)}, status=500)


class CourseBulkExportView(APIView):
    """Export all students' data for a course in one file"""

    permission_classes = [IsAuthenticated]

    def get(self, request, course_id, file_format=None):
        print(
            f"=== HIT BULK EXPORT VIEW === course_id={course_id}, file_format={file_format}"
        )
        print(f"Query params: {request.query_params}")
        return self._export_bulk_report(request, course_id, file_format)

    def post(self, request, course_id, file_format=None):
        print(
            f"=== HIT BULK EXPORT VIEW POST === course_id={course_id}, file_format={file_format}"
        )
        return self._export_bulk_report(request, course_id, file_format)

    def _export_bulk_report(self, request, course_id, file_format=None):
        try:
            print("=== BULK EXPORT REQUEST ===")
            print(f"User: {request.user}")
            print(f"Course ID: {course_id}")

            # Check URL parameter first, then query parameter, then request body, default to csv
            export_format = (
                file_format
                or request.query_params.get("format")
                or request.data.get("format", "csv")
            ).lower()

            print(f"Export format: {export_format}")

            if export_format not in ["pdf", "docx", "csv"]:
                return Response(
                    {"error": "Invalid format. Use 'pdf', 'docx', or 'csv'."},
                    status=status.HTTP_400_BAD_REQUEST,
                )

            # Get and validate course
            try:
                course = get_object_or_404(Course, id=course_id)
                print(f"Found course: {course.code} - {course.name}")
            except Exception as e:
                print(f"Course not found: {e}")
                return Response(
                    {"error": "Course not found"}, status=status.HTTP_404_NOT_FOUND
                )

            # Check permissions
            instructor_perms = CourseInstructor.objects.filter(
                course=course, user=request.user, accepted=True
            )

            if not instructor_perms.exists():
                return Response(
                    {
                        "error": "You don't have permission to export reports for this course"
                    },
                    status=status.HTTP_403_FORBIDDEN,
                )

            # Get all students for this course
            students = Student.objects.filter(course=course)

            if not students.exists():
                return Response(
                    {"error": "No students found in this course"},
                    status=status.HTTP_404_NOT_FOUND,
                )

            # Import the export generators
            from .helpers import (
                CSVExportGenerator,
                DOCXReportGenerator,
                PDFReportGenerator,
            )

            # Generate the appropriate export
            if export_format == "csv":
                response = CSVExportGenerator.generate_bulk_student_reports_csv(
                    students, course
                )
            elif export_format == "docx":
                response = DOCXReportGenerator.generate_bulk_student_reports_docx(
                    course, students
                )
            else:  # pdf
                response = PDFReportGenerator.generate_bulk_student_reports_pdf(
                    course, students
                )

            print(f"Bulk export generated successfully: {export_format}")
            return response

        except Exception as e:
            print(f"Bulk export error: {str(e)}")
            import traceback

            traceback.print_exc()
            return Response(
                {"error": f"Failed to generate bulk export: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )
